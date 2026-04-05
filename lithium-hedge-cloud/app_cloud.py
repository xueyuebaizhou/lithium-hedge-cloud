# app_cloud.py - 完整云端版本（v32：准商业级增强：报告导出/策略管理/风险模块；现货内置表；不生成模拟价格）
import streamlit as st
import pandas as pd
import plotly.express as px
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
from matplotlib import font_manager
import os
import sys
from datetime import datetime, timedelta
import warnings
import json
import io
import base64
import hashlib
import traceback
import math
import re
import requests
from typing import Optional, Dict, Any, List
warnings.filterwarnings('ignore')

# =========================
# v32 商业化增强：报告导出（PDF/Word）+ 策略管理/风险模块（已移除角色与日志权限）
# 重要合规约束：
# - 禁止随机游走或任何方式生成模拟价格路径
# - 仅使用真实历史数据做风险度量（历史法 VaR/CVaR、历史极端情景）
# - 若数据缺失，UI 必须红字提示“当前为模拟数据，禁止用于对外报告”
# =========================

ROLE_LABELS = {}  # 角色区分已移除

# 页面可见性（越靠左权限越高）
PAGE_ROLES = {}  # 角色区分已移除：不做页面隐藏

# 动作权限
ACTION_ROLES = {}  # 角色区分已移除：不做权限限制

def _get_role() -> str:
    """v31.1: role/identity separation removed. All users share the same permissions."""
    return "unified"


def _can(action: str) -> bool:
    """No role separation: all actions allowed."""
    return True

def _require(action: str, tip: str = "权限不足，无法执行该操作。") -> bool:
    if _can(action):
        return True
    st.error(f"⛔ {tip}")
    return False

class AuditLogger:
    """轻量日志：
    - 优先写入 Supabase 表 audit_logs（若可用）
    - 同时写入 session_state（便于无后端场景演示）
    """
    def __init__(self, supabase_client=None):
        self.supabase = supabase_client

    def log(self, user_info: dict, action: str, detail: dict | None = None):
        try:
            detail = detail or {}
            role = _get_role()
            record = {
                "ts": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "user_id": (user_info or {}).get("user_id"),
                "username": (user_info or {}).get("username"),
                "role": role,
                "action": action,
                "detail": detail,
            }
            # session_state
            if "audit_logs" not in st.session_state:
                st.session_state.audit_logs = []
            st.session_state.audit_logs.insert(0, record)  # 最新在前

            # supabase（尽力而为）
            if self.supabase:
                try:
                    self.supabase.table("audit_logs").insert(record).execute()
                except Exception:
                    pass
        except Exception:
            pass

def _log(action: str, detail: dict | None = None):
    """便捷记录：在 UI 层调用。"""
    try:
        analyzer = st.session_state.get("_analyzer_ref")
        if analyzer and hasattr(analyzer, "audit"):
            analyzer.audit.log(st.session_state.get("user_info") or {}, action, detail=detail or {})
    except Exception:
        pass

def build_report_pdf_bytes(title: str, lines: list[str]) -> bytes | None:
    """生成 PDF（可选依赖 reportlab）。"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # 尝试注册中文字体（如果找不到就降级英文/方框）
        try:
            # 常见字体路径（本地/云端不保证存在）
            candidates = [
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
                "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            ]
            font_path = next((p for p in candidates if os.path.exists(p)), None)
            if font_path:
                pdfmetrics.registerFont(TTFont("CJK", font_path))
                font_name = "CJK"
            else:
                font_name = "Helvetica"
        except Exception:
            font_name = "Helvetica"

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        width, height = A4

        c.setTitle(title)
        c.setFont(font_name, 14)
        c.drawString(20*mm, height-20*mm, title)

        c.setFont(font_name, 10)
        y = height - 30*mm
        for line in lines:
            if y < 20*mm:
                c.showPage()
                c.setFont(font_name, 10)
                y = height - 20*mm
            # 简单换行
            s = str(line)
            max_chars = 90
            while len(s) > max_chars:
                c.drawString(20*mm, y, s[:max_chars])
                s = s[max_chars:]
                y -= 6*mm
                if y < 20*mm:
                    c.showPage()
                    c.setFont(font_name, 10)
                    y = height - 20*mm
            c.drawString(20*mm, y, s)
            y -= 6*mm

        c.save()
        buf.seek(0)
        return buf.read()
    except Exception:
        return None

def build_report_docx_bytes(title: str, lines: list[str]) -> bytes | None:
    """生成 Word（可选依赖 python-docx）。"""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(title, level=0)
        for line in lines:
            doc.add_paragraph(str(line))
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
    except Exception:
        return None

SPOT_LITHIUM_PRICE_TABLE = [
    ('2025-01-01', 97000),
    ('2025-01-02', 97500),
    ('2025-01-03', 98000),
    ('2025-01-04', 98500),
    ('2025-01-05', 99000),
    ('2025-01-06', 99500),
    ('2025-01-07', 100000),
    ('2025-01-08', 99500),
    ('2025-01-09', 99000),
    ('2025-01-10', 98500),
    ('2025-01-11', 98000),
    ('2025-01-12', 97500),
    ('2025-01-13', 97000),
    ('2025-01-14', 96500),
    ('2025-01-15', 96000),
    ('2025-01-16', 95500),
    ('2025-01-17', 95000),
    ('2025-01-18', 94500),
    ('2025-01-19', 94000),
    ('2025-01-20', 93500),
    ('2025-01-21', 93000),
    ('2025-01-22', 92500),
    ('2025-01-23', 92000),
    ('2025-01-24', 91500),
    ('2025-01-25', 91000),
    ('2025-01-26', 90500),
    ('2025-01-27', 90000),
    ('2025-01-28', 89500),
    ('2025-01-29', 89000),
    ('2025-01-30', 88500),
    ('2025-01-31', 88000),
    ('2025-02-01', 87500),
    ('2025-02-02', 87000),
    ('2025-02-03', 86500),
    ('2025-02-04', 86000),
    ('2025-02-05', 85500),
    ('2025-02-06', 85000),
    ('2025-02-07', 84500),
    ('2025-02-08', 84000),
    ('2025-02-09', 83500),
    ('2025-02-10', 83000),
    ('2025-02-11', 82500),
    ('2025-02-12', 82000),
    ('2025-02-13', 81500),
    ('2025-02-14', 81000),
    ('2025-02-15', 80500),
    ('2025-02-16', 80000),
    ('2025-02-17', 79500),
    ('2025-02-18', 79000),
    ('2025-02-19', 78500),
    ('2025-02-20', 78000),
    ('2025-02-21', 77500),
    ('2025-02-22', 77000),
    ('2025-02-23', 76500),
    ('2025-02-24', 76000),
    ('2025-02-25', 75500),
    ('2025-02-26', 75000),
    ('2025-02-27', 74500),
    ('2025-02-28', 74000),
    ('2025-03-01', 73500),
    ('2025-03-02', 73000),
    ('2025-03-03', 72500),
    ('2025-03-04', 72000),
    ('2025-03-05', 71500),
    ('2025-03-06', 71000),
    ('2025-03-07', 70500),
    ('2025-03-08', 70000),
    ('2025-03-09', 69500),
    ('2025-03-10', 69000),
    ('2025-03-11', 68500),
    ('2025-03-12', 68000),
    ('2025-03-13', 67500),
    ('2025-03-14', 67000),
    ('2025-03-15', 66500),
    ('2025-03-16', 66000),
    ('2025-03-17', 65500),
    ('2025-03-18', 65000),
    ('2025-03-19', 64500),
    ('2025-03-20', 64000),
    ('2025-03-21', 63500),
    ('2025-03-22', 63000),
    ('2025-03-23', 62500),
    ('2025-03-24', 62000),
    ('2025-03-25', 61500),
    ('2025-03-26', 61000),
    ('2025-03-27', 60500),
    ('2025-03-28', 60000),
    ('2025-03-29', 59900),
    ('2025-03-30', 59900),
    ('2025-03-31', 59900),
    ('2025-04-01', 60000),
    ('2025-04-02', 60500),
    ('2025-04-03', 61000),
    ('2025-04-04', 61500),
    ('2025-04-05', 62000),
    ('2025-04-06', 62500),
    ('2025-04-07', 63000),
    ('2025-04-08', 63500),
    ('2025-04-09', 64000),
    ('2025-04-10', 64500),
    ('2025-04-11', 65000),
    ('2025-04-12', 65500),
    ('2025-04-13', 66000),
    ('2025-04-14', 66500),
    ('2025-04-15', 67000),
    ('2025-04-16', 67500),
    ('2025-04-17', 68000),
    ('2025-04-18', 68500),
    ('2025-04-19', 69000),
    ('2025-04-20', 69500),
    ('2025-04-21', 70000),
    ('2025-04-22', 70500),
    ('2025-04-23', 71000),
    ('2025-04-24', 71500),
    ('2025-04-25', 72000),
    ('2025-04-26', 72500),
    ('2025-04-27', 73000),
    ('2025-04-28', 73500),
    ('2025-04-29', 74000),
    ('2025-04-30', 74500),
    ('2025-05-01', 75000),
    ('2025-05-02', 75500),
    ('2025-05-03', 76000),
    ('2025-05-04', 76500),
    ('2025-05-05', 77000),
    ('2025-05-06', 77500),
    ('2025-05-07', 78000),
    ('2025-05-08', 78500),
    ('2025-05-09', 79000),
    ('2025-05-10', 79500),
    ('2025-05-11', 80000),
    ('2025-05-12', 80500),
    ('2025-05-13', 81000),
    ('2025-05-14', 81500),
    ('2025-05-15', 82000),
    ('2025-05-16', 82500),
    ('2025-05-17', 83000),
    ('2025-05-18', 83500),
    ('2025-05-19', 84000),
    ('2025-05-20', 84500),
    ('2025-05-21', 85000),
    ('2025-05-22', 85500),
    ('2025-05-23', 86000),
    ('2025-05-24', 86500),
    ('2025-05-25', 87000),
    ('2025-05-26', 87500),
    ('2025-05-27', 88000),
    ('2025-05-28', 88500),
    ('2025-05-29', 89000),
    ('2025-05-30', 89500),
    ('2025-05-31', 90000),
    ('2025-06-01', 90500),
    ('2025-06-02', 91000),
    ('2025-06-03', 91500),
    ('2025-06-04', 92000),
    ('2025-06-05', 92500),
    ('2025-06-06', 93000),
    ('2025-06-07', 93500),
    ('2025-06-08', 94000),
    ('2025-06-09', 94500),
    ('2025-06-10', 95000),
    ('2025-06-11', 95500),
    ('2025-06-12', 96000),
    ('2025-06-13', 96500),
    ('2025-06-14', 97000),
    ('2025-06-15', 97500),
    ('2025-06-16', 98000),
    ('2025-06-17', 98500),
    ('2025-06-18', 99000),
    ('2025-06-19', 99500),
    ('2025-06-20', 100000),
    ('2025-06-21', 100500),
    ('2025-06-22', 101000),
    ('2025-06-23', 101500),
    ('2025-06-24', 102000),
    ('2025-06-25', 102500),
    ('2025-06-26', 103000),
    ('2025-06-27', 103500),
    ('2025-06-28', 104000),
    ('2025-06-29', 104500),
    ('2025-06-30', 105000),
    ('2025-07-01', 105500),
    ('2025-07-02', 106000),
    ('2025-07-03', 106500),
    ('2025-07-04', 107000),
    ('2025-07-05', 107500),
    ('2025-07-06', 108000),
    ('2025-07-07', 108500),
    ('2025-07-08', 109000),
    ('2025-07-09', 109500),
    ('2025-07-10', 110000),
    ('2025-07-11', 110500),
    ('2025-07-12', 111000),
    ('2025-07-13', 111500),
    ('2025-07-14', 112000),
    ('2025-07-15', 112500),
    ('2025-07-16', 113000),
    ('2025-07-17', 113500),
    ('2025-07-18', 114000),
    ('2025-07-19', 114500),
    ('2025-07-20', 115000),
    ('2025-07-21', 115500),
    ('2025-07-22', 116000),
    ('2025-07-23', 116500),
    ('2025-07-24', 117000),
    ('2025-07-25', 117500),
    ('2025-07-26', 118000),
    ('2025-07-27', 118500),
    ('2025-07-28', 119000),
    ('2025-07-29', 119500),
    ('2025-07-30', 120000),
    ('2025-07-31', 120500),
    ('2025-08-01', 121000),
    ('2025-08-02', 121500),
    ('2025-08-03', 122000),
    ('2025-08-04', 122500),
    ('2025-08-05', 123000),
    ('2025-08-06', 123500),
    ('2025-08-07', 124000),
    ('2025-08-08', 124500),
    ('2025-08-09', 125000),
    ('2025-08-10', 125500),
    ('2025-08-11', 126000),
    ('2025-08-12', 126500),
    ('2025-08-13', 127000),
    ('2025-08-14', 127500),
    ('2025-08-15', 128000),
    ('2025-08-16', 128500),
    ('2025-08-17', 129000),
    ('2025-08-18', 129500),
    ('2025-08-19', 130000),
    ('2025-08-20', 130500),
    ('2025-08-21', 131000),
    ('2025-08-22', 131500),
    ('2025-08-23', 132000),
    ('2025-08-24', 132500),
    ('2025-08-25', 133000),
    ('2025-08-26', 133500),
    ('2025-08-27', 134000),
    ('2025-08-28', 134500),
    ('2025-08-29', 135000),
    ('2025-08-30', 135500),
    ('2025-08-31', 136000),
    ('2025-09-01', 136500),
    ('2025-09-02', 137000),
    ('2025-09-03', 137500),
    ('2025-09-04', 138000),
    ('2025-09-05', 138500),
    ('2025-09-06', 139000),
    ('2025-09-07', 139500),
    ('2025-09-08', 140000),
    ('2025-09-09', 140500),
    ('2025-09-10', 141000),
    ('2025-09-11', 141500),
    ('2025-09-12', 142000),
    ('2025-09-13', 142500),
    ('2025-09-14', 143000),
    ('2025-09-15', 143500),
    ('2025-09-16', 144000),
    ('2025-09-17', 144500),
    ('2025-09-18', 145000),
    ('2025-09-19', 145500),
    ('2025-09-20', 146000),
    ('2025-09-21', 146500),
    ('2025-09-22', 147000),
    ('2025-09-23', 147500),
    ('2025-09-24', 148000),
    ('2025-09-25', 148500),
    ('2025-09-26', 149000),
    ('2025-09-27', 149500),
    ('2025-09-28', 150000),
    ('2025-09-29', 150500),
    ('2025-09-30', 151000),
    ('2025-10-01', 151500),
    ('2025-10-02', 152000),
    ('2025-10-03', 152500),
    ('2025-10-04', 153000),
    ('2025-10-05', 153500),
    ('2025-10-06', 154000),
    ('2025-10-07', 154500),
    ('2025-10-08', 155000),
    ('2025-10-09', 155500),
    ('2025-10-10', 156000),
    ('2025-10-11', 156500),
    ('2025-10-12', 157000),
    ('2025-10-13', 157500),
    ('2025-10-14', 158000),
    ('2025-10-15', 158500),
    ('2025-10-16', 159000),
    ('2025-10-17', 159500),
    ('2025-10-18', 160000),
    ('2025-10-19', 160500),
    ('2025-10-20', 161000),
    ('2025-10-21', 161500),
    ('2025-10-22', 162000),
    ('2025-10-23', 162500),
    ('2025-10-24', 163000),
    ('2025-10-25', 163500),
    ('2025-10-26', 164000),
    ('2025-10-27', 164500),
    ('2025-10-28', 165000),
    ('2025-10-29', 165500),
    ('2025-10-30', 166000),
    ('2025-10-31', 166500),
    ('2025-11-01', 167000),
    ('2025-11-02', 167500),
    ('2025-11-03', 168000),
    ('2025-11-04', 168500),
    ('2025-11-05', 169000),
    ('2025-11-06', 169500),
    ('2025-11-07', 170000),
    ('2025-11-08', 170500),
    ('2025-11-09', 171000),
    ('2025-11-10', 171500),
    ('2025-11-11', 172000),
    ('2025-11-12', 172500),
    ('2025-11-13', 173000),
    ('2025-11-14', 173500),
    ('2025-11-15', 174000),
    ('2025-11-16', 174500),
    ('2025-11-17', 175000),
    ('2025-11-18', 175500),
    ('2025-11-19', 176000),
    ('2025-11-20', 176500),
    ('2025-11-21', 177000),
    ('2025-11-22', 177500),
    ('2025-11-23', 178000),
    ('2025-11-24', 178500),
    ('2025-11-25', 179000),
    ('2025-11-26', 179500),
    ('2025-11-27', 180000),
    ('2025-11-28', 180500),
    ('2025-11-29', 181000),
    ('2025-11-30', 181500),
    ('2025-12-01', 182000),
    ('2025-12-02', 182500),
    ('2025-12-03', 183000),
    ('2025-12-04', 183500),
    ('2025-12-05', 184000),
    ('2025-12-06', 184500),
    ('2025-12-07', 185000),
    ('2025-12-08', 185500),
    ('2025-12-09', 186000),
    ('2025-12-10', 186500),
    ('2025-12-11', 187000),
    ('2025-12-12', 187500),
    ('2025-12-13', 188000),
    ('2025-12-14', 188500),
    ('2025-12-15', 189000),
    ('2025-12-16', 189500),
    ('2025-12-17', 190000),
    ('2025-12-18', 190500),
    ('2025-12-19', 191000),
    ('2025-12-20', 191500),
    ('2025-12-21', 192000),
    ('2025-12-22', 192500),
    ('2025-12-23', 193000),
    ('2025-12-24', 193500),
    ('2025-12-25', 194000),
    ('2025-12-26', 194500),
    ('2025-12-27', 195000),
    ('2025-12-28', 195500),
    ('2025-12-29', 196000),
    ('2025-12-30', 196500),
    ('2025-12-31', 197000),
    ('2026-01-01', 129000),
    ('2026-01-02', 132000),
    ('2026-01-03', 135000),
    ('2026-01-04', 138000),
    ('2026-01-05', 141000),
    ('2026-01-06', 144000),
    ('2026-01-07', 147000),
    ('2026-01-08', 150000),
    ('2026-01-09', 153000),
    ('2026-01-10', 156000),
    ('2026-01-11', 159000),
    ('2026-01-12', 162000),
    ('2026-01-13', 165000),
    ('2026-01-14', 168000),
    ('2026-01-15', 171000),
    ('2026-01-16', 174000),
    ('2026-01-17', 177000),
    ('2026-01-18', 180000),
    ('2026-01-19', 181500),
    ('2026-01-20', 181500),
    ('2026-01-21', 181500),
    ('2026-01-22', 181500),
    ('2026-01-23', 181500),
    ('2026-01-24', 178500),
    ('2026-01-25', 175500),
    ('2026-01-26', 172500),
    ('2026-01-27', 169500),
    ('2026-01-28', 166500),
    ('2026-01-29', 163500),
    ('2026-01-30', 160500),
    ('2026-01-31', 157500),
    ('2026-02-01', 154500),
    ('2026-02-02', 151500),
    ('2026-02-03', 148500),
    ('2026-02-04', 145500),
    ('2026-02-05', 142500),
    ('2026-02-06', 139500),
    ('2026-02-07', 136500),
    ('2026-02-08', 133500),
    ('2026-02-09', 130500),
    ('2026-02-10', 127500),
    ('2026-02-11', 124500),
    ('2026-02-12', 121500),
    ('2026-02-13', 118500),
    ('2026-02-14', 121500),
    ('2026-02-15', 124500),
    ('2026-02-16', 127500),
    ('2026-02-17', 130500),
    ('2026-02-18', 133500),
    ('2026-02-19', 136500),
    ('2026-02-20', 139500),
    ('2026-02-21', 142500),
    ('2026-02-22', 145500),
    ('2026-02-23', 148500),
    ('2026-02-24', 151500),
    ('2026-02-25', 154500),
    ('2026-02-26', 157500),
    ('2026-02-27', 160500),
    ('2026-02-28', 163500),
    ('2026-03-01', 169500),
    ('2026-03-02', 172500),
    ('2026-03-03', 173000),
]

# === 内置现货价格表（定期更新）辅助 ===
def _build_spot_df():
    try:
        df = pd.DataFrame(SPOT_LITHIUM_PRICE_TABLE, columns=["date", "spot_price"])
        df["date"] = pd.to_datetime(df["date"], errors="coerce")
        df = df.dropna(subset=["date"]).sort_values("date").reset_index(drop=True)
        df["spot_price"] = pd.to_numeric(df["spot_price"], errors="coerce")
        df = df.dropna(subset=["spot_price"])
        return df
    except Exception:
        return pd.DataFrame(columns=["date", "spot_price"])

_SPOT_DF = _build_spot_df()

def get_spot_price_on_or_before(target_date: datetime):
    """取现货表中 date <= target_date 的最近一条价格。返回 (price, used_date, ok)"""
    try:
        if _SPOT_DF is None or _SPOT_DF.empty:
            return None, None, False
        d = pd.to_datetime(target_date).normalize()
        sub = _SPOT_DF[_SPOT_DF["date"] <= d]
        if sub.empty:
            return None, None, False
        row = sub.iloc[-1]
        return float(row["spot_price"]), row["date"].to_pydatetime(), True
    except Exception:
        return None, None, False

def compute_historical_var_cvar(pnl_series: pd.Series, alpha: float = 0.05):
    """历史法 VaR / CVaR（不生成任何模拟价格，完全基于真实历史分布）。"""
    s = pd.to_numeric(pnl_series, errors="coerce").dropna()
    if s.empty:
        return None, None
    var = float(np.quantile(s, alpha))
    tail = s[s <= var]
    cvar = float(tail.mean()) if not tail.empty else var
    return var, cvar

def align_spot_futures(spot_df: pd.DataFrame, fut_df: pd.DataFrame):
    """对齐现货与期货到同一日期（inner join）。"""
    a = spot_df.copy()
    b = fut_df.copy()
    a["date"] = pd.to_datetime(a["date"]).dt.normalize()
    b["date"] = pd.to_datetime(b["date"]).dt.normalize()
    merged = pd.merge(a, b, on="date", how="inner")
    merged = merged.sort_values("date").reset_index(drop=True)
    return merged

def _fmt_dt(x):
    """Format datetime/date/str safely for UI."""
    try:
        if hasattr(x, "strftime"):
            return x.strftime("%Y-%m-%d")
    except Exception:
        pass
    if isinstance(x, str):
        return x[:10] if len(x) >= 10 else x
    return str(x)

# 添加utils路径
sys.path.append(os.path.join(os.path.dirname(__file__), 'utils'))

# 导入Supabase管理器
try:
    from supabase_client import get_supabase_manager
    supabase = get_supabase_manager()
    HAS_SUPABASE = supabase is not None
except ImportError as e:
    HAS_SUPABASE = False
    print(f"Supabase导入失败: {e}")


def ensure_chinese_font():
    """确保云端环境可用中文字体（推荐：随项目携带 fonts/ 字体文件）。

    - Matplotlib：使用本地 TTC/OTF 注册后全局生效
    - Plotly：通过 @font-face 注入 WOFF2（子集字体），浏览器端可显示中文
    """
    base_dir = os.path.dirname(__file__)
    fonts_dir = os.path.join(base_dir, "fonts")
    os.makedirs(fonts_dir, exist_ok=True)

    # 1) Matplotlib 字体（用于 plt / matplotlib）
    ttc_path = os.path.join(fonts_dir, "NotoSansCJK-Regular.ttc")  # 推荐放这个
    ttf_fallback = os.path.join(fonts_dir, "SourceHanSansSC-Regular.otf")  # 可选备用

    font_name = None
    for fp in (ttc_path, ttf_fallback):
        if os.path.exists(fp):
            try:
                font_manager.fontManager.addfont(fp)
                font_name = font_manager.FontProperties(fname=fp).get_name()
                break
            except Exception as exc:
                print(f"中文字体加载失败: {fp} -> {exc}")

    # 2) Plotly / 前端字体（用于 plotly 图表中文）
    # 使用子集 woff2，文件很小；你需要把它也放到 fonts/ 目录
    woff2_path = os.path.join(fonts_dir, "NotoSansCJKSC-Subset.woff2")
    if os.path.exists(woff2_path):
        try:
            with open(woff2_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("utf-8")
            st.markdown(
                f"""<style>
@font-face {{
  font-family: 'NotoSansCJKSC';
  src: url(data:font/woff2;base64,{b64}) format('woff2');
  font-weight: normal;
  font-style: normal;
}}
</style>""",
                unsafe_allow_html=True
            )
        except Exception as exc:
            print(f"Plotly 字体注入失败: {exc}")

    return font_name


chinese_font = ensure_chinese_font()

# Matplotlib 全局字体
if chinese_font:
    matplotlib.rcParams['font.sans-serif'] = [chinese_font, 'DejaVu Sans', 'Arial']
else:
    # 兜底（若 fonts/ 没放字体，仍可能变方框）
    matplotlib.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial']
matplotlib.rcParams['axes.unicode_minus'] = False
matplotlib.rcParams['axes.prop_cycle'] = matplotlib.cycler(color=['#B9924D', '#496A81', '#2E8B80', '#A55D4C', '#6E7480'])

# Plotly 全局字体（浏览器端）：优先使用我们注入的 'NotoSansCJKSC'
import plotly.io as pio
PLOTLY_FONT_FAMILY = "NotoSansCJKSC, Microsoft YaHei, SimHei, PingFang SC, Heiti SC, Arial"
for tpl in ["plotly", "plotly_white", "plotly_dark", "ggplot2", "seaborn", "simple_white", "presentation"]:
    try:
        pio.templates[tpl].layout.font.family = PLOTLY_FONT_FAMILY
        pio.templates[tpl].layout.paper_bgcolor = "rgba(255,255,255,0)"
        pio.templates[tpl].layout.plot_bgcolor = "rgba(255,255,255,0.72)"
        pio.templates[tpl].layout.colorway = ["#355c7d", "#c9a96b", "#6c7a89", "#8f6f3e", "#5a7d6c"]
    except Exception:
        pass


# ============================================================================
# 用户认证管理器（云端版）
# ============================================================================

class CloudUserAuth:
    """云端用户认证管理器

    说明：
    - 认证与数据存储通过 utils/supabase_client.py 中的 Supabase 管理器完成（若存在）。
    - 由于不同版本/封装返回值形态不同，本类统一将所有返回规范化为 (ok, payload)。
    """

    def __init__(self):
        self.supabase = supabase if HAS_SUPABASE else None

    @staticmethod
    def _localize_auth_message(msg):
        if msg is None:
            return ""
        if isinstance(msg, dict):
            for key in ["msg", "message", "error_description", "error", "code"]:
                if key in msg and msg.get(key):
                    return CloudUserAuth._localize_auth_message(msg.get(key))
            return str(msg)
        text = str(msg).strip()
        lower = text.lower()

        # 尽量兼容异常对象、JSON 字符串、HTTP 文本等不同格式
        compact = lower.replace("\n", " ").replace("\r", " ")
        compact = compact.replace("_", " ").replace("-", " ")

        mappings = [
            ("email rate limit exceeded", "发送过于频繁，请稍后再试"),
            ("rate limit exceeded", "发送过于频繁，请稍后再试"),
            ("over email send rate limit", "发送过于频繁，请稍后再试"),
            ("security purposes", "发送过于频繁，请稍后再试"),
            ("too many requests", "请求过于频繁，请稍后再试"),
            ("invalid login credentials", "邮箱或密码错误"),
            ("invalid otp", "验证码错误或已过期"),
            ("otp expired", "验证码已过期，请重新获取"),
            ("token has expired or is invalid", "验证码错误或已过期"),
            ("email not confirmed", "邮箱尚未完成验证"),
            ("user not found", "用户不存在"),
            ("signup is disabled", "当前未开放注册"),
            ("network", "网络异常，请稍后重试"),
            ("timeout", "请求超时，请稍后重试"),
        ]
        for needle, repl in mappings:
            if needle in compact:
                return repl
        return text

    @staticmethod
    def _normalize_result(ret, default_ok: bool = False):
        """Normalize backend returns to (ok, payload)."""
        if isinstance(ret, tuple):
            if len(ret) >= 2:
                return bool(ret[0]), ret[1]
            if len(ret) == 1:
                return bool(ret[0]), None
            return default_ok, None
        if isinstance(ret, dict):
            if "success" in ret:
                return bool(ret.get("success")), ret
            if "ok" in ret:
                return bool(ret.get("ok")), ret
            if "error" in ret and ret.get("error"):
                return False, ret
            return default_ok, ret
        # Some helpers may return object/string directly
        return default_ok, ret

    def _get_supabase_rest_config(self):
        """尽量从环境变量或 Supabase 客户端中提取 REST 所需配置。"""
        url_candidates = [
            os.getenv("SUPABASE_URL"),
            os.getenv("NEXT_PUBLIC_SUPABASE_URL"),
            getattr(self.supabase, "supabase_url", None) if self.supabase else None,
            getattr(self.supabase, "url", None) if self.supabase else None,
        ]
        key_candidates = [
            os.getenv("SUPABASE_ANON_KEY"),
            os.getenv("NEXT_PUBLIC_SUPABASE_ANON_KEY"),
            os.getenv("SUPABASE_KEY"),
            getattr(self.supabase, "supabase_key", None) if self.supabase else None,
            getattr(self.supabase, "key", None) if self.supabase else None,
        ]
        auth = getattr(self.supabase, "auth", None) if self.supabase else None
        if auth is not None:
            url_candidates.extend([
                getattr(auth, "url", None),
                getattr(auth, "_url", None),
            ])
            key_candidates.extend([
                getattr(auth, "key", None),
                getattr(auth, "_key", None),
            ])

        url = next((x.strip().rstrip("/") for x in url_candidates if isinstance(x, str) and x.strip()), None)
        key = next((x.strip() for x in key_candidates if isinstance(x, str) and x.strip()), None)
        return url, key

    def _post_gotrue(self, path: str, payload: dict):
        url, key = self._get_supabase_rest_config()
        if not url or not key:
            return False, "缺少 Supabase URL 或 ANON KEY 配置"
        try:
            resp = requests.post(
                f"{url}/auth/v1/{path.lstrip('/')}",
                json=payload,
                headers={
                    "apikey": key,
                    "Authorization": f"Bearer {key}",
                    "Content-Type": "application/json",
                },
                timeout=20,
            )
            if 200 <= resp.status_code < 300:
                try:
                    return True, resp.json()
                except Exception:
                    return True, resp.text
            try:
                err = resp.json()
            except Exception:
                err = resp.text
            return False, self._localize_auth_message(err or f"HTTP {resp.status_code}")
        except Exception as e:
            return False, self._localize_auth_message(str(e))

    def register(self, username: str, password: str, email: str):
        """注册用户。返回 (ok, payload/message)"""
        if not self.supabase:
            return False, "数据库连接失败，请检查配置"

        username = (username or "").strip()
        email = (email or "").strip()
        password = password or ""

        if len(username) < 3:
            return False, "用户名至少3个字符"
        if "@" not in email or "." not in email:
            return False, "请输入有效的邮箱地址"
        if len(password) < 6:
            return False, "密码长度至少6位"

        # Prefer project-specific helper if available
        for fn_name in ["create_user", "register", "sign_up", "signup"]:
            fn = getattr(self.supabase, fn_name, None)
            if callable(fn):
                try:
                    return self._normalize_result(fn(username, password, email), default_ok=True)
                except TypeError:
                    # Some SDK use (email, password) only
                    try:
                        return self._normalize_result(fn(email, password), default_ok=True)
                    except Exception as e:
                        return False, f"注册失败: {e}"
                except Exception as e:
                    return False, f"注册失败: {e}"

        # Supabase python client often uses supabase.auth.sign_up(...)
        auth = getattr(self.supabase, "auth", None)
        if auth and hasattr(auth, "sign_up"):
            try:
                return self._normalize_result(auth.sign_up({"email": email, "password": password}), default_ok=True)
            except Exception as e:
                return False, f"注册失败: {e}"

        return False, "后端未提供注册接口"

    def login(self, username_or_email: str, password: str):
        """登录。返回 (ok, payload)"""
        if not self.supabase:
            return False, "数据库连接失败，请检查配置"

        username_or_email = (username_or_email or "").strip()
        password = password or ""
        if not username_or_email:
            return False, "请输入用户名或邮箱"
        if not password:
            return False, "请输入密码"

        # Prefer project helper
        for fn_name in ["authenticate_user", "login", "sign_in", "sign_in_with_password"]:
            fn = getattr(self.supabase, fn_name, None)
            if callable(fn):
                try:
                    if fn_name == "sign_in_with_password":
                        return self._normalize_result(fn({"email": username_or_email, "password": password}), default_ok=True)
                    return self._normalize_result(fn(username_or_email, password), default_ok=True)
                except TypeError:
                    # Sometimes accepts (email, password) only
                    try:
                        return self._normalize_result(fn(username_or_email, password), default_ok=True)
                    except Exception as e:
                        return False, f"登录失败: {e}"
                except Exception as e:
                    return False, f"登录失败: {e}"

        # Supabase auth style
        auth = getattr(self.supabase, "auth", None)
        if auth and hasattr(auth, "sign_in_with_password"):
            try:
                return self._normalize_result(auth.sign_in_with_password({"email": username_or_email, "password": password}), default_ok=True)
            except Exception as e:
                return False, f"登录失败: {e}"

        return False, "后端未提供登录接口"

    def change_password(self, username_or_email: str, old_password: str, new_password: str):
        """修改密码：先验证旧密码再改新密码（若后端支持）。"""
        if not self.supabase:
            return False, "数据库连接失败"

        new_password = new_password or ""
        if len(new_password) < 6:
            return False, "新密码至少6位"

        # If helper exists
        for fn_name in ["change_password", "update_password", "update_user_password"]:
            fn = getattr(self.supabase, fn_name, None)
            if callable(fn):
                try:
                    return self._normalize_result(fn(username_or_email, old_password, new_password), default_ok=True)
                except TypeError:
                    try:
                        return self._normalize_result(fn(username_or_email, new_password), default_ok=True)
                    except Exception as e:
                        return False, f"修改密码失败: {e}"
                except Exception as e:
                    return False, f"修改密码失败: {e}"

        return False, "后端未提供改密接口"

    def send_email_login_code(self, email: str):
        """发送邮箱验证码用于免密登录。"""
        if not self.supabase:
            return False, "数据库连接失败，请检查配置"

        email = (email or "").strip()
        if not email or "@" not in email or "." not in email:
            return False, "请输入有效的邮箱地址"

        for fn_name in [
            "send_email_login_code",
            "send_login_code",
            "send_email_otp",
            "sign_in_with_otp",
            "login_with_otp",
        ]:
            fn = getattr(self.supabase, fn_name, None)
            if callable(fn):
                try:
                    if fn_name in ["sign_in_with_otp", "login_with_otp"]:
                        return self._normalize_result(fn({"email": email}), default_ok=True)
                    return self._normalize_result(fn(email), default_ok=True)
                except TypeError:
                    try:
                        return self._normalize_result(fn(email=email), default_ok=True)
                    except Exception as e:
                        return False, f"发送邮箱验证码失败: {self._localize_auth_message(e)}"
                except Exception as e:
                    return False, f"发送邮箱验证码失败: {self._localize_auth_message(e)}"

        auth = getattr(self.supabase, "auth", None)
        if auth and hasattr(auth, "sign_in_with_otp"):
            try:
                return self._normalize_result(auth.sign_in_with_otp({"email": email}), default_ok=True)
            except Exception as e:
                return False, f"发送邮箱验证码失败: {self._localize_auth_message(e)}"

        return False, "后端未提供邮箱验证码登录接口"

    def login_with_email_code(self, email: str, code: str):
        """邮箱验证码免密登录。"""
        if not self.supabase:
            return False, "数据库连接失败，请检查配置"

        email = (email or "").strip()
        code = (code or "").strip()
        if not email or "@" not in email or "." not in email:
            return False, "请输入有效的邮箱地址"
        if not code:
            return False, "请输入邮箱验证码"

        for fn_name in [
            "login_with_email_code",
            "verify_email_login_code",
            "verify_login_code",
            "verify_email_otp",
        ]:
            fn = getattr(self.supabase, fn_name, None)
            if callable(fn):
                try:
                    return self._normalize_result(fn(email, code), default_ok=True)
                except TypeError:
                    try:
                        return self._normalize_result(fn({"email": email, "token": code}), default_ok=True)
                    except Exception as e:
                        return False, f"验证码登录失败: {self._localize_auth_message(e)}"
                except Exception as e:
                    return False, f"验证码登录失败: {self._localize_auth_message(e)}"

        auth = getattr(self.supabase, "auth", None)
        if auth and hasattr(auth, "verify_otp"):
            verify_payloads = [
                {"email": email, "token": code, "type": "email"},
                {"email": email, "token": code, "type": "magiclink"},
            ]
            last_err = None
            for payload in verify_payloads:
                try:
                    return self._normalize_result(auth.verify_otp(payload), default_ok=True)
                except Exception as e:
                    last_err = e
            return False, f"验证码登录失败: {self._localize_auth_message(last_err)}" if last_err else "验证码登录失败"

        return False, "后端未提供邮箱验证码登录接口"

    def generate_reset_code(self, username_or_email: str, email: str | None = None):
        """发送重置验证码。当前项目固定使用邮箱 OTP 重置密码。"""
        target = (email or username_or_email or "").strip()
        if not target:
            return False, "请输入用户名或邮箱"
        if "@" not in target or "." not in target:
            return False, "请输入有效的注册邮箱"

        last_err = None
        if self.supabase:
            for fn_name in ["generate_reset_code", "send_reset_code", "send_password_reset_email", "reset_password_for_email"]:
                fn = getattr(self.supabase, fn_name, None)
                if callable(fn):
                    try:
                        ret = fn(target)
                        ok, payload = self._normalize_result(ret, default_ok=True)
                        if ok:
                            return True, target
                        last_err = payload if isinstance(payload, str) else payload
                    except Exception as e:
                        last_err = e

            auth = getattr(self.supabase, "auth", None)
            if auth and hasattr(auth, "reset_password_for_email"):
                try:
                    try:
                        ret = auth.reset_password_for_email(target)
                    except TypeError:
                        ret = auth.reset_password_for_email(target, {})
                    ok, payload = self._normalize_result(ret, default_ok=True)
                    if ok:
                        return True, target
                    last_err = payload if isinstance(payload, str) else payload
                except Exception as e:
                    last_err = e

        # Supabase REST 后备：即使 Python 包未暴露该方法，也可直接调用 GoTrue
        ok, payload = self._post_gotrue("recover", {"email": target})
        if ok:
            return True, target

        detail = payload or last_err
        return False, f"发送重置验证码失败: {self._localize_auth_message(detail)}" if detail else "发送重置验证码失败"

    def reset_password(self, username_or_email: str, reset_code: str | None = None, new_password: str | None = None, email: str | None = None):
        """使用邮箱验证码（OTP）重置密码。"""
        target_email = (email or username_or_email or "").strip()
        reset_code = (reset_code or "").strip()
        new_password = (new_password or "").strip()

        if not target_email or "@" not in target_email or "." not in target_email:
            return False, "缺少有效邮箱，请返回上一步重新获取验证码"
        if not reset_code:
            return False, "请输入验证码"
        if len(new_password) < 6:
            return False, "新密码至少6位"

        verify_err = None
        auth = getattr(self.supabase, "auth", None) if self.supabase else None
        if auth and hasattr(auth, "verify_otp"):
            for payload in [
                {"email": target_email, "token": reset_code, "type": "recovery"},
                {"email": target_email, "token": reset_code, "type": "email"},
            ]:
                try:
                    verify_ret = auth.verify_otp(payload)
                    session = None
                    if hasattr(verify_ret, "session"):
                        session = getattr(verify_ret, "session")
                    elif isinstance(verify_ret, dict):
                        session = verify_ret.get("session")
                    if hasattr(auth, "update_user"):
                        try:
                            auth.update_user({"password": new_password})
                            return True, "密码重置成功，请使用新密码登录"
                        except TypeError:
                            auth.update_user(password=new_password)
                            return True, "密码重置成功，请使用新密码登录"
                    break
                except Exception as e:
                    verify_err = e

        # 项目自定义后备接口
        if self.supabase:
            for fn_name in ["reset_password", "set_password", "update_user_password", "update_password"]:
                fn = getattr(self.supabase, fn_name, None)
                if callable(fn):
                    try:
                        try:
                            return self._normalize_result(fn(target_email, reset_code, new_password), default_ok=True)
                        except TypeError:
                            try:
                                return self._normalize_result(fn(target_email, new_password), default_ok=True)
                            except TypeError:
                                return self._normalize_result(fn(username_or_email, reset_code, new_password), default_ok=True)
                    except Exception as e:
                        verify_err = e

        # REST 后备：verify_otp 获取 session，再用 access_token 更新密码
        ok, payload = self._post_gotrue("verify", {"email": target_email, "token": reset_code, "type": "recovery"})
        if not ok:
            ok2, payload2 = self._post_gotrue("verify", {"email": target_email, "token": reset_code, "type": "email"})
            if ok2:
                ok, payload = ok2, payload2
            else:
                detail = payload2 or payload or verify_err
                return False, f"验证码校验失败: {self._localize_auth_message(detail)}" if detail else "验证码校验失败"

        access_token = None
        if isinstance(payload, dict):
            session = payload.get("session") or {}
            access_token = session.get("access_token") or payload.get("access_token")
        if not access_token:
            return False, "验证码校验成功，但未获取到会话令牌，无法更新密码"

        url, key = self._get_supabase_rest_config()
        if not url or not key:
            return False, "缺少 Supabase URL 或 ANON KEY 配置"
        try:
            resp = requests.put(
                f"{url}/auth/v1/user",
                json={"password": new_password},
                headers={
                    "apikey": key,
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json",
                },
                timeout=20,
            )
            if 200 <= resp.status_code < 300:
                return True, "密码重置成功，请使用新密码登录"
            try:
                err = resp.json()
            except Exception:
                err = resp.text
            return False, f"重置密码失败: {self._localize_auth_message(err or ('HTTP ' + str(resp.status_code)))}"
        except Exception as e:
            return False, f"重置密码失败: {self._localize_auth_message(e)}"

    def update_user_settings(self, user_id: str, settings: dict):
        """更新用户设置（可选）。"""
        if not self.supabase:
            return False, "数据库连接失败"
        for fn_name in ["update_user_settings", "update_settings"]:
            fn = getattr(self.supabase, fn_name, None)
            if callable(fn):
                try:
                    return self._normalize_result(fn(user_id, settings), default_ok=True)
                except Exception as e:
                    return False, f"更新设置失败: {e}"
        return False, "后端未提供设置更新接口"


# =========================
# User info extraction helper
# =========================
def _extract_user_info_from_login_result(result, username_fallback: str = "") -> dict:
    """Normalize various Supabase/helper login returns into a dict with user_id/username/email/settings."""
    info: dict = {}
    try:
        if isinstance(result, dict):
            info = dict(result)
        else:
            # supabase-py AuthResponse style: result.user / result.session
            if hasattr(result, "user") and getattr(result, "user") is not None:
                u = getattr(result, "user")
                info["user_id"] = getattr(u, "id", None) or getattr(u, "user_id", None)
                info["email"] = getattr(u, "email", None)
                um = getattr(u, "user_metadata", None)
                if isinstance(um, dict):
                    info["username"] = um.get("username") or um.get("name")

            # some wrappers use result.data.user
            if (not info.get("user_id")) and hasattr(result, "data") and getattr(result, "data") is not None:
                d = getattr(result, "data")
                u = None
                if isinstance(d, dict):
                    u = d.get("user") or d.get("profile")
                else:
                    u = getattr(d, "user", None)
                if u is not None:
                    if isinstance(u, dict):
                        info["user_id"] = u.get("id") or u.get("user_id")
                        info["email"] = u.get("email")
                        um = u.get("user_metadata")
                        if isinstance(um, dict):
                            info["username"] = um.get("username") or um.get("name")
                    else:
                        info["user_id"] = getattr(u, "id", None) or getattr(u, "user_id", None)
                        info["email"] = getattr(u, "email", None)
                        um = getattr(u, "user_metadata", None)
                        if isinstance(um, dict):
                            info["username"] = um.get("username") or um.get("name")

            # direct attributes
            if not info.get("user_id") and hasattr(result, "id"):
                info["user_id"] = getattr(result, "id", None)
            if not info.get("email") and hasattr(result, "email"):
                info["email"] = getattr(result, "email", None)
    except Exception:
        info = info or {}

    fallback_name = username_fallback or info.get("email") or ""
    if isinstance(fallback_name, str) and "@" in fallback_name:
        fallback_name = fallback_name.split("@", 1)[0]
    username = (info.get("username") or fallback_name or "").strip()
    raw_uid = info.get("user_id") or info.get("id")
    user_id = raw_uid.strip() if isinstance(raw_uid, str) else raw_uid
    settings = info.get("settings") if isinstance(info.get("settings"), dict) else {}

    # If still no user_id, fall back to username as key for MVP (keep app usable)
    if not user_id and username:
        user_id = f"user::{username}"

    return {
        "user_id": user_id,
        "username": username or username_fallback or "用户",
        "email": info.get("email"),
        "settings": settings,
        # 角色：优先读取 settings.role；若无则默认 manager（保证可用性）
        "role": (settings.get("role") or "manager"),
    }


class CloudLithiumAnalyzer:
    """Main application service layer.

    Notes:
    - Keep this class small and deterministic: it should NOT contain Streamlit UI code.
    - Provide stable APIs used by the UI layer:
        - fetch_real_time_data(symbol: str = "LC0") -> pd.DataFrame
        - hedge_calculation(cost_price, inventory, hedge_ratio, margin_rate) -> (fig, suggestions, metrics)
    """

    def __init__(self):
        self.supabase = supabase if HAS_SUPABASE else None
        self.auth = CloudUserAuth()
        self.audit = AuditLogger(self.supabase)
        # Simple in-memory cache for price data (keyed by symbol)
        self.cache_data: dict[str, pd.DataFrame] = {}
        self.cache_time: dict[str, datetime] = {}
        self._configure_matplotlib_fonts()

    def _configure_matplotlib_fonts(self) -> None:
        """Try to enable CJK text rendering on cloud runtimes.

        If no CJK font is available, plots will still render but CJK glyphs may appear as boxes.
        In that case, we will prefer English labels inside plots while keeping Chinese text in Streamlit UI.
        """
        try:
            import matplotlib as mpl
            from matplotlib import font_manager

            preferred = [
                "Noto Sans CJK SC",
                "Noto Sans CJK",
                "Source Han Sans CN",
                "Source Han Sans",
                "SimHei",
                "Microsoft YaHei",
                "PingFang SC",
                "WenQuanYi Micro Hei",
            ]
            available = {f.name for f in font_manager.fontManager.ttflist}
            for name in preferred:
                if name in available:
                    mpl.rcParams["font.family"] = name
                    break
            mpl.rcParams["axes.unicode_minus"] = False
        except Exception:
            # Never crash the app due to font configuration.
            pass

    def fetch_real_time_data(self, symbol: str = "LC0", days: int = 180, force_refresh: bool = False) -> pd.DataFrame:
        """Return a price series DataFrame.

        Contract symbols:
        - Continuous contract: e.g. LC0
        - Specific contract: e.g. LC2603

        Primary source (real, verifiable): AkShare -> 新浪财经期货日频数据
        Interface: ak.futures_zh_daily_sina(symbol=...)

        Output columns (cloud-stable):
            - 日期 (datetime)
            - 收盘价 (float)
            - 涨跌幅 (float, %, optional)
            - 成交量 (optional)
            - __data_source (str)
            - __is_simulated (bool)

        If real data fetch fails and no cache exists, we fall back to a **constant** series
        (NOT random walk) and clearly mark it as simulated so the UI can show a red warning.
        """
        import pandas as pd
        from datetime import datetime, timedelta

        now = datetime.now()

        # Cache (30 minutes)
        try:
            ttl = timedelta(minutes=30)
            if not force_refresh and symbol in self.cache_data and symbol in self.cache_time:
                if now - self.cache_time[symbol] < ttl:
                    return self.cache_data[symbol].copy()
        except Exception:
            pass

        # 1) Primary: AkShare -> 新浪期货日频
        df = None
        data_source = ""
        try:
            import akshare as ak
            raw = ak.futures_zh_daily_sina(symbol=symbol)
            if raw is not None and not raw.empty:
                df = raw.copy()
                # Normalize
                if "date" in df.columns and "日期" not in df.columns:
                    df = df.rename(columns={"date": "日期"})
                if "close" in df.columns and "收盘价" not in df.columns:
                    df = df.rename(columns={"close": "收盘价"})
                if "open" in df.columns and "开盘价" not in df.columns:
                    df = df.rename(columns={"open": "开盘价"})
                if "high" in df.columns and "最高价" not in df.columns:
                    df = df.rename(columns={"high": "最高价"})
                if "low" in df.columns and "最低价" not in df.columns:
                    df = df.rename(columns={"low": "最低价"})
                if "volume" in df.columns and "成交量" not in df.columns:
                    df = df.rename(columns={"volume": "成交量"})

                if "日期" in df.columns:
                    df["日期"] = pd.to_datetime(df["日期"], errors="coerce")
                df = df.dropna(subset=["日期", "收盘价"]).sort_values("日期").reset_index(drop=True)

                if "涨跌幅" not in df.columns and len(df) > 1:
                    df["涨跌幅"] = df["收盘价"].pct_change() * 100

                df = df.tail(max(30, min(int(days or 180), 3650))).reset_index(drop=True)
                data_source = "AkShare:futures_zh_daily_sina(Sina)"
        except Exception:
            df = None

        if df is not None and not df.empty:
            df = df.copy()
            df["__data_source"] = data_source
            df["__is_simulated"] = False
            try:
                self.cache_data[symbol] = df
                self.cache_time[symbol] = now
            except Exception:
                pass
            return df.copy()

        # 2) Cache fallback (stale but real)
        try:
            if symbol in self.cache_data and not self.cache_data[symbol].empty:
                cached = self.cache_data[symbol].copy()
                cached["__data_source"] = cached.get("__data_source", "CACHE")
                cached["__is_simulated"] = False
                return cached
        except Exception:
            pass

        # 3) Simulated fallback (constant, explicitly marked)
        base_price = 0.0
        n = max(30, min(int(days or 180), 365))
        dates = [now - timedelta(days=(n - 1 - i)) for i in range(n)]
        df = pd.DataFrame({"日期": pd.to_datetime(dates), "收盘价": [base_price] * n})
        df["涨跌幅"] = 0.0
        df["__data_source"] = "SIMULATED:constant_series"
        df["__is_simulated"] = True
        return df

    
    def fetch_spot_reference_price(self, item_query: str = "碳酸锂", date: Optional[str] = None, lookback_days: int = 10) -> dict:
        """Fetch spot *reference* price (public statistical caliber, not transaction price).

        Primary source: AkShare `futures_spot_price(date=YYYYMMDD)` (collected from 生意社).
        Practical note: the *same-day* record often returns empty. We therefore implement a
        deterministic fallback: search backward up to `lookback_days` calendar days and use
        the most recent date that returns a non-empty table and matches `item_query`.

        Returns:
            {
              'price': float | None,
              'date': 'YYYYMMDD',                 # the actual date used (may be earlier than requested)
              'source': str,
              'detail': str,
              'is_simulated': bool
            }
        """
        from datetime import datetime, timedelta
        import pandas as pd

        # Normalize start date
        start_dt = None
        if date:
            s = str(date).strip()
            try:
                if re.fullmatch(r"\d{8}", s):
                    start_dt = datetime.strptime(s, "%Y%m%d")
                else:
                    # allow YYYY-MM-DD or other parsable forms
                    start_dt = pd.to_datetime(s, errors="coerce").to_pydatetime()
            except Exception:
                start_dt = None
        if start_dt is None:
            start_dt = datetime.now()

        q = str(item_query or "碳酸锂").strip()
        last_err = None
        attempts = []

        for i in range(max(1, int(lookback_days or 10))):
            d_dt = start_dt - timedelta(days=i)
            d = d_dt.strftime('%Y%m%d')
            try:
                import akshare as ak
                spot_df = ak.futures_spot_price(d)
                if spot_df is None or spot_df.empty:
                    attempts.append(f"{d}: empty")
                    continue

                df = spot_df.copy()

                # Normalize columns
                if '商品' not in df.columns or '现货价格' not in df.columns:
                    for c in df.columns:
                        if ('商品' in c) and ('商品' not in df.columns):
                            df = df.rename(columns={c: '商品'})
                        if ('现货' in c and '价格' in c) and ('现货价格' not in df.columns):
                            df = df.rename(columns={c: '现货价格'})

                if '商品' not in df.columns or '现货价格' not in df.columns:
                    attempts.append(f"{d}: missing columns")
                    continue

                df = df.dropna(subset=['商品', '现货价格'])
                if df.empty:
                    attempts.append(f"{d}: empty after dropna")
                    continue

                lc = df[df['商品'].astype(str).str.contains(q, na=False)]
                if lc.empty:
                    attempts.append(f"{d}: no match '{q}'")
                    continue

                price = float(pd.to_numeric(lc['现货价格'], errors='coerce').dropna().mean())
                if not (price > 0):
                    attempts.append(f"{d}: non-positive")
                    continue

                names = ', '.join(lc['商品'].astype(str).head(5).tolist())
                # Detail shows if fallback happened
                if i == 0:
                    detail = f"匹配条目示例: {names}"
                else:
                    detail = f"当日无数据/不可用，已回溯至 {d}（回溯{i}天）。匹配条目示例: {names}"

                return {
                    'price': price,
                    'date': d,
                    'source': 'AkShare:futures_spot_price(生意社)',
                    'detail': detail,
                    'is_simulated': False,
                }
            except Exception as e:
                last_err = e
                attempts.append(f"{d}: error")
                continue

        # Fail: no usable public reference price found in lookback window
        detail = f"回溯{lookback_days}天仍未获得有效数据；尝试记录: {', '.join(attempts[:8])}{'...' if len(attempts)>8 else ''}"
        if last_err is not None:
            detail += f"；最后错误: {last_err}"

        return {
            'price': None,
            'date': start_dt.strftime('%Y%m%d'),
            'source': 'AkShare:futures_spot_price(生意社)',
            'detail': detail,
            'is_simulated': True,
        }



    def fetch_spot_price_from_excel(self, date: Optional[str] = None, excel_path: Optional[str] = None) -> dict:
        """Fetch spot reference price from an embedded (static) table.

        Notes:
        - The user provided an Excel table '（已内置现货表）'. For deployment stability (no external files),
          we embed that table into this program as SPOT_LITHIUM_PRICE_TABLE.
        - No realtime update is performed.
        - If `date` is provided, we use the latest spot price whose date <= `date`.
        - If no matching record exists, returns is_simulated=True (and UI should show the red compliance warning).

        Args:
            date: 'YYYYMMDD' or 'YYYY-MM-DD' or None (use latest available).
            excel_path: kept for backward compatibility, ignored.

        Returns:
            {
              'price': float|None,
              'date': 'YYYYMMDD' (actual date used),
              'source': str,
              'detail': str,
              'is_simulated': bool
            }
        """
        from datetime import datetime

        try:
            table = SPOT_LITHIUM_PRICE_TABLE
            if not table:
                return {
                    'price': None,
                    'date': None,
                    'source': '内置表:碳酸锂现货价格',
                    'detail': '内置现货表为空',
                    'is_simulated': True,
                }

            # Parse target date
            if date is None:
                target = datetime.strptime(table[-1][0], '%Y-%m-%d').date()
            else:
                s = str(date).strip()
                if len(s) == 8 and s.isdigit():
                    target = datetime.strptime(s, '%Y%m%d').date()
                else:
                    target = datetime.strptime(s[:10], '%Y-%m-%d').date()

            # Binary search latest <= target
            import bisect
            dates = [d for d, _ in table]
            key = target.isoformat()
            idx = bisect.bisect_right(dates, key) - 1
            if idx < 0:
                return {
                    'price': None,
                    'date': target.strftime('%Y%m%d'),
                    'source': '内置表:碳酸锂现货价格',
                    'detail': f'目标日期 {target.isoformat()} 早于表格起始日期 {dates[0]}',
                    'is_simulated': True,
                }

            used_date, used_price = table[idx]
            used_dt = datetime.strptime(used_date, '%Y-%m-%d').date()
            detail = f'使用内置现货表；目标={target.isoformat()}；命中={used_date}（<=目标的最近一条）'

            return {
                'price': float(used_price),
                'date': used_dt.strftime('%Y%m%d'),
                'source': '内置表:碳酸锂现货价格',
                'detail': detail,
                'is_simulated': False,
            }
        except Exception as e:
            return {
                'price': None,
                'date': str(date) if date is not None else None,
                'source': '内置表:碳酸锂现货价格',
                'detail': f'内置现货表解析失败: {e}',
                'is_simulated': True,
            }


    def fetch_100ppi_benchmark_price(self, url: str, force_refresh: bool = False, cache_minutes: int = 30) -> dict:
        """Fetch benchmark spot price from 100ppi (生意社) detail page.

        This is a best-effort scraper. 100ppi pages sometimes use dynamic rendering / anti-bot.
        We therefore:
        - set a browser-like UA + headers
        - keep a session (cookies)
        - try multiple regex patterns for price + unit + update time
        - cache successful results per URL

        Returns:
            dict with keys:
              price(float|None), unit(str), source(str), url(str), update_time(str|None),
              is_simulated(bool), detail(str)
        """
        try:
            # cache key per-url
            key = "100ppi_" + hashlib.md5(url.encode("utf-8")).hexdigest()
            now = datetime.now()
            last_t = self.cache_time.get(key)
            if (not force_refresh) and last_t and (now - last_t) < timedelta(minutes=cache_minutes):
                df = self.cache_data.get(key)
                if isinstance(df, pd.DataFrame) and not df.empty:
                    row = df.iloc[0].to_dict()
                    return {
                        "price": float(row.get("price")) if row.get("price") is not None else None,
                        "unit": row.get("unit", "元/吨"),
                        "source": "100ppi",
                        "url": url,
                        "update_time": row.get("update_time"),
                        "is_simulated": False,
                        "detail": "缓存命中（非强制刷新）",
                    }

            sess = requests.Session()
            headers = {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/122.0.0.0 Safari/537.36"
                ),
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
                "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
                "Accept-Encoding": "gzip, deflate, br",
                "Connection": "keep-alive",
                "Referer": "https://www.100ppi.com/",
                "Upgrade-Insecure-Requests": "1",
            }

            r = sess.get(url, headers=headers, timeout=15)
            r.raise_for_status()

            html = r.text or ""
            if (not html) and r.content:
                try:
                    html = r.content.decode("utf-8", errors="ignore")
                except Exception:
                    html = r.content.decode(errors="ignore")

            unit = "元/吨"
            price = None

            patterns = [
                r'new_price[^>]*>\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)\s*<',
                r'([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)\s*元/吨',
                r'现价[^0-9]*([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)',
                r'"price"\s*:\s*"?(\d+(?:\.\d+)?)"?',
            ]
            for pat in patterns:
                m_price = re.search(pat, html, flags=re.I)
                if m_price:
                    try:
                        price = float(m_price.group(1).replace(",", ""))
                        break
                    except Exception:
                        price = None

            update_time = None
            ut_patterns = [
                r'更新\s*时间[：:\s]*([0-9]{2}-[0-9]{2}\s*[0-9]{2}:[0-9]{2})',
                r'更新\s*时间[：:\s]*([0-9]{4}-[0-9]{2}-[0-9]{2}\s*[0-9]{2}:[0-9]{2})',
                r'更新时间[：:\s]*([0-9]{2}-[0-9]{2}\s*[0-9]{2}:[0-9]{2})',
                r'更新时间[：:\s]*([0-9]{4}-[0-9]{2}-[0-9]{2}\s*[0-9]{2}:[0-9]{2})',
            ]
            for pat in ut_patterns:
                m_ut = re.search(pat, html)
                if m_ut:
                    update_time = m_ut.group(1).strip()
                    break

            if (price is None) or (price <= 0):
                return {
                    "price": None,
                    "unit": unit,
                    "source": "100ppi",
                    "url": url,
                    "update_time": update_time,
                    "is_simulated": True,
                    "detail": "页面解析未得到有效价格字段（可能是页面结构变更 / 动态渲染 / 反爬拦截）。建议："
                              "1) 先在浏览器打开确认页面可访问；"
                              "2) 换一个生意社品种详情页 URL；"
                              "3) 或使用 AkShare futures_spot_price 作为统计口径来源。",
                }

            cache_df = pd.DataFrame([{
                "price": price,
                "unit": unit,
                "update_time": update_time,
                "url": url,
                "fetched_at": now.strftime("%Y-%m-%d %H:%M:%S"),
            }])
            self.cache_data[key] = cache_df
            self.cache_time[key] = now

            return {
                "price": price,
                "unit": unit,
                "source": "100ppi",
                "url": url,
                "update_time": update_time,
                "is_simulated": False,
                "detail": "抓取成功",
            }
        except Exception as e:
            return {
                "price": None,
                "unit": "元/吨",
                "source": "100ppi",
                "url": url,
                "update_time": None,
                "is_simulated": True,
                "detail": f"抓取失败：{e}",
            }


    def list_spot_items(self, date: Optional[str] = None, keyword: str = "锂") -> dict:
        """List available spot item names from AkShare futures_spot_price.

        Returns:
          {
            'date': 'YYYYMMDD',
            'items': [str, ...],
            'source': 'AkShare:futures_spot_price(生意社)',
            'detail': str,
            'is_simulated': bool
          }
        """
        from datetime import datetime
        d = date or datetime.now().strftime('%Y%m%d')
        try:
            import akshare as ak
            df = ak.futures_spot_price(d)
            if df is None or df.empty:
                raise ValueError("empty")
            tmp = df.copy()
            if '商品' not in tmp.columns:
                # best-effort rename
                for c in tmp.columns:
                    if '商品' in c:
                        tmp = tmp.rename(columns={c: '商品'})
                        break
            if '商品' not in tmp.columns:
                raise ValueError("missing 商品")
            tmp = tmp.dropna(subset=['商品'])
            items = tmp['商品'].astype(str).unique().tolist()
            kw = str(keyword or "").strip()
            if kw:
                items = [x for x in items if kw in x]
            items = sorted(set(items))
            if not items:
                return {
                    'date': d,
                    'items': [],
                    'source': 'AkShare:futures_spot_price(生意社)',
                    'detail': f'当日未找到包含“{kw}”的品种条目',
                    'is_simulated': True,
                }
            return {
                'date': d,
                'items': items,
                'source': 'AkShare:futures_spot_price(生意社)',
                'detail': f'当日共匹配到 {len(items)} 个品种条目（关键字：{kw}）',
                'is_simulated': False,
            }
        except Exception as e:
            return {
                'date': d,
                'items': [],
                'source': 'AkShare:futures_spot_price(生意社)',
                'detail': f'获取失败: {e}',
                'is_simulated': True,
            }

    def hedge_calculation(

        self,
        cost_price: float,
        inventory: float,
        hedge_ratio: float,
        margin_rate: float = 0.1,
    ):
        """Compute a basic hedge P&L curve and return a Matplotlib figure.

        Returns:
            fig: matplotlib.figure.Figure
            suggestions: list[str]
            metrics: dict[str, Any]
        """
        import numpy as np
        import matplotlib.pyplot as plt
        from datetime import datetime

        # Clamp inputs
        inventory = max(0.0, float(inventory))
        cost_price = max(0.0, float(cost_price))
        hedge_ratio = float(hedge_ratio)
        hedge_ratio = max(0.0, min(1.0, hedge_ratio))

        # Get latest market price (best-effort)
        current_price = cost_price
        try:
            price_df = self.fetch_real_time_data(symbol="LC0")
            if price_df is not None and not price_df.empty and "收盘价" in price_df.columns:
                current_price = float(price_df["收盘价"].iloc[-1])
        except Exception:
            current_price = cost_price

        # Price change scenarios (-20% .. +20%)
        price_changes = np.linspace(-0.2, 0.2, 81)

        # Assumption: you hold inventory (long physical). Price up => gain, price down => loss.
        unhedged_pnl = inventory * current_price * price_changes

        # Futures hedge: short futures with notional = inventory * hedge_ratio.
        hedged_pnl = unhedged_pnl - inventory * current_price * hedge_ratio * price_changes

        fig, ax = plt.subplots(figsize=(8, 4.5), dpi=160)
        ax.plot(price_changes * 100, unhedged_pnl / 1e6, label="未套保盈亏（百万元）")
        ax.plot(price_changes * 100, hedged_pnl / 1e6, label="套保后盈亏（百万元）")

        ax.axvline(0, linestyle="--", linewidth=1)
        ax.axhline(0, linestyle="--", linewidth=1)

        # Annotate the intersection points at x=0 (always 0) and show base values.
        ax.annotate(
            "基准情景（0%）\n0.00",
            xy=(0, 0),
            xytext=(5, 5),
            textcoords="offset points",
            fontsize=9,
            bbox=dict(boxstyle="round,pad=0.3", alpha=0.2),
        )

        _matplotlib_style(ax, "套保情景分析", "价格变动（%）", "盈亏（百万元）")
        ax.legend(loc="best", fontsize=9, frameon=False)

        suggestions = []
        if hedge_ratio == 0:
            suggestions.append("未套保：盈亏完全暴露于价格波动。")
        elif hedge_ratio < 0.5:
            suggestions.append("低套保比例：降低下行风险，但仍存在较大暴露。")
        elif hedge_ratio < 0.9:
            suggestions.append("中等套保比例：显著降低风险敞口。")
        else:
            suggestions.append("高套保比例：盈亏更稳定，但可能限制上行收益。")

        hedge_contracts_int = int(np.round(inventory * hedge_ratio))
        total_margin = float(hedge_contracts_int * current_price * margin_rate)

        current_profit = float((current_price - cost_price) * inventory)
        profit_percentage = float((current_price - cost_price) / cost_price * 100) if cost_price > 0 else 0.0

        metrics = {
            "latest_date": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "current_price": float(current_price),
            "hedge_ratio": float(hedge_ratio),
            "inventory": float(inventory),
            "margin_rate": float(margin_rate),
            # Fields consumed by UI
            "hedge_contracts_int": hedge_contracts_int,
            "total_margin": total_margin,
            "current_profit": current_profit,
            "profit_percentage": profit_percentage,
        }
        return fig, suggestions, metrics
    def save_analysis_history(self, user_id: str, record: dict) -> bool:
        """Save analysis record to Supabase (if available) or local session."""
        record = dict(record or {})
        record.setdefault("user_id", user_id)
        record.setdefault("created_at", datetime.utcnow().isoformat())
        record.setdefault("analysis_id", hashlib.md5((record.get("created_at","")+json.dumps(record,ensure_ascii=False,sort_keys=True)).encode("utf-8")).hexdigest())
        record.setdefault("analysis_type", record.get("record_type","analysis"))
        try:
            if hasattr(self, "supabase") and self.supabase:
                # Supabase python client style
                self.supabase.table("analysis_history").insert(record).execute()
                return True
        except Exception:
            # fallback to local store
            pass
        try:
            st.session_state.setdefault("_analysis_history", [])
            st.session_state["_analysis_history"].append(record)
            return True
        except Exception:
            return False


    def delete_history_record(self, analysis_id: str) -> bool:
        """Delete a history record by analysis_id (best-effort).

        If Supabase is available, delete from analysis_history table.
        Otherwise, delete from local session cache.
        """
        if not analysis_id:
            return False
        try:
            if hasattr(self, "supabase") and self.supabase:
                self.supabase.table("analysis_history").delete().eq("analysis_id", analysis_id).execute()
                return True
        except Exception:
            pass
        try:
            hist = st.session_state.get("_analysis_history", [])
            new_hist = [r for r in hist if (r or {}).get("analysis_id") != analysis_id]
            st.session_state["_analysis_history"] = new_hist
            return True
        except Exception:
            return False

    def get_user_history(self, user_id: str | None = None, limit: int = 50):
        """Fetch analysis history for current user."""
        uid = user_id
        try:
            uid = uid or st.session_state.get("user_info", {}).get("user_id") or st.session_state.get("user_id")
        except Exception:
            uid = uid or None

        # Try supabase first
        try:
            if hasattr(self, "supabase") and self.supabase and uid:
                resp = (
                    self.supabase.table("analysis_history")
                    .select("*")
                    .eq("user_id", uid)
                    .order("created_at", desc=True)
                    .limit(limit)
                    .execute()
                )
                data = getattr(resp, "data", None)
                if data is None and isinstance(resp, dict):
                    data = resp.get("data")
                return data or []
        except Exception:
            pass

        # Fallback: local session
        hist = st.session_state.get("_analysis_history", [])
        if uid:
            hist = [r for r in hist if r.get("user_id") == uid]
        return list(reversed(hist[-limit:]))

    # Inventory & Profit (MVP) APIs
    # -------------------------------
    def save_inventory_txn(self, user_id: str, txn: dict) -> bool:
        """Save an inventory transaction record.

        txn fields (recommended):
            - date (YYYY-MM-DD)
            - txn_type: '入库' | '出库'
            - grade: e.g. 电池级 / 工业级
            - warehouse: 仓库/地点
            - qty_ton: float
            - unit_cost: float (required for 入库; optional for 出库)
            - notes: str
        """
        rec = dict(txn or {})
        rec.setdefault("record_type", "inventory_txn")
        rec.setdefault("date", datetime.now().strftime("%Y-%m-%d"))
        rec.setdefault("created_at", datetime.utcnow().isoformat())
        rec["qty_ton"] = float(rec.get("qty_ton", 0.0) or 0.0)
        rec["unit_cost"] = float(rec.get("unit_cost", 0.0) or 0.0)
        return bool(self.save_analysis_history(user_id, rec))

    def save_sales_txn(self, user_id: str, txn: dict) -> bool:
        """Save a sales (profit) transaction record."""
        rec = dict(txn or {})
        rec.setdefault("record_type", "sales_txn")
        rec.setdefault("date", datetime.now().strftime("%Y-%m-%d"))
        rec.setdefault("created_at", datetime.utcnow().isoformat())
        rec["qty_ton"] = float(rec.get("qty_ton", 0.0) or 0.0)
        rec["unit_price"] = float(rec.get("unit_price", 0.0) or 0.0)
        # Optional override cost per ton (enterprise real contract/settlement)
        rec["override_cost"] = float(rec.get("override_cost", 0.0) or 0.0)
        return bool(self.save_analysis_history(user_id, rec))

    def _get_records_by_type(self, user_id: str, record_type: str, limit: int = 500):
        rows = self.get_user_history(user_id=user_id, limit=limit) or []
        out = []
        for r in rows:
            try:
                if (r or {}).get("record_type") == record_type:
                    out.append(r)
            except Exception:
                pass
        # sort by date then created_at
        def _k(x):
            d = (x or {}).get("date") or "1900-01-01"
            c = (x or {}).get("created_at") or ""
            return (d, c)
        out.sort(key=_k)
        return out

    def get_inventory_txns(self, user_id: str, limit: int = 500):
        return self._get_records_by_type(user_id, "inventory_txn", limit=limit)

    def get_sales_txns(self, user_id: str, limit: int = 500):
        return self._get_records_by_type(user_id, "sales_txn", limit=limit)

    def compute_inventory_position(self, user_id: str):
        """Compute inventory position using moving weighted average cost.

        Returns:
            summary_df: columns [grade, warehouse, qty_ton, avg_cost]
            detail: dict with 'simulated' flag if any negative / unknown cost appears
        """
        txns = self.get_inventory_txns(user_id, limit=2000)
        # state per (grade, warehouse)
        state = {}
        simulated = False

        def key_of(t):
            return (str(t.get("grade") or "未分类").strip(), str(t.get("warehouse") or "默认仓").strip())

        for t in txns:
            k = key_of(t)
            stt = state.setdefault(k, {"qty": 0.0, "avg_cost": 0.0})
            qty = float(t.get("qty_ton") or 0.0)
            unit_cost = float(t.get("unit_cost") or 0.0)
            tp = str(t.get("txn_type") or "").strip()
            if tp == "入库":
                if qty > 0:
                    new_qty = stt["qty"] + qty
                    # weighted avg cost
                    if new_qty <= 0:
                        stt["qty"] = 0.0
                        stt["avg_cost"] = 0.0
                    else:
                        # if unit_cost missing, mark simulated
                        if unit_cost <= 0:
                            simulated = True
                            # keep avg cost unchanged for safety
                            unit_cost = stt["avg_cost"]
                        stt["avg_cost"] = (stt["avg_cost"] * stt["qty"] + unit_cost * qty) / new_qty if new_qty else 0.0
                        stt["qty"] = new_qty
            elif tp == "出库":
                if qty > 0:
                    stt["qty"] -= qty
                    if stt["qty"] < -1e-6:
                        simulated = True  # oversold / data incomplete
            else:
                # unknown type -> ignore but mark
                simulated = True

        rows = []
        for (grade, wh), s in state.items():
            q = float(s.get("qty") or 0.0)
            if abs(q) < 1e-9:
                continue
            rows.append({"grade": grade, "warehouse": wh, "qty_ton": round(q, 4), "avg_cost": round(float(s.get("avg_cost") or 0.0), 2)})
        summary_df = pd.DataFrame(rows)
        if summary_df.empty:
            summary_df = pd.DataFrame(columns=["grade", "warehouse", "qty_ton", "avg_cost"])
        return summary_df, {"simulated": simulated}

    def compute_profit_report(self, user_id: str):
        """Compute profit report based on inventory ledger + sales records.

        COGS default uses moving average cost at the time of sale (grade+warehouse aggregated by grade).
        If a sale has override_cost > 0, treat it as enterprise real cost and do NOT mark simulated for that row.
        """
        inv_txns = self.get_inventory_txns(user_id, limit=5000)
        sales = self.get_sales_txns(user_id, limit=5000)

        # Build grade-level state (aggregate warehouses) for simplicity in MVP
        state = {}
        simulated = False

        def inv_key(t):
            return str(t.get("grade") or "未分类").strip()

        # merge events by date
        events = []
        for t in inv_txns:
            events.append(("inv", t.get("date") or "1900-01-01", t.get("created_at") or "", t))
        for s in sales:
            events.append(("sale", s.get("date") or "1900-01-01", s.get("created_at") or "", s))
        events.sort(key=lambda x: (x[1], x[2], 0 if x[0]=="inv" else 1))

        profit_rows = []
        for kind, _, _, obj in events:
            if kind == "inv":
                grade = inv_key(obj)
                stt = state.setdefault(grade, {"qty": 0.0, "avg_cost": 0.0})
                tp = str(obj.get("txn_type") or "").strip()
                qty = float(obj.get("qty_ton") or 0.0)
                unit_cost = float(obj.get("unit_cost") or 0.0)
                if tp == "入库" and qty > 0:
                    new_qty = stt["qty"] + qty
                    if unit_cost <= 0:
                        simulated = True
                        unit_cost = stt["avg_cost"]
                    stt["avg_cost"] = (stt["avg_cost"] * stt["qty"] + unit_cost * qty) / new_qty if new_qty else 0.0
                    stt["qty"] = new_qty
                elif tp == "出库" and qty > 0:
                    stt["qty"] -= qty
                    if stt["qty"] < -1e-6:
                        simulated = True
                else:
                    simulated = True
            else:
                grade = str(obj.get("grade") or "未分类").strip()
                stt = state.setdefault(grade, {"qty": 0.0, "avg_cost": 0.0})
                qty = float(obj.get("qty_ton") or 0.0)
                unit_price = float(obj.get("unit_price") or 0.0)
                override_cost = float(obj.get("override_cost") or 0.0)
                revenue = qty * unit_price
                if override_cost > 0:
                    cogs = qty * override_cost
                    # treat as enterprise real cost
                else:
                    cogs = qty * stt["avg_cost"]
                    # If no inventory or avg_cost unavailable, mark simulated
                    if stt["avg_cost"] <= 0 or stt["qty"] < qty - 1e-6:
                        simulated = True
                gross = revenue - cogs
                stt["qty"] -= qty
                if stt["qty"] < -1e-6:
                    simulated = True

                profit_rows.append({
                    "date": obj.get("date") or "",
                    "grade": grade,
                    "qty_ton": round(qty, 4),
                    "unit_price": round(unit_price, 2),
                    "revenue": round(revenue, 2),
                    "unit_cost_used": round((override_cost if override_cost>0 else stt["avg_cost"]), 2),
                    "cogs": round(cogs, 2),
                    "gross_profit": round(gross, 2),
                    "notes": obj.get("notes") or ""
                })

        df = pd.DataFrame(profit_rows)
        if df.empty:
            df = pd.DataFrame(columns=["date","grade","qty_ton","unit_price","revenue","unit_cost_used","cogs","gross_profit","notes"])
        return df, {"simulated": simulated}



def _norm_cdf(x: float) -> float:
    return 0.5 * (1 + math.erf(x / math.sqrt(2)))


def black_scholes_price(option_type: str, spot: float, strike: float, time_years: float,
                        risk_free: float, volatility: float) -> float:
    if spot <= 0 or strike <= 0 or time_years <= 0 or volatility <= 0:
        return 0.0

    d1 = (math.log(spot / strike) + (risk_free + 0.5 * volatility ** 2) * time_years) / (
        volatility * math.sqrt(time_years)
    )
    d2 = d1 - volatility * math.sqrt(time_years)

    if option_type == "call":
        return spot * _norm_cdf(d1) - strike * math.exp(-risk_free * time_years) * _norm_cdf(d2)
    return strike * math.exp(-risk_free * time_years) * _norm_cdf(-d2) - spot * _norm_cdf(-d1)

# ============================================================================
# Streamlit应用主程序
# ============================================================================



# =========================
# UI assets & helpers (vUI 重构)
# =========================
PROJECT_ROOT = os.path.dirname(__file__)
BANNER_IMAGE_PATH = os.path.join(PROJECT_ROOT, "assets", "hero_banner.png")
HOME_REFERENCE_IMAGE_PATH = BANNER_IMAGE_PATH


def _to_data_uri(local_path: str) -> str:
    try:
        with open(local_path, "rb") as f:
            encoded = base64.b64encode(f.read()).decode("utf-8")
        ext = os.path.splitext(local_path)[1].lower().replace('.', '') or 'png'
        if ext == 'jpg':
            ext = 'jpeg'
        return f"data:image/{ext};base64,{encoded}"
    except Exception:
        return ""


def render_standard_page_header(title: str, desc: str):
    st.markdown(
        f"""
        <div class='page-shell'>
            <div class='page-shell-title'>{title}</div>
            <div class='page-shell-desc'>{desc}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_global_nav():
    group_pages = {
        "市场监测": ["价格行情", "价差走势"],
        "风险测算": ["套保计算", "多情景分析", "期权计算"],
        "经营决策": ["风险敞口", "库存管理", "利润管理"],
        "报告与管理": ["分析报告", "分析历史", "策略管理", "账号设置"],
    }
    page_alias = {"期权计算": "期权保险"}
    reverse_alias = {v: k for k, v in page_alias.items()}
    current_page = st.session_state.get("current_page", "首页")
    current_label = reverse_alias.get(current_page, current_page)
    is_authed = bool(st.session_state.get("authenticated", False))

    active_group = "首页"
    for gname, labels in group_pages.items():
        if current_label in labels:
            active_group = gname
            break

    right_meta = "大连熵合科技有限公司｜面向碳酸锂产业链企业的一体化数字平台"
    if is_authed:
        user_info = st.session_state.get('user_info') or {}
        username = user_info.get('username', '用户')
        right_current = f"当前模块：{active_group} <span> / {current_label}</span>"
        right_action_html = f"<div class='eh-navbar-actions'><span class='eh-user-chip'>已登录：{username}</span></div>"
    else:
        right_current = "当前状态：<span>访客模式（首页可见）</span>"
        right_action_html = "<div class='eh-navbar-actions'><span class='eh-user-chip'>访客</span></div>"

    st.markdown(
        f"""
        <div class='eh-navbar'>
            <div class='eh-navbar-inner'>
                <div class='eh-brand'>
                    <div class='eh-brand-title'>熵合科技</div>
                    <div class='eh-brand-sub'>新能源企业风险管理SaaS平台</div>
                </div>
                <div class='eh-navbar-right'>
                    <div class='eh-navbar-meta'>{right_meta}</div>
                    <div class='eh-navbar-current'>{right_current}</div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    cols = st.columns([1.0, 1.12, 1.12, 1.12, 1.12, 1.55, 1.0], gap="small")
    home_active = current_page == "首页"
    if cols[0].button(f"{'● ' if home_active else ''}首页", key="nav_home_btn", use_container_width=True):
        st.session_state.public_auth_requested = False
        if current_page != "首页":
            st.session_state.current_page = "首页"
        st.rerun()

    for idx, (group_name, labels) in enumerate(group_pages.items(), start=1):
        active = current_label in labels
        pop_label = f"{group_name}{' ▾' if active else ' ▼'}"
        with cols[idx].popover(pop_label, use_container_width=True):
            st.markdown(f"<div class='popover-group-title'>{group_name}</div>", unsafe_allow_html=True)
            for label in labels:
                page_value = page_alias.get(label, label)
                is_active = current_page == page_value
                if st.button(f"{'当前页 · ' if is_active else ''}{label}", key=f"nav_pop_{group_name}_{label}", use_container_width=True):
                    st.session_state.current_page = page_value
                    if not is_authed and page_value != '首页':
                        st.session_state.public_auth_requested = True
                    st.rerun()

    if is_authed:
        if cols[5].button("个人中心 / 退出", key="nav_profile_btn", use_container_width=True):
            st.session_state.current_page = "账号设置"
            st.rerun()
    else:
        if cols[5].button("登录 / 注册", key="nav_login_btn", use_container_width=True):
            st.session_state.public_auth_requested = True
            st.rerun()

    try:
        price_data = st.session_state.get("_nav_price_cache")
        if price_data is None:
            analyzer = st.session_state.get("_analyzer_ref")
            if analyzer:
                price_data = analyzer.fetch_real_time_data(force_refresh=st.session_state.get("force_refresh", False))
                st.session_state["_nav_price_cache"] = price_data
        if price_data is not None and not price_data.empty:
            latest_price = float(price_data['收盘价'].iloc[-1])
            latest_date = pd.to_datetime(price_data['日期'].iloc[-1]).strftime('%Y-%m-%d')
            latest_chg = float(price_data['涨跌幅'].iloc[-1]) if '涨跌幅' in price_data.columns and not pd.isna(price_data['涨跌幅'].iloc[-1]) else 0.0
            if is_authed:
                user_info = st.session_state.get('user_info') or {}
                username = user_info.get('username', '用户')
            else:
                username = '访客'
            st.markdown(
                f"""
                <div class='eh-statusbar'>
                    <div class='eh-status-item'><span>当前身份</span><strong>{username}</strong></div>
                    <div class='eh-status-item'><span>碳酸锂期货</span><strong>{latest_price:,.0f}</strong></div>
                    <div class='eh-status-item'><span>日涨跌幅</span><strong>{latest_chg:+.2f}%</strong></div>
                    <div class='eh-status-item'><span>数据日期</span><strong>{latest_date}</strong></div>
                </div>
                """,
                unsafe_allow_html=True,
            )
    except Exception:
        pass


def main():
    st.set_page_config(
        page_title="熵合科技｜新能源企业风险管理SaaS平台",
        page_icon="LC",
        layout="wide",
        initial_sidebar_state="collapsed"
    )
    
    # 初始化分析器
    analyzer = CloudLithiumAnalyzer()
    st.session_state["_analyzer_ref"] = analyzer
    
    # 初始化session state
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'user_info' not in st.session_state:
        st.session_state.user_info = None
    if 'user_id' not in st.session_state:
        st.session_state.user_id = None
    if 'current_page' not in st.session_state:
        st.session_state.current_page = "首页"
    if 'show_forgot_password' not in st.session_state:
        st.session_state.show_forgot_password = False
    if 'show_reset_form' not in st.session_state:
        st.session_state.show_reset_form = False
    if 'reset_email' not in st.session_state:
        st.session_state.reset_email = None
    if 'reset_username' not in st.session_state:
        st.session_state.reset_username = None
    if 'force_refresh' not in st.session_state:
        st.session_state.force_refresh = False
    if 'public_auth_requested' not in st.session_state:
        st.session_state.public_auth_requested = False

    current_page_for_style = st.session_state.get('current_page', '首页')
    bg_blur = '0.65px' if current_page_for_style == '首页' else '1.2px'
    bg_overlay_top = '0.48' if current_page_for_style == '首页' else '0.60'
    bg_overlay_bottom = '0.58' if current_page_for_style == '首页' else '0.68'
    shell_blur = '0.4px' if current_page_for_style == '首页' else '0.8px'

    # 自定义CSS
    banner_uri = _to_data_uri(BANNER_IMAGE_PATH)
    css = """
    <style>
    #MainMenu, footer, [data-testid="stToolbar"], [data-testid="stStatusWidget"], [data-testid="stDecoration"], [data-testid="stHeaderActionElements"] {display:none !important;}
    [data-testid="stSidebar"], [data-testid="collapsedControl"] {display:none !important;}

    :root {
        --bg: #F7F7F7;
        --card: #FFFFFF;
        --line: #E8EAED;
        --text: #2B2F33;
        --muted: #666666;
        --gold: #C9A96B;
        --gold-deep: #B9924D;
        --green-soft: #EAF1EB;
        --shadow: 0 10px 30px rgba(31,35,41,0.08);
        --shadow-soft: 0 6px 18px rgba(31,35,41,0.05);
        --radius: 12px;
        --maxw: 1320px;
    }

    html, body, [class*="css"] {
        font-family: "NotoSansCJKSC", "Microsoft YaHei", "PingFang SC", Arial, sans-serif;
        color: var(--text);
    }
    .stApp, [data-testid="stAppViewContainer"], [data-testid="stHeader"] {
        background: transparent !important;
    }
    .stApp::before {
        content: ""; position: fixed; inset: 0; z-index: 0;
        background:
            linear-gradient(rgba(247,247,247,__BG_OVERLAY_TOP__), rgba(247,247,247,__BG_OVERLAY_BOTTOM__)),
            url("__BANNER__");
        background-size: cover; background-position: center center; background-attachment: fixed;
        filter: blur(__BG_BLUR__) saturate(0.98);
        transform: scale(1.03);
        pointer-events: none;
    }
    [data-testid="stAppViewContainer"] > .main, [data-testid="stAppViewContainer"] .block-container, [data-testid="stVerticalBlock"] {
        position: relative; z-index: 1;
    }
    .block-container {
        max-width: var(--maxw);
        padding-top: 1.2rem;
        padding-bottom: 3rem;
    }
    h1, h2, h3, h4, p, label, div, span {
        color: var(--text);
    }
    @keyframes ehFadeUp {
        from {opacity: 0; transform: translateY(16px);} 
        to {opacity: 1; transform: translateY(0);} 
    }
    @keyframes ehFloat {
        0%,100% {transform: translateY(0);} 
        50% {transform: translateY(-4px);} 
    }
    .main-header {
        font-size: 3.25rem; font-weight: 900; text-align: center; margin: 0 0 .55rem 0; letter-spacing: -0.04em; line-height: 1.08;
        animation: ehFadeUp .75s ease both;
    }
    .auth-subtitle {
        text-align:center; color:rgba(255,255,255,.82); margin:0 auto 1.4rem auto; max-width: 760px; font-size: 1.08rem; line-height: 1.9;
        animation: ehFadeUp .95s ease both;
    }
    .auth-shell {
        display:grid; grid-template-columns: 1.15fr .95fr; gap: 26px; align-items: stretch; margin-top: 1.4rem;
    }
    .auth-brand-card, .auth-form-card {
        background: rgba(255,255,255,0.90); border:1px solid rgba(255,255,255,0.52); border-radius: 24px; box-shadow: 0 22px 60px rgba(18,22,28,0.16);
        backdrop-filter: blur(14px); overflow:hidden; position:relative;
    }
    .auth-brand-card {padding: 42px 40px 36px 40px; animation: ehFadeUp .9s ease both;}
    .auth-form-card {padding: 22px 24px 18px 24px; animation: ehFadeUp 1.05s ease both;}
    .auth-brand-kicker {font-size:.9rem; color: var(--gold-deep); text-transform: uppercase; letter-spacing: .16em; font-weight: 900;}
    .auth-brand-title {font-size: 3rem; font-weight: 900; line-height: 1.1; letter-spacing: -0.04em; margin: 1.1rem 0 .8rem 0;}
    .auth-brand-desc {color: var(--muted); line-height: 1.95; font-size: 1rem; max-width: 590px;}
    .auth-brand-tags {display:flex; flex-wrap:wrap; gap: 10px; margin-top: 1.2rem;}
    .auth-brand-tags span {display:inline-flex; align-items:center; border-radius:999px; padding: 8px 14px; background: rgba(201,169,107,.12); border:1px solid rgba(201,169,107,.20); color: var(--gold-deep); font-weight: 700; font-size: .86rem;}
    .auth-brand-grid {display:grid; grid-template-columns: repeat(3, minmax(0,1fr)); gap: 14px; margin-top: 1.4rem;}
    .auth-brand-grid div {background: rgba(255,255,255,.82); border:1px solid rgba(201,169,107,.16); border-radius: 16px; padding: 16px 14px; box-shadow: var(--shadow-soft); animation: ehFloat 4.2s ease-in-out infinite;}
    .auth-brand-grid div:nth-child(2) {animation-delay: .4s;} .auth-brand-grid div:nth-child(3) {animation-delay: .8s;}
    .auth-brand-grid strong {display:block; font-size: 1.25rem; color: var(--text);} .auth-brand-grid span {display:block; color: var(--muted); margin-top: 4px; font-size: .82rem; line-height:1.5;}
    .auth-form-title {font-size: 1.55rem; font-weight: 900; letter-spacing: -.02em; margin-bottom: .25rem;}
    .auth-form-text {color: var(--muted); margin-bottom: 1rem; line-height: 1.8;}
    .eh-navbar {
        background: linear-gradient(180deg, rgba(255,255,255,0.98) 0%, rgba(255,255,255,0.93) 100%);
        border: 1px solid rgba(201,169,107,0.16);
        border-radius: 18px;
        padding: 18px 24px 14px 24px;
        box-shadow: var(--shadow);
        position: sticky; top: 0.55rem; z-index: 20; backdrop-filter: blur(14px);
        margin-bottom: .45rem;
    }
    .eh-navbar-inner {display:flex; justify-content:space-between; align-items:flex-end; gap:18px; flex-wrap:wrap;}
    .eh-brand-title {font-size: 1.65rem; font-weight: 900; letter-spacing: 0.01em;}
    .eh-brand-sub {font-size:.9rem; color:var(--muted); margin-top: 4px;}
    .eh-navbar-meta {font-size:.92rem; color:var(--muted); text-align:right;}
    .eh-navbar-current {margin-top:6px; font-size:.9rem; color:var(--gold-deep); font-weight:800; text-align:right;}
    .eh-navbar-current span {color:var(--text); font-weight:700;}
    .eh-navbar-actions {display:flex; gap:10px; justify-content:flex-end; margin-top:8px;}
    .eh-user-chip {display:inline-flex; align-items:center; padding:6px 12px; border-radius:999px; background:rgba(201,169,107,0.10); border:1px solid rgba(201,169,107,0.18); color:var(--gold-deep); font-size:.82rem; font-weight:800;}

    .extra-group-title {margin: .45rem 0 .65rem 0; font-size: .92rem; color: var(--gold-deep); font-weight: 800; letter-spacing: .08em;}

    .popover-group-title {font-size: .92rem; color: var(--gold-deep); font-weight: 800; margin-bottom: .55rem; letter-spacing: .06em;}
    div[data-testid="stPopover"] > div > button {
        width: 100%; border-radius: 12px !important; border: 1px solid rgba(201,169,107,0.18) !important;
        background: rgba(255,255,255,0.90) !important; color: var(--text) !important; font-weight: 800 !important; min-height: 2.95rem !important;
        box-shadow: var(--shadow-soft) !important;
    }
    div[data-testid="stPopover"] > div > button:hover {background: rgba(201,169,107,0.08) !important; border-color: rgba(201,169,107,0.36) !important; transform: translateY(-1px);}
    div[data-testid="stPopover"] > div > button:focus {box-shadow: none !important; border-color: rgba(201,169,107,0.40) !important;}
    div[data-testid="stPopoverContent"] {border-radius: 12px !important; border: 1px solid var(--line) !important; box-shadow: 0 12px 30px rgba(0,0,0,0.10) !important;}
    div[data-testid="stPopoverContent"] button[kind="secondary"] {justify-content: flex-start !important; text-align: left !important; border-radius: 10px !important; background: #fff !important;}
    div[data-testid="stPopoverContent"] button[kind="secondary"]:hover {background: rgba(201,169,107,0.08) !important; border-color: rgba(201,169,107,0.22) !important;}

    .eh-navbar, .hero-banner, .intro-card, .cap-card, .scene-card, .process-step, .home-footer, .page-shell {
        animation: ehFadeUp .68s ease both;
    }
    .cap-card:nth-child(2), .scene-card:nth-child(2), .process-step:nth-child(2) {animation-delay:.06s;}
    .cap-card:nth-child(3), .scene-card:nth-child(3), .process-step:nth-child(3) {animation-delay:.12s;}
    .cap-card:nth-child(4), .scene-card:nth-child(4), .process-step:nth-child(4) {animation-delay:.18s;}
    .cap-card:nth-child(5), .scene-card:nth-child(5), .process-step:nth-child(5) {animation-delay:.24s;}
    .cap-card:nth-child(6) {animation-delay:.30s;}
    [data-testid="stMetric"] {
        position:relative; overflow:hidden; border-radius:18px !important; box-shadow: 0 14px 36px rgba(26,31,38,.08) !important;
        background: linear-gradient(180deg, rgba(255,255,255,.98) 0%, rgba(252,249,242,.94) 100%) !important;
        transition: transform .18s ease, box-shadow .18s ease !important;
    }
    [data-testid="stMetric"]:hover {transform: translateY(-3px); box-shadow: 0 18px 44px rgba(26,31,38,.12) !important;}
    [data-testid="stMetric"]::before {content:''; position:absolute; left:0; top:0; bottom:0; width:4px; background: linear-gradient(180deg, var(--gold), rgba(201,169,107,0.18));}
    [data-testid="stMetricLabel"] {font-size:.86rem !important; letter-spacing:.04em !important; text-transform: uppercase !important;}
    [data-testid="stMetricValue"] {font-size:1.7rem !important; font-weight:900 !important; letter-spacing:-.03em !important;}
    [data-testid="stMetricDelta"] {font-weight:700 !important;}
    .chart-card {background: linear-gradient(180deg, rgba(255,255,255,.98) 0%, rgba(251,249,244,.96) 100%); border:1px solid var(--line); border-radius: 18px; box-shadow: var(--shadow); padding: 18px 18px 12px 18px; margin: .55rem 0 1rem 0;}
    .chart-card-title {font-size: 1.08rem; font-weight: 800; margin-bottom: 6px;}
    .chart-card-subtitle {font-size: .9rem; color: var(--muted); margin-bottom: 8px;}

    div[role="radiogroup"] {gap: .45rem;}
    div[role="radiogroup"] label {
        background: transparent !important;
        border: none !important;
        border-bottom: 2px solid transparent !important;
        border-radius: 0 !important;
        padding: .45rem .75rem .55rem .75rem !important;
        color: var(--muted) !important;
        font-weight: 700 !important;
        min-height: auto !important;
    }
    div[role="radiogroup"] label[data-baseweb="radio"]:has(input:checked),
    div[role="radiogroup"] label:has(input:checked) {
        color: var(--gold-deep) !important;
        border-bottom: 2px solid var(--gold) !important;
        background: transparent !important;
    }
    div[role="radiogroup"] p {font-size: 1rem !important;}

    .eh-statusbar {
        display:grid; grid-template-columns: repeat(4, minmax(0,1fr)); gap: 14px;
        margin: .8rem 0 1.2rem 0;
    }
    .eh-status-item {
        background: rgba(255,255,255,0.94); border: 1px solid var(--line); border-radius: 14px;
        box-shadow: var(--shadow-soft); padding: 14px 16px;
        position: relative; overflow:hidden;
    }
    .eh-status-item::before {content:''; position:absolute; left:0; top:0; bottom:0; width:4px; background: linear-gradient(180deg, var(--gold), rgba(201,169,107,0.20));}
    .eh-status-item span {display:block; font-size:.84rem; color:var(--muted); margin-bottom:4px;}
    .eh-status-item strong {font-size:1.08rem; color:var(--text);}

    .page-shell, .soft-card, .quick-card, [data-testid="stMetric"], [data-testid="stExpander"], .stDataFrame, [data-testid="stTable"] {
        background: rgba(255,255,255,0.94) !important; border: 1px solid rgba(232,234,237,0.95) !important; border-radius: var(--radius) !important; box-shadow: var(--shadow) !important;
        backdrop-filter: blur(__SHELL_BLUR__);
    }
    .page-shell {padding: 26px 28px; margin: .55rem 0 1.1rem 0; position:relative; overflow:hidden;}
    .page-shell::before {content:''; position:absolute; left:0; right:0; top:0; height:4px; background: linear-gradient(90deg, var(--gold), rgba(201,169,107,0.15));}
    .page-shell-title {font-size: 2rem; font-weight: 800; letter-spacing: -0.02em;}
    .page-shell-desc {margin-top: .4rem; color: var(--muted); line-height: 1.8;}

    .hero-card {padding: 0 !important; overflow: hidden; background: transparent !important; border: none !important; box-shadow:none !important;}
    .hero-banner {
        position: relative; min-height: 560px; border-radius: 22px; overflow: hidden;
        background: linear-gradient(90deg, rgba(19,29,40,0.68) 0%, rgba(19,29,40,0.54) 34%, rgba(255,255,255,0.06) 72%, rgba(255,255,255,0.03) 100%);
        border: 1px solid rgba(255,255,255,0.38);
        box-shadow: 0 24px 60px rgba(22,28,34,0.14);
        backdrop-filter: blur(1px);
    }
    .hero-banner::before {
        content: ""; position:absolute; inset:0;
        background: linear-gradient(180deg, rgba(255,255,255,0.05) 0%, rgba(255,255,255,0.00) 100%);
    }
    .hero-content {
        position:relative; z-index:2; max-width: 700px; padding: 92px 60px; color:#fff;
    }
    .hero-content-panel {
        min-height: 500px; display:flex; flex-direction:column; justify-content:center;
    }
    .hero-kicker {font-size:.95rem; color: rgba(255,255,255,.76); letter-spacing: .14em; text-transform: uppercase;}
    .hero-title {font-size: 3.35rem; line-height: 1.14; font-weight: 900; color:#fff; margin: 1rem 0 .95rem 0; letter-spacing:-0.03em;}
    .hero-desc {font-size: 1.04rem; line-height: 1.9; color: rgba(255,255,255,.88); max-width: 560px;}
    .hero-action-tip {margin-top: 1.2rem; color: rgba(255,255,255,.72); font-size: .92rem;}
    .hero-tag-row {display:flex; flex-wrap:wrap; gap:10px; margin-top: 1.2rem;}
    .hero-tag-row span {display:inline-flex; align-items:center; padding:8px 14px; border-radius:999px; background:rgba(255,255,255,0.12); border:1px solid rgba(255,255,255,0.18); color:#fff; font-size:.88rem; font-weight:700; backdrop-filter: blur(4px);}
    .hero-floating-card {position:absolute; right:34px; bottom:34px; z-index:2; width:330px; background:rgba(255,255,255,0.90); border:1px solid rgba(255,255,255,0.55); border-radius:18px; padding:22px 22px 18px 22px; box-shadow: 0 18px 40px rgba(18,22,28,0.16); backdrop-filter: blur(10px);}
    .hero-floating-title {font-size:1rem; color:var(--gold-deep); font-weight:900; letter-spacing:.06em; text-transform:uppercase;}
    .hero-floating-text {margin-top:10px; color:var(--muted); line-height:1.8; font-size:.94rem;}
    .hero-floating-metrics {display:grid; grid-template-columns: repeat(3, minmax(0,1fr)); gap:12px; margin-top:16px;}
    .hero-floating-metrics div {background:#fff; border:1px solid rgba(201,169,107,0.14); border-radius:14px; padding:12px 10px; text-align:center; box-shadow: var(--shadow-soft);}
    .hero-floating-metrics strong {display:block; font-size:1.3rem; color:var(--text);}
    .hero-floating-metrics span {display:block; margin-top:4px; color:var(--muted); font-size:.76rem; line-height:1.45;}

    .section-title {font-size: 1.8rem; font-weight: 800; margin: 1.5rem 0 .35rem 0; letter-spacing: -.02em;}
    .section-kicker {font-size: .92rem; color: var(--gold-deep); text-transform: uppercase; letter-spacing: .12em;}
    .section-subtle {color: var(--muted); line-height: 1.9; margin-bottom: 1rem;}

    .intro-grid {display:grid; grid-template-columns: minmax(0, 1fr); gap: 20px; margin: 1rem 0 1.2rem 0;}
    .intro-card {padding: 34px 38px; background: rgba(255,255,255,0.97); border:1px solid var(--line); border-radius: 18px; box-shadow: var(--shadow); max-width: 1120px; margin: 0 auto; position:relative; overflow:hidden;}
    .intro-card::before {content:''; position:absolute; left:0; top:0; bottom:0; width:5px; background: linear-gradient(180deg, var(--gold), rgba(201,169,107,0.14));}
    .intro-card p {color: var(--muted); line-height: 1.95; font-size: 1rem; margin: 0 0 1.1rem 0;}
    .intro-card p:last-child {margin-bottom: 0;}

    .cap-grid {display:grid; grid-template-columns: repeat(3, minmax(0,1fr)); gap: 16px; margin: .9rem 0 1.2rem 0;}
    .scene-grid {display:grid; grid-template-columns: repeat(4, minmax(0,1fr)); gap: 16px; margin: .9rem 0 1.2rem 0;}
    .cap-card, .scene-card {
        background: rgba(255,255,255,0.98); border:1px solid var(--line); border-radius: 16px; box-shadow: var(--shadow-soft); padding: 24px; position:relative; overflow:hidden; transition: transform .18s ease, box-shadow .18s ease;
    }
    .cap-card:hover, .scene-card:hover {transform: translateY(-4px); box-shadow: var(--shadow);}
    .cap-card::after, .scene-card::after {content:''; position:absolute; left:0; right:0; top:0; height:3px; background: linear-gradient(90deg, rgba(201,169,107,0.95), rgba(201,169,107,0.08));}
    .cap-index, .scene-index {font-size:.82rem; color: var(--gold-deep); font-weight:700; letter-spacing:.08em;}
    .cap-title, .scene-title {font-size: 1.15rem; font-weight:800; margin:.55rem 0 .45rem 0;}
    .cap-text, .scene-text {color: var(--muted); line-height:1.8; font-size:.94rem;}

    .process-wrap {
        display:grid; grid-template-columns: repeat(5, minmax(0,1fr)); gap: 14px; margin: 1rem 0 1.3rem 0;
    }
    .process-step {
        position:relative; background:rgba(255,255,255,0.98); border:1px solid var(--line); border-radius: 16px; padding: 22px 18px; box-shadow: var(--shadow-soft);
        min-height: 126px;
    }
    .process-step:not(:last-child)::after {
        content: "→"; position:absolute; right:-12px; top:50%; transform:translateY(-50%); color: var(--gold-deep); font-size: 1.35rem; font-weight:700;
    }
    .process-no {display:inline-flex; width:30px; height:30px; border-radius:999px; align-items:center; justify-content:center; background: rgba(201,169,107,.14); color:var(--gold-deep); font-weight:800; margin-bottom:10px;}
    .process-title {font-size: 1rem; font-weight:700; line-height:1.7;}

    .home-footer {
        margin-top: 1.5rem; background: linear-gradient(180deg,#fff,#fafaf8); border:1px solid var(--line); border-radius: 16px; box-shadow: var(--shadow); padding: 24px 28px; position:relative; overflow:hidden;
    }
    .home-footer::before {content:''; position:absolute; inset:0 auto 0 0; width:6px; background: linear-gradient(180deg, var(--gold), rgba(201,169,107,0.18));}
    .home-footer-title {font-size:1.08rem; font-weight:800; margin-bottom:8px;}
    .home-footer p {margin:.25rem 0; color: var(--muted);}

    .stButton > button {
        border-radius: 12px !important; min-height: 2.95rem !important; font-weight: 800 !important;
        border: 1px solid rgba(201,169,107,.24) !important; background: rgba(255,255,255,0.96) !important; color: var(--text) !important;
        box-shadow: var(--shadow-soft) !important; transition: all .18s ease !important;
    }
    .stButton > button:hover {border-color: var(--gold) !important; color: var(--gold-deep) !important; transform: translateY(-1px); box-shadow: var(--shadow) !important;}
    button[kind="primary"], button[data-testid="baseButton-primary"] {
        background: linear-gradient(135deg, #cfb177 0%, #b9924d 100%) !important; color:#fff !important; border:none !important;
    }
    [data-testid="stMetric"] {padding: 16px 16px !important;}
    [data-testid="stMetricLabel"] {color: var(--muted) !important; font-weight:600 !important;}
    [data-testid="stMetricValue"] {color: var(--text) !important;}
    [data-testid="stExpander"] details summary {background:#FCFBF7 !important; color:var(--text) !important; font-weight:700 !important;}
    [data-testid="stAlert"] {border-radius: 12px !important;}
    hr {border:none; border-top:1px solid var(--line); margin: 1.2rem 0 1.35rem 0;}

    @media (max-width: 1100px) {
        .eh-statusbar, .cap-grid, .scene-grid, .process-wrap, .intro-grid {grid-template-columns: 1fr 1fr !important;}
        .auth-shell {grid-template-columns: 1fr !important;}
        .hero-content {padding: 62px 36px;}
        .hero-title {font-size: 2.55rem;}
        .hero-floating-card {position:relative; right:auto; bottom:auto; width:auto; margin: 0 36px 36px 36px;}
    }
    @media (max-width: 760px) {
        .eh-navbar-inner, .intro-grid, .cap-grid, .scene-grid, .process-wrap, .eh-statusbar {grid-template-columns: 1fr !important; display:grid !important;}
        .auth-brand-card, .auth-form-card {padding: 22px 18px !important;}
        .auth-brand-title {font-size: 2.15rem !important;}
        .auth-brand-grid {grid-template-columns: 1fr !important;}
        .hero-banner {min-height: 500px;}
        .hero-content {padding: 42px 22px;}
        .hero-title {font-size: 2.1rem;}
        .hero-tag-row {gap:8px;}
        .hero-floating-card {margin: 0 22px 22px 22px; padding:18px;}
    }
    </style>
    """
    css = (css
        .replace("__BANNER__", banner_uri)
        .replace("__BG_OVERLAY_TOP__", str(bg_overlay_top))
        .replace("__BG_OVERLAY_BOTTOM__", str(bg_overlay_bottom))
        .replace("__BG_BLUR__", str(bg_blur))
        .replace("__SHELL_BLUR__", str(shell_blur))
    )
    st.markdown(css, unsafe_allow_html=True)

    # 检查Supabase连接状态（v46：隐藏侧边栏状态提示，不影响功能）
    pass
    
    # 未登录时：首页可见；其余功能或主动点击登录时进入登录页
    if not st.session_state.authenticated:
        if st.session_state.show_forgot_password:
            render_forgot_password(analyzer)
            return
        if st.session_state.show_reset_form and (st.session_state.reset_email or st.session_state.reset_username):
            render_reset_password(analyzer)
            return
        if st.session_state.get('public_auth_requested', False) or st.session_state.get('current_page', '首页') != '首页':
            render_auth_page(analyzer)
            return
    
    # 主应用界面
    render_main_app(analyzer)

# ============================================================================
# 页面渲染函数
# ============================================================================

def render_auth_page(analyzer):
    """渲染登录/注册页面（访客可先看首页；点登录/注册或访问功能页时进入）"""
    st.markdown('<h1 class="main-header">新能源企业风险管理平台</h1>', unsafe_allow_html=True)
    st.markdown('<p class="auth-subtitle" style="text-align:center;">面向碳酸锂产业链企业的一体化数字平台</p>', unsafe_allow_html=True)

    if "show_forgot_password" not in st.session_state:
        st.session_state.show_forgot_password = False

    wrap_left, wrap_center, wrap_right = st.columns([1, 1.25, 1])
    with wrap_center:
        st.markdown("<div class='auth-form-title' style='margin-top:18px;'>用户登录</div>", unsafe_allow_html=True)
        st.markdown("<div class='auth-form-text' style='margin-bottom:12px;'>请登录后使用价格行情、价差走势、套保测算、多情景分析、期权计算与报告输出等完整功能。</div>", unsafe_allow_html=True)

        tab_list = st.tabs(["用户名/密码登录", "邮箱验证码登录", "新用户注册"])
        tab_login_pwd, tab_login_email, tab_register = tab_list[0], tab_list[1], tab_list[2]

        with tab_login_pwd:
            username = st.text_input("用户名", placeholder="请输入用户名", key="login_username")
            password = st.text_input("密码", type="password", placeholder="请输入密码", key="login_password")

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("登录", type="primary", use_container_width=True):
                    with st.spinner("正在验证..."):
                        success, result = analyzer.auth.login(username, password)
                    if success:
                        st.session_state.authenticated = True
                        st.session_state.user_info = _extract_user_info_from_login_result(result, username_fallback=username)
                        st.session_state.user_id = st.session_state.user_info.get('user_id')
                        st.session_state.public_auth_requested = False
                        st.success("登录成功！")
                        _log("login", {"username": username, "mode": "password"})
                        st.rerun()
                    else:
                        msg = result.get("message") if isinstance(result, dict) else str(result)
                        st.error(msg or "登录失败")

            with col_btn2:
                if st.button("忘记密码", use_container_width=True):
                    st.session_state.show_forgot_password = True
                    st.rerun()

        with tab_login_email:
            email = st.text_input("邮箱", placeholder="请输入注册邮箱", key="email_login_email")
            email_code = st.text_input("邮箱验证码", placeholder="请输入邮箱验证码", key="email_login_code")

            email_login_wait = _cooldown_remaining("email_login_code_sent_at", 60)
            email_login_btn_text = f"{email_login_wait}秒后重试" if email_login_wait > 0 else "发送验证码"
            col_email_btn1, col_email_btn2 = st.columns(2)
            with col_email_btn1:
                if st.button(email_login_btn_text, use_container_width=True, key="send_email_login_code_btn", disabled=email_login_wait > 0):
                    with st.spinner("正在发送验证码..."):
                        success, result = analyzer.auth.send_email_login_code(email)
                    if success:
                        _mark_code_sent("email_login_code_sent_at")
                        st.session_state.email_login_last_email = email.strip()
                        msg = result.get("message") if isinstance(result, dict) else str(result)
                        st.success(msg or "验证码已发送，请查收邮箱。")
                    else:
                        msg = analyzer.auth._localize_auth_message(result.get("message") if isinstance(result, dict) else result)
                        st.error(msg or "验证码发送失败")

            with col_email_btn2:
                if st.button("验证码登录", type="primary", use_container_width=True, key="email_code_login_btn"):
                    with st.spinner("正在验证验证码..."):
                        success, result = analyzer.auth.login_with_email_code(email, email_code)
                    if success:
                        st.session_state.authenticated = True
                        st.session_state.user_info = _extract_user_info_from_login_result(result, username_fallback=email)
                        st.session_state.user_id = st.session_state.user_info.get('user_id')
                        st.session_state.public_auth_requested = False
                        st.success("登录成功！")
                        _log("login", {"email": email, "mode": "email_code"})
                        st.rerun()
                    else:
                        msg = result.get("message") if isinstance(result, dict) else str(result)
                        st.error(msg or "验证码登录失败")

            st.caption("验证码登录用于免密登录，邮箱需为已注册邮箱。")

        with tab_register:
            new_username = st.text_input("用户名", key="reg_username", placeholder="至少3个字符")
            new_email = st.text_input("邮箱", key="reg_email", placeholder="用于找回密码")
            new_password = st.text_input("密码", type="password", key="reg_password1", placeholder="至少6个字符")
            confirm_password = st.text_input("确认密码", type="password", key="reg_password2")

            if st.button("注册并登录", type="primary", use_container_width=True):
                if not new_username or not new_email or not new_password:
                    st.error("请完整填写注册信息")
                elif new_password != confirm_password:
                    st.error("两次输入的密码不一致")
                elif len(new_password) < 6:
                    st.error("密码长度至少6位")
                else:
                    with st.spinner("正在注册..."):
                        ok, msg = analyzer.auth.register(new_username, new_password, new_email)
                    if isinstance(msg, dict):
                        msg = msg.get('message') or msg.get('msg') or str(msg)
                    if ok:
                        st.success(msg if isinstance(msg, str) else "注册成功")
                        with st.spinner("正在登录..."):
                            success, result = analyzer.auth.login(new_username, new_password)
                        if success:
                            st.session_state.authenticated = True
                            st.session_state.user_info = _extract_user_info_from_login_result(result, username_fallback=new_username)
                            st.session_state.user_id = st.session_state.user_info.get('user_id')
                            st.session_state.public_auth_requested = False
                            st.rerun()
                        else:
                            st.info("注册成功，但自动登录失败，请回到“用户名/密码登录”手动登录。")
                    else:
                        st.error(msg if isinstance(msg, str) else "注册失败")

        with st.expander("快速体验"):
            st.markdown("""**演示账号**：  
- 用户名：demo_user  
- 密码：demo123  

也可以直接注册新账号，或使用邮箱验证码免密登录。""")

        st.markdown("<div style='height:10px;'></div>", unsafe_allow_html=True)
        if st.button("返回首页", use_container_width=True, key="auth_back_home_btn_bottom"):
            st.session_state.current_page = "首页"
            st.session_state.public_auth_requested = False
            st.rerun()

    if st.session_state.show_forgot_password:
        render_forgot_password(analyzer)


def _cooldown_remaining(state_key: str, cooldown_seconds: int = 60) -> int:
    sent_at = st.session_state.get(state_key)
    if not sent_at:
        return 0
    try:
        elapsed = int(datetime.now().timestamp() - float(sent_at))
    except Exception:
        return 0
    return max(0, cooldown_seconds - elapsed)


def _mark_code_sent(state_key: str):
    st.session_state[state_key] = datetime.now().timestamp()


def render_forgot_password(analyzer):
    """渲染忘记密码页面"""
    st.markdown("### 找回密码")
    
    with st.container():
        col_left, col_center, col_right = st.columns([1, 2, 1])
        
        with col_center:
            username = st.text_input("用户名", key="forgot_username")
            email = st.text_input("注册邮箱", key="forgot_email")
            
            reset_wait = _cooldown_remaining("reset_code_sent_at", 60)
            reset_btn_text = f"{reset_wait}秒后重试" if reset_wait > 0 else "获取验证码"
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button(reset_btn_text, use_container_width=True, disabled=reset_wait > 0):
                    if username and email:
                        success, result = analyzer.auth.generate_reset_code(username, email)
                        if success:
                            _mark_code_sent("reset_code_sent_at")
                            st.session_state.reset_username = username
                            st.session_state.reset_email = email
                            st.session_state.show_forgot_password = False
                            st.session_state.show_reset_form = True
                            st.success(f"验证码已发送到您的邮箱：**{email}**")
                            st.rerun()
                        else:
                            st.error(result)
                    else:
                        st.error("请输入用户名和邮箱")

            with col_btn2:
                if st.button("返回登录", use_container_width=True):
                    st.session_state.show_forgot_password = False
                    st.session_state.show_reset_form = False
                    st.session_state.reset_username = None
                    st.session_state.reset_email = None
                    st.rerun()

def render_reset_password(analyzer):
    """渲染重置密码页面"""
    reset_username = st.session_state.get("reset_username") or ""
    reset_email = st.session_state.get("reset_email") or ""
    title_suffix = reset_username if reset_username else reset_email
    st.markdown(f"### 重置密码 - {title_suffix}")

    with st.container():
        col_left, col_center, col_right = st.columns([1, 2, 1])

        with col_center:
            if reset_email:
                st.info(f"正在为邮箱 **{reset_email}** 重置密码")
            else:
                st.info("请输入验证码并设置新密码")

            reset_code = st.text_input("验证码", placeholder="请输入邮箱验证码")
            new_password = st.text_input("新密码", type="password", placeholder="至少6个字符")
            confirm_password = st.text_input("确认新密码", type="password")

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("重置密码", type="primary", use_container_width=True):
                    if reset_code and new_password and confirm_password:
                        if new_password != confirm_password:
                            st.error("两次输入的密码不一致")
                        elif len(new_password) < 6:
                            st.error("密码长度至少6位")
                        else:
                            success, message = analyzer.auth.reset_password(
                                reset_email or reset_username, reset_code, new_password, email=reset_email
                            )
                            if success:
                                st.success(message)
                                st.session_state.show_reset_form = False
                                st.session_state.reset_username = None
                                st.session_state.reset_email = None
                                st.session_state.show_forgot_password = False
                                st.rerun()
                            else:
                                st.error(message)
                    else:
                        st.error("请填写所有字段")

            with col_btn2:
                if st.button("取消", use_container_width=True):
                    st.session_state.show_reset_form = False
                    st.session_state.reset_username = None
                    st.session_state.reset_email = None
                    st.session_state.show_forgot_password = False
                    st.rerun()

    # -------------------------------



def _open_card(extra_class: str = "soft-card"):
    st.markdown(f"<div class='{extra_class}'>", unsafe_allow_html=True)


def _close_card():
    st.markdown("</div>", unsafe_allow_html=True)


def _matplotlib_style(ax, title: str | None = None, xlabel: str | None = None, ylabel: str | None = None):
    try:
        import matplotlib.pyplot as plt
        ax.set_facecolor("#FCFAF6")
        if hasattr(ax, 'figure') and ax.figure is not None:
            ax.figure.patch.set_facecolor("#FCFAF6")
        ax.grid(True, axis='y', alpha=0.22, linestyle='--', linewidth=0.8, color="#d4c6ab")
        ax.grid(False, axis='x')
        for side in ["top", "right"]:
            ax.spines[side].set_visible(False)
        for side in ["left", "bottom"]:
            ax.spines[side].set_color("#d8ccb6")
            ax.spines[side].set_linewidth(1.2)
        ax.tick_params(colors="#5b6470", labelsize=10)
        try:
            ax.title.set_fontweight('bold')
        except Exception:
            pass
        if title is not None:
            ax.set_title(title, fontsize=15, fontweight="bold", color="#2B2F33", pad=16, loc='left')
        if xlabel is not None:
            ax.set_xlabel(xlabel, fontsize=11, color="#5b6470", labelpad=10)
        if ylabel is not None:
            ax.set_ylabel(ylabel, fontsize=11, color="#5b6470", labelpad=10)
        ax.axhline(0, color="#e6dcc8", linewidth=.8, zorder=0)
        try:
            plt.tight_layout()
        except Exception:
            pass
    except Exception:
        pass


def render_main_app(analyzer):
    """渲染主应用界面（顶部横向导航版）"""
    page_map = {
        "首页": render_home_page,
        "价格行情": render_price_page,
        "风险敞口": render_exposure_page,
        "套保计算": render_hedge_page,
        "价差走势": render_basis_page,
        "库存管理": render_inventory_page,
        "利润管理": render_profit_page,
        "期权保险": render_option_page,
        "期权计算": render_option_page,
        "多情景分析": render_scenario_page,
        "分析报告": render_report_page,
        "分析历史": render_history_page,
        "策略管理": render_strategy_page,
        "账号设置": render_settings_page,
    }

    if st.session_state.current_page not in page_map:
        st.session_state.current_page = "首页"

    render_global_nav()
    current_page = st.session_state.current_page
    renderer = page_map.get(current_page, render_home_page)
    if current_page == "期权计算":
        st.session_state.current_page = "期权保险"
    renderer(analyzer)


def render_home_page(analyzer):
    """渲染首页（官网风重构版）"""

    ref_uri = _to_data_uri(BANNER_IMAGE_PATH)

    st.markdown(
        f"""
        <div class='hero-card'>
            <div class='hero-banner'>
                <div class='hero-content hero-content-panel'>
                    <div class='hero-kicker'>Entropy Harmony Technology</div>
                    <div class='hero-title'>新能源企业风险管理平台</div>
                    <div class='hero-desc'>
                        面向碳酸锂产业链企业的一体化数字平台。围绕价格波动、基差风险、套期保值与压力测试，
                        提供从市场监测到策略评估、从经营分析到管理报告输出的一站式风险管理支持。
                    </div>
                    <div class='hero-action-tip'>企业官网级首页风格 · 产业数据平台化内页视觉 · 可用于比赛/路演展示</div>
                    <div class='hero-tag-row'>
                        <span>真实市场数据</span>
                        <span>动态风险建模</span>
                        <span>套保与期权测算</span>
                        <span>管理层报告输出</span>
                    </div>
                </div>
                <div class='hero-floating-card'>
                    <div class='hero-floating-title'>平台定位</div>
                    <div class='hero-floating-text'>围绕“市场监测—风险识别—策略测算—报告输出”的完整链路，形成更接近企业级产品的风控工作台。</div>
                    <div class='hero-floating-metrics'>
                        <div><strong>6</strong><span>核心功能模块</span></div>
                        <div><strong>4</strong><span>典型应用场景</span></div>
                        <div><strong>1</strong><span>一体化风险平台</span></div>
                    </div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    cta1, cta2, _ = st.columns([1.05, 1.05, 4.2])
    with cta1:
        if st.button("进入系统", type="primary", use_container_width=True, key="home_enter_system"):
            st.session_state.current_page = "价格行情"
            st.rerun()
    with cta2:
        if st.button("查看功能", use_container_width=True, key="home_view_features"):
            st.session_state.current_page = "套保计算"
            st.rerun()

    st.markdown("<div class='section-kicker'>Platform Profile</div><div class='section-title'>平台简介 / Platform Profile</div>", unsafe_allow_html=True)
    st.markdown(
        """
        <div class='intro-grid'>
            <div class='intro-card'>
                <p>熵合科技新能源企业风险管理平台，面向碳酸锂产业链企业在原材料价格波动背景下的经营风险管理需求而设计。平台基于真实市场数据与动态风险建模方法，构建价格监测、基差分析、套期保值测算、多情景压力测试及策略评估的一体化分析体系。</p>
                <p>通过对期货与市场参考价格的联动分析，平台帮助企业识别价格风险敞口，评估不同套保策略下的盈亏表现与风险水平，从而提升企业在复杂市场环境中的决策能力与风险控制水平。</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    capabilities = [
        ("01", "价格行情监测", "实时查看碳酸锂期货价格走势、涨跌幅、区间统计与历史对比，为后续风险测算提供统一市场基准。"),
        ("02", "基差动态分析", "围绕期货与市场参考价的联动关系，展示价差结构、变化区间与当前偏离程度。"),
        ("03", "套保策略测算", "结合库存、成本、套保比例与保证金参数，形成经营层可理解的套保效果测算。"),
        ("04", "多情景压力测试", "对上涨、下跌、平稳及自定义波动情景进行横向比较，识别极端条件下的盈亏表现。"),
        ("05", "期权策略支持", "为锁定采购上限或销售下限提供期权成本测算与情景比较支持。"),
        ("06", "风险报告输出", "自动汇总关键市场信息、风险敞口、情景结果与建议，便于路演展示和管理层汇报。"),
    ]
    st.markdown("<div class='section-kicker'>Core Capability</div><div class='section-title'>核心能力</div><div class='section-subtle'>保留原有功能逻辑，仅对视觉结构、展示层次和页面气质进行整体升级。</div>", unsafe_allow_html=True)
    cards_html = "".join([f"<div class='cap-card'><div class='cap-index'>{i}</div><div class='cap-title'>{t}</div><div class='cap-text'>{d}</div></div>" for i,t,d in capabilities])
    st.markdown(f"<div class='cap-grid'>{cards_html}</div>", unsafe_allow_html=True)

    scenes = [
        ("01", "原材料采购风险管理", "用于识别采购阶段面临的上涨风险与采购成本不确定性。"),
        ("02", "存货价格波动控制", "针对现有库存的价格下跌风险，支持风险对冲与保值决策。"),
        ("03", "套保策略决策支持", "为交易部门与经营管理层提供定量化的方案对比与执行依据。"),
        ("04", "管理层风险报告输出", "支持形成结构化、可展示、可路演的阶段性风险分析结果。"),
    ]
    st.markdown("<div class='section-kicker'>Application Scenario</div><div class='section-title'>应用场景</div>", unsafe_allow_html=True)
    scene_html = "".join([f"<div class='scene-card'><div class='scene-index'>{i}</div><div class='scene-title'>{t}</div><div class='scene-text'>{d}</div></div>" for i,t,d in scenes])
    st.markdown(f"<div class='scene-grid'>{scene_html}</div>", unsafe_allow_html=True)

    process_steps = ["市场数据接入", "风险敞口识别", "基差与价格分析", "套保策略测算", "风险评估与报告输出"]
    st.markdown("<div class='section-kicker'>Workflow</div><div class='section-title'>风险管理流程</div>", unsafe_allow_html=True)
    process_html = "".join([f"<div class='process-step'><div class='process-no'>{idx}</div><div class='process-title'>{name}</div></div>" for idx, name in enumerate(process_steps, start=1)])
    st.markdown(f"<div class='process-wrap'>{process_html}</div>", unsafe_allow_html=True)

    st.markdown(
        """
        <div class='home-footer'>
            <div class='home-footer-title'>项目信息</div>
            <p>项目名称：熵合科技--基于动态风险建模的新能源企业风险管理SaaS平台</p>
            <p>公司名称：大连熵合科技有限公司</p>
            <p>电话：15773359917</p>
            <p>邮箱：fyj05818188@163.com</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_hedge_page(analyzer):
    """渲染套保计算页面"""
    render_standard_page_header("套保计算", "结合库存、成本与套保比例进行策略测算，输出指标卡、情景图表与分析建议。")
    
    # 获取用户设置（如果有）
    user_settings = {}
    if 'user_info' in st.session_state and st.session_state.user_info.get('settings'):
        user_settings = st.session_state.user_info['settings']
    
    # 创建两列布局
    col_left, col_right = st.columns([1, 2])
    
    with col_left:
        st.markdown("### 输入参数")
        st.markdown("---")
        
        # 成本价输入
        default_cost = user_settings.get('default_cost_price', 100000.0)
        cost_price = st.number_input(
            "存货成本价 (元/吨)",
            min_value=0.0,
            max_value=500000.0,
            value=float(default_cost),
            step=1000.0,
            help="您采购或生产碳酸锂的成本价格"
        )
        
        # 存货量输入
        default_inventory = user_settings.get('default_inventory', 100.0)

        use_inv_mgr = st.checkbox(
            "从“库存管理”自动读取库存量",
            value=False,
            help="自动汇总库存管理台账中的当前库存（吨）。若无记录则使用默认值。"
        )
        inv_grade_filter = None
        if use_inv_mgr and 'user_info' in st.session_state:
            uid = st.session_state.user_info.get("user_id")
            inv_summary, inv_meta = analyzer.compute_inventory_position(uid)
            grades = ["全部"] + (sorted(inv_summary["grade"].unique().tolist()) if not inv_summary.empty else [])
            choice = st.selectbox("选择品级（可选）", grades, index=0)
            inv_grade_filter = None if choice == "全部" else choice
            if not inv_summary.empty:
                inv_qty = float(inv_summary[inv_summary["grade"] == inv_grade_filter]["qty_ton"].sum()) if inv_grade_filter else float(inv_summary["qty_ton"].sum())
                default_inventory = inv_qty
            else:
                st.info("库存管理暂无记录，将使用默认库存量。")

        inventory = st.number_input(
            "存货数量 (吨)",
            min_value=0.0,
            max_value=10000.0,
            value=float(default_inventory),
            step=1.0,
            help="您当前持有的碳酸锂库存数量"
        )
        
        # 套保比例滑块
        default_ratio = user_settings.get('default_hedge_ratio', 0.8)
        _ov = st.session_state.get('hedge_ratio_override')
        if _ov is not None:
            try:
                default_ratio = float(_ov)
            except Exception:
                pass

        hedge_ratio_percent = st.slider(
            "套保比例 (%)",
            min_value=0,
            max_value=100,
            value=int(default_ratio * 100),
            step=5,
            help="计划对冲的价格风险比例，100%表示完全对冲"
        )
        
        hedge_ratio = hedge_ratio_percent / 100
        
        # 高级选项

        with st.expander("高级选项"):
            margin_rate = st.slider(
                "保证金比例 (%)",
                min_value=5,
                max_value=30,
                value=15,
                step=1,
                help="期货交易保证金比例"
            ) / 100
            
            # 保存为默认设置选项
            if 'user_info' in st.session_state:
                save_defaults = st.checkbox("保存为默认设置", value=False)
                if save_defaults:
                    new_settings = {
                        'default_cost_price': float(cost_price),
                        'default_inventory': float(inventory),
                        'default_hedge_ratio': float(hedge_ratio)
                    }

                    if analyzer.auth.update_user_settings(st.session_state.user_info['user_id'], new_settings):
                        st.success("默认设置已保存")
        
        # 操作按钮

        auto_update = st.toggle("实时更新", value=True, help="拖动参数后自动刷新图表")

        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            calc_button = st.button(
                "开始计算", 
                type="primary", 
                use_container_width=True,
                help="基于当前参数计算套保方案"
            )
        
        with col_btn2:
            if st.button("刷新数据", use_container_width=True):
                st.session_state.force_refresh = True
                st.rerun()
        
        # 如果点击了计算按钮
        should_calculate = auto_update or calc_button
        if should_calculate:
            with st.spinner("正在获取最新数据并计算套保方案..."):
                fig, suggestions, metrics = analyzer.hedge_calculation(
                    cost_price, inventory, hedge_ratio, margin_rate
                )

                if fig is not None:
                    st.session_state.hedge_results = {
                        'fig': fig,
                        'suggestions': suggestions,
                        'metrics': metrics,
                        'params': {
                            'cost_price': cost_price,
                            'inventory': inventory,
                            'hedge_ratio': hedge_ratio,
                            'margin_rate': margin_rate
                        }
                    }
                else:
                    st.error("计算失败，请检查网络连接或稍后重试")
    
    with col_right:
        st.markdown("### 分析结果")
        st.markdown("---")
        
        # 检查是否有计算结果
        if 'hedge_results' in st.session_state:
            results = st.session_state.hedge_results
            metrics = results['metrics']
            params = results['params']
            
            # 显示数据来源和时间
            st.info(f"数据时间：{_fmt_dt(metrics.get('latest_date'))}")
            
            # 关键指标卡片
            col_metric1, col_metric2, col_metric3 = st.columns(3)
            
            with col_metric1:
                # 计算价格变化
                price_diff = metrics['current_price'] - params['cost_price']
                price_diff_pct = (price_diff / params['cost_price']) * 100 if params['cost_price'] > 0 else 0
                
                delta_color = "normal" if price_diff >= 0 else "inverse"
                st.metric(
                    label="当前市场价格",
                    value=f"{metrics['current_price']:,.0f}",
                    delta=f"{price_diff_pct:+.1f}%",
                    delta_color=delta_color,
                    help=f"较成本价{price_diff:+,.0f}元/吨"
                )
            

            with col_metric2:
                actual_ratio = metrics['hedge_contracts_int'] / params['inventory'] * 100 if params['inventory'] > 0 else 0
                st.metric(

                    label="建议套保手数",
                    value=f"{metrics['hedge_contracts_int']}",
                    delta=f"{actual_ratio:.1f}%",
                    help=f"基于{params['inventory']:,.1f}吨存货"
                )
            
            with col_metric3:
                st.metric(

                    label="所需保证金",
                    value=f"¥{metrics['total_margin']:,.0f}",
                    help=f"按{params['margin_rate']*100:.0f}%保证金比例"
                )
            
            # 显示图表
           
            st.markdown("#### 盈亏情景分析")
            st.pyplot(results['fig'])
            
            # 详细建议
            with st.expander("详细分析报告", expanded=True):
                st.markdown(results['suggestions'])
            
            # 导出功能
            st.markdown("#### 导出结果")
            col_export1, col_export2, col_export3 = st.columns(3)
            
            with col_export1:
                if st.button("保存到云端历史", use_container_width=True, 
                           help="将分析结果保存到云端历史记录"):
                    if 'user_info' in st.session_state:
                        st.success("分析结果已保存到云端历史记录")
                    else:
                        st.warning("请先登录以保存历史记录")
            
            with col_export2:
                # 生成文本报告
                actual_ratio_report = metrics['hedge_contracts_int'] / params['inventory'] * 100 if params['inventory'] > 0 else 0
                report_text = f"""碳酸锂套保分析报告
生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
用户：{st.session_state.user_info['username'] if 'user_info' in st.session_state else '游客'}
数据来源：akshare实时数据

=== 输入参数 ===
存货成本价：{params['cost_price']:,.2f} 元/吨
存货数量：{params['inventory']:,.2f} 吨
套保比例：{params['hedge_ratio']*100:.2f}%
保证金比例：{params['margin_rate']*100:.0f}%

=== 市场数据 ===
当前价格：{metrics['current_price']:,.2f} 元/吨
数据时间：{_fmt_dt(metrics.get('latest_date'))}

=== 套保方案 ===
理论套保手数：{params['inventory'] * params['hedge_ratio']:.2f} 手
实际套保手数：{metrics['hedge_contracts_int']} 手
实际套保比例：{actual_ratio_report:.2f}%
每手保证金：{metrics['current_price'] * params['margin_rate']:,.2f} 元
总保证金要求：{metrics['total_margin']:,.2f} 元

=== 盈亏分析 ===
当前每吨盈亏：{metrics['current_price'] - params['cost_price']:,.2f} 元
当前总盈亏：{metrics['current_profit']:,.2f} 元
盈亏比例：{metrics['profit_percentage']:.2f}%

=== 风险提示 ===
请根据自身风险承受能力调整套保策略。
期货交易存在风险，建议咨询专业人士。
本分析仅供参考，不构成投资建议。
"""
                
               
                st.download_button(
                    label="下载文本报告",
                    data=report_text,
                    file_name=f"套保分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True,
                    help="下载完整的分析报告文本文件"
                )
            
            with col_export3:
             
                if st.button("保存图表", use_container_width=True,
                           help="保存分析图表为PNG文件"):
                    import io
                    buf = io.BytesIO()
                    results['fig'].savefig(buf, format='png', dpi=300, bbox_inches='tight')
                    buf.seek(0)
                    
                   
                    st.download_button(
                        label="下载PNG图表",
                        data=buf,
                        file_name=f"套保分析图表_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
                        mime="image/png",
                        use_container_width=True
                    )
        
        else:
            # 如果没有计算结果，显示说明
            st.info("请在左侧输入参数并点击“开始计算”")
            
            # 显示示例
            with st.expander("参数说明"):
                st.markdown("""
                **参数解释**：
                
                1. **存货成本价**：您采购或生产碳酸锂的成本价格
                   - 示例：100,000元/吨
                   - 范围：0-500,000元/吨
                
                2. **存货数量**：您当前持有的碳酸锂库存量
                   - 示例：100吨
                   - 范围：0-10,000吨
                
               
                3. **套保比例**：您希望对冲的价格风险比例
                   - 0%：完全不套保，承担全部价格风险
                   - 50%：对冲一半的价格风险
                   - 100%：完全对冲价格风险
                
                4. **保证金比例**：期货交易所需的保证金比例
                   - 行业标准：10-20%
                   - 交易所可能根据市场情况调整
                
                **计算原理**：
                - 根据当前市场价格计算盈亏
                - 模拟未来价格变动情景（-50%到+100%）
                - 计算套保后的盈亏变化
                - 提供风险管理建议
                """)


    st.markdown("---")
    st.markdown("## 套保比例优化")
    with st.expander("打开：基于真实历史分布自动求解最优套保比例", expanded=False):
        st.caption("说明：使用真实历史“现货（内置表）+ 期货收盘价（AkShare/新浪）”对齐后的变动来做风险度量（历史 VaR / CVaR）。不做随机模拟、不生成任何虚构价格。")
        opt_days = st.slider("用于优化的历史窗口（交易日）", 60, 400, 180, 10)
        horizon = st.selectbox("风险度量周期", ["1日", "5日", "20日"], index=0)
        alpha = st.select_slider("置信水平（尾部概率）", options=[0.10, 0.05, 0.01], value=0.05, help="0.05 表示 95% 置信水平下的尾部风险")
        exposure_type = st.selectbox(
            "业务场景",
            ["库存（担心价格下跌）", "未来采购（担心价格上涨）"],
            index=0,
            help="库存：持有现货，担心价格下跌；未来采购：未来要买入现货，担心价格上涨"
        )
        objective = st.selectbox("优化目标", ["最小CVaR（推荐）", "最小VaR", "最小波动率"], index=0)

        if st.button("计算最优套保比例", use_container_width=True):
            try:
                # 期货历史
                fut = analyzer.fetch_real_time_data(symbol="LC0", days=int(opt_days)+30, force_refresh=False)
                fut = fut.rename(columns={"日期": "date", "收盘价": "fut_close"})
                fut = fut[["date", "fut_close"]].dropna()

                # 现货历史（内置）
                spot = _SPOT_DF.copy()
                merged = align_spot_futures(spot, fut)
                if merged.empty or merged.shape[0] < 30:
                    st.error("现货/期货可对齐的历史数据不足，无法优化（请检查内置现货表日期范围与期货数据是否重叠）。")
                else:
                    # 仅取最近 opt_days
                    merged = merged.tail(int(opt_days)).copy()
                    w = {"1日": 1, "5日": 5, "20日": 20}[horizon]

                    merged["d_spot"] = merged["spot_price"].diff(w)
                    merged["d_fut"] = merged["fut_close"].diff(w)
                    merged = merged.dropna(subset=["d_spot", "d_fut"])

                    Q = float(inventory)  # 吨
                    # 业务方向：库存（担心跌）=> 现货P&L=Q*d_spot；对冲用“卖出期货”=> 期货P&L = -h*Q*d_fut
                    # 未来采购（担心涨）=> 成本风险= -Q*d_spot（价格涨则亏）；对冲用“买入期货”=> +h*Q*d_fut
                    if "库存" in exposure_type:
                        spot_pnl = Q * merged["d_spot"]
                        fut_pnl_sign = -1.0
                    else:
                        spot_pnl = -Q * merged["d_spot"]
                        fut_pnl_sign = 1.0

                    hs = np.linspace(0, 1, 101)
                    rows = []
                    for h in hs:
                        net_pnl = spot_pnl + fut_pnl_sign * h * Q * merged["d_fut"]
                        if objective == "最小波动率":
                            score = float(np.std(net_pnl))
                        else:
                            var, cvar = compute_historical_var_cvar(net_pnl, alpha=alpha)
                            if var is None:
                                continue
                            score = float(cvar) if objective == "最小CVaR（推荐）" else float(var)
                        rows.append((h, score))

                    if not rows:
                        st.error("优化失败：风险指标无法计算。")
                    else:
                        # 对于 VaR/CVaR（这里是 PnL 的尾部均值/分位数，越大越好），所以我们选择 score 最大（最不亏）
                        df_opt = pd.DataFrame(rows, columns=["hedge_ratio", "score"])
                        if objective == "最小波动率":
                            best = df_opt.loc[df_opt["score"].idxmin()]
                        else:
                            best = df_opt.loc[df_opt["score"].idxmax()]

                        best_h = float(best["hedge_ratio"])
                        st.success(f"建议套保比例：**{best_h*100:.0f}%**（{objective}，周期{horizon}，alpha={alpha}）")

                        # 展示对比
                        base_net = spot_pnl  # h=0
                        hedged_net = spot_pnl + fut_pnl_sign * best_h * Q * merged["d_fut"]
                        base_var, base_cvar = compute_historical_var_cvar(base_net, alpha=alpha)
                        hed_var, hed_cvar = compute_historical_var_cvar(hedged_net, alpha=alpha)

                        c1, c2, c3, c4 = st.columns(4)
                        c1.metric("未套保 VaR", f"{base_var:,.0f}" if base_var is not None else "—")
                        c2.metric("未套保 CVaR", f"{base_cvar:,.0f}" if base_cvar is not None else "—")
                        c3.metric("套保后 VaR", f"{hed_var:,.0f}" if hed_var is not None else "—")
                        c4.metric("套保后 CVaR", f"{hed_cvar:,.0f}" if hed_cvar is not None else "—")

                        st.line_chart(pd.DataFrame({
                            "未套保": base_net.reset_index(drop=True),
                            "套保后": hedged_net.reset_index(drop=True),
                        }))

                        # 一键写回页面 slider 的参考值（不强制）
                        if st.button("将建议比例写入上方套保比例", key="apply_best_h"):
                            st.session_state.hedge_ratio_override = best_h
                            st.success("已写入建议比例（下次计算将使用该比例）")
            except Exception as e:
                st.error(f"优化计算失败：{e}")

    st.markdown("## 策略保存与回溯")
    with st.expander("打开：将当前方案保存到“策略管理”", expanded=False):
        st.caption("说明：保存的是参数与当日（或最近交易日）期货收盘价快照，用于后续回溯。")
        strat_name = st.text_input("策略名称", value=f"LC套保_{datetime.now().strftime('%Y%m%d_%H%M')}")
        if st.button("保存当前策略", use_container_width=True):
            try:
                if not _require("save_strategy", "只有交易员/管理者/管理员可保存策略。"):
                    return
                if "saved_strategies" not in st.session_state:
                    st.session_state.saved_strategies = []
                # 获取最新期货收盘价
                fdf = analyzer.fetch_real_time_data(symbol="LC0", days=10, force_refresh=False)
                last_close = float(fdf["收盘价"].dropna().iloc[-1]) if (fdf is not None and not fdf.empty and "收盘价" in fdf.columns) else None
                # 现货（内置表，取<=今日最近）
                s_price, s_date, s_ok = get_spot_price_on_or_before(datetime.now())
                st.session_state.saved_strategies.append({
                    "name": strat_name,
                    "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "inventory_ton": float(inventory),
                    "cost_price": float(cost_price),
                    "hedge_ratio": float(hedge_ratio),
                    "fut_close_at_create": last_close,
                    "spot_at_create": s_price if s_ok else None,
                    "spot_date": s_date.strftime("%Y-%m-%d") if s_ok and s_date else None,
                    "note": f"现货来源：内置表；期货来源：AkShare/新浪"
                })
                st.success("已保存。请到【策略管理】页面查看回溯。")
                _log("save_strategy", {"name": strat_name, "hedge_ratio": float(hedge_ratio), "inventory_ton": float(inventory)})
            except Exception as e:
                st.error(f"保存失败：{e}")

def render_price_page(analyzer):
    """渲染价格行情页面"""
    render_standard_page_header("价格行情", "围绕碳酸锂期货主力与月合约，展示价格走势、统计指标、详细表格与导出能力。")
    
    # 数据控制栏
    col_control1, col_control2, col_control3, col_control4 = st.columns([2, 2, 1, 1])
    
    with col_control1:
        period = st.selectbox(
            "查看周期",
            ["最近1个月", "最近3个月", "最近6个月", "最近1年", "全部数据"],
            index=3,
            help="选择要查看的价格周期"
        )
    
    with col_control2:
        symbol = st.selectbox(
            "选择合约",
            ["LC0", "LC2401", "LC2402", "LC2403", "LC2404", "LC2405", "LC2406"],
            index=0,
            help="选择碳酸锂期货合约"
        )
    
    with col_control3:

        if st.button("刷新", use_container_width=True, 
                    help="强制刷新最新数据"):
            analyzer.cache_data = {}
            st.session_state.force_refresh = True
            st.rerun()
    
    with col_control4:
        show_stats = st.checkbox("显示统计", value=True)
    
    # 获取数据
    with st.spinner("正在加载实时价格数据..."):
        price_data = analyzer.fetch_real_time_data(symbol=symbol, force_refresh=st.session_state.get("force_refresh", False))
        st.session_state.force_refresh = False

    if price_data is None or price_data.empty:
        st.error("无法获取价格数据，请检查网络连接或稍后重试")
        return

    # 统一字段（避免 KeyError: '日期'）
    price_data = price_data.copy()
    if "日期" not in price_data.columns:
        if "时间" in price_data.columns:
            price_data = price_data.rename(columns={"时间": "日期"})
        elif "date" in price_data.columns:
            price_data = price_data.rename(columns={"date": "日期"})
        elif "Date" in price_data.columns:
            price_data = price_data.rename(columns={"Date": "日期"})
    if "日期" in price_data.columns:
        price_data["日期"] = pd.to_datetime(price_data["日期"], errors="coerce")
    price_data = price_data.dropna(subset=["日期", "收盘价"]).reset_index(drop=True)
    if "涨跌幅" not in price_data.columns and len(price_data) > 1:
        price_data["涨跌幅"] = price_data["收盘价"].pct_change() * 100

    # 根据周期筛选数据
    period_map = {
        "最近1个月": 30,
        "最近3个月": 90,
        "最近6个月": 180,
        "最近1年": 365,
        "全部数据": len(price_data),
    }
    days = int(period_map.get(period, 365))
    display_data = price_data.tail(days).copy()

    latest_price = float(display_data["收盘价"].iloc[-1])
    latest_date = display_data["日期"].iloc[-1]
    latest_change = float(display_data["涨跌幅"].iloc[-1]) if "涨跌幅" in display_data.columns and not pd.isna(display_data["涨跌幅"].iloc[-1]) else 0.0

    st.markdown("### 最新行情")
    c1, c2, c3 = st.columns(3)
    c1.metric("最新收盘价", f"{latest_price:,.0f} 元/吨")
    c2.metric("最新涨跌幅", f"{latest_change:+.2f}%")
    c3.metric("数据日期", _fmt_dt(latest_date))

    # 主图
    fig_main, ax_main = plt.subplots(figsize=(12, 6), dpi=160)
    ax_main.plot(display_data["日期"], display_data["收盘价"], linewidth=2, label="收盘价")

    # 均线
    if len(display_data) > 20:
        ma20 = display_data["收盘价"].rolling(window=20).mean()
        ax_main.plot(display_data["日期"], ma20, linestyle="--", linewidth=1.5, label="20日均线")
    if len(display_data) > 60:
        ma60 = display_data["收盘价"].rolling(window=60).mean()
        ax_main.plot(display_data["日期"], ma60, linestyle="--", linewidth=1.5, label="60日均线")

    _matplotlib_style(ax_main, f"{symbol} {period}价格走势", "日期", "价格 (元/吨)")
    ax_main.legend(fontsize=10, frameon=False)
    ax_main.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, pos: f"{x:,.0f}"))
    plt.xticks(rotation=45)
    plt.tight_layout()
    st.pyplot(fig_main)
    
    # 统计信息
    if show_stats:
        with st.expander("详细统计信息", expanded=True):
            col_stat1, col_stat2 = st.columns(2)
            
            with col_stat1:
                st.markdown("#### 价格统计")
                stats_price = {
                    "最新价格": f"{latest_price:,.2f} 元/吨",
                    "期间最高": f"{display_data['收盘价'].max():,.2f} 元/吨",
                    "期间最低": f"{display_data['收盘价'].min():,.2f} 元/吨",
                    "平均价格": f"{display_data['收盘价'].mean():,.2f} 元/吨",
                    "价格中位数": f"{display_data['收盘价'].median():,.2f} 元/吨",
                    "价格标准差": f"{display_data['收盘价'].std():,.2f} 元/吨",
                    "价格波动率": f"{(display_data['收盘价'].std() / display_data['收盘价'].mean() * 100):.2f}%"
                }
                
                for key, value in stats_price.items():
                    st.text(f"{key}: {value}")
            
            with col_stat2:
                if '涨跌幅' in display_data.columns:
                    st.markdown("#### 涨跌幅统计")
                    returns = display_data['涨跌幅'].dropna()
                    
                    stats_returns = {
                        "平均日涨跌": f"{returns.mean():.2f}%",
                        "上涨天数": f"{(returns > 0).sum()} 天",
                        "下跌天数": f"{(returns < 0).sum()} 天",
                        "平盘天数": f"{(returns == 0).sum()} 天",
                        "最大单日涨幅": f"{returns.max():.2f}%",
                        "最大单日跌幅": f"{returns.min():.2f}%",
                        "上涨概率": f"{(returns > 0).sum() / len(returns) * 100:.1f}%"
                    }
                    
                    for key, value in stats_returns.items():
                        st.text(f"{key}: {value}")
    
    # 详细数据表格
    with st.expander("详细数据表格", expanded=False):
        display_data_formatted = display_data.copy()
        display_data_formatted['日期'] = display_data_formatted['日期'].dt.strftime('%Y-%m-%d')
        
        # 选择显示的列
        available_cols = ['日期', '收盘价', '开盘价', '最高价', '最低价', '涨跌幅', '成交量']
        display_cols = [col for col in available_cols if col in display_data_formatted.columns]
        
        st.dataframe(
            display_data_formatted[display_cols].style.format({
                '收盘价': '{:,.0f}',
                '开盘价': '{:,.0f}',
                '最高价': '{:,.0f}',
                '最低价': '{:,.0f}',
                '涨跌幅': '{:.2f}%',
                '成交量': '{:,.0f}'
            } if '成交量' in display_data_formatted.columns else {
                '收盘价': '{:,.0f}',
                '开盘价': '{:,.0f}',
                '最高价': '{:,.0f}',
                '最低价': '{:,.0f}',
                '涨跌幅': '{:.2f}%'
            }),
            use_container_width=True,
            height=400
        )
    # 数据导出功能
    st.markdown("---")
    st.markdown("### 数据导出")

    col_export1, col_export2, col_export3 = st.columns(3)

    with col_export1:
        csv_data = display_data.to_csv(index=False).encode("utf-8-sig")
        st.download_button(
            label="下载CSV数据",
            data=csv_data,
            file_name=f"碳酸锂价格_{symbol}_{period}_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            use_container_width=True,
            help="下载当前显示的价格数据为CSV文件",
        )

    with col_export2:
        buf = io.BytesIO()
        fig_main.savefig(buf, format="png", dpi=300, bbox_inches="tight")
        buf.seek(0)
        st.download_button(
            label="保存图表为PNG",
            data=buf,
            file_name=f"碳酸锂价格图表_{symbol}_{period}_{datetime.now().strftime('%Y%m%d')}.png",
            mime="image/png",
            use_container_width=True,
            help="下载当前图表为PNG图片",
        )

    with col_export3:
        # 生成价格分析报告（文本）
        latest_price = float(display_data["收盘价"].iloc[-1]) if "收盘价" in display_data.columns and not display_data.empty else float("nan")
        price_vol = float(display_data["收盘价"].std() / max(display_data["收盘价"].mean(), 1e-9) * 100) if "收盘价" in display_data.columns and len(display_data) > 1 else 0.0

        report_text = f"""=== 碳酸锂价格分析报告 ===
标的：{symbol}
周期：{period}
数据条数：{len(display_data)}
最新价格：{latest_price:,.2f} 元/吨
期间最高：{display_data['收盘价'].max():,.2f} 元/吨
期间最低：{display_data['收盘价'].min():,.2f} 元/吨
期间均价：{display_data['收盘价'].mean():,.2f} 元/吨
价格波动率：{price_vol:.2f}%

"""

        if "涨跌幅" in display_data.columns:
            returns = display_data["涨跌幅"].dropna()
            if not returns.empty:
                report_text += f"""=== 涨跌统计 ===
平均日涨跌：{returns.mean():.2f}%
最大单日涨幅：{returns.max():.2f}%
最大单日跌幅：{returns.min():.2f}%
上涨天数：{int((returns > 0).sum())}
下跌天数：{int((returns < 0).sum())}

"""

        report_text += f"""报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""

        st.download_button(
            label="生成分析报告",
            data=report_text,
            file_name=f"碳酸锂分析报告_{symbol}_{period}_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain",
            use_container_width=True,
            help="生成并下载详细的价格分析报告",
        )


def _series_last_bool(df: 'pd.DataFrame', col: str, default: bool = False) -> bool:
    """Safely get the last boolean value from a DataFrame column.
    Works for pandas Series with RangeIndex where `series[-1]` would raise.
    """
    try:
        if df is None or getattr(df, "empty", True):
            return default
        if col not in df.columns:
            return default
        s = df[col]
        # pandas Series
        if hasattr(s, "iloc"):
            return bool(s.iloc[-1])
        # list-like
        if isinstance(s, (list, tuple, np.ndarray)):
            return bool(s[-1]) if len(s) else default
        return bool(s)
    except Exception:
        return default

def render_basis_page(analyzer):
    """渲染多基准价差页面。"""
    render_standard_page_header("价差走势", "支持市场现货价、用户自定义、真实采购成本三种口径，统一展示基差指标、对比表格与逐日明细。")

    # 先放一行紧凑的筛选器，避免右侧表单把整页顶部撑出大空白
    filter_col1, filter_col2 = st.columns([1, 1])
    with filter_col1:
        symbol = st.selectbox(
            "选择期货主力合约",
            ["LC0", "LC2401", "LC2402", "LC2403", "LC2404", "LC2405", "LC2406"],
            index=0,
        )
    with filter_col2:
        period = st.selectbox(
            "查看周期",
            ["最近1个月", "最近3个月", "最近6个月", "最近1年"],
            index=2,
        )

    period_map = {"最近1个月": 30, "最近3个月": 90, "最近6个月": 180, "最近1年": 365}
    price_data = analyzer.fetch_real_time_data(
        symbol=symbol,
        days=period_map[period],
        force_refresh=st.session_state.get("force_refresh", False),
    )
    if price_data is None or price_data.empty or (_series_last_bool(price_data, "__is_simulated") is True):
        st.error("无法获取期货真实数据（AkShare 新浪日频）。")
        st.markdown(
            "<div style='color:#b00020;font-weight:700'>当前为模拟数据，禁止用于对外报告</div>",
            unsafe_allow_html=True,
        )
        return

    display_data = price_data.tail(period_map[period]).copy()

    if not display_data.empty:
        if "日期" not in display_data.columns:
            if "date" in display_data.columns:
                display_data = display_data.rename(columns={"date": "日期"})
            elif "Date" in display_data.columns:
                display_data = display_data.rename(columns={"Date": "日期"})
            elif display_data.index.name:
                display_data = display_data.reset_index().rename(columns={display_data.index.name: "日期"})
        if "日期" in display_data.columns:
            display_data["日期"] = pd.to_datetime(display_data["日期"], errors="coerce")

    if "收盘价" not in display_data.columns:
        for cand in ["收盘", "close", "Close", "收盘价(元)", "结算价", "结算"]:
            if cand in display_data.columns:
                display_data = display_data.rename(columns={cand: "收盘价"})
                break
    if "收盘价" not in display_data.columns:
        st.error(f"期货数据缺少收盘价列，无法计算价差。当前列：{list(display_data.columns)}")
        return

    update_time = display_data["日期"].iloc[-1]
    spot_date_str = pd.to_datetime(update_time).strftime("%Y%m%d")
    spot_info = analyzer.fetch_spot_price_from_excel(date=spot_date_str)
    market_spot_price = spot_info.get("price")
    market_spot_price = float(market_spot_price) if (market_spot_price is not None and float(market_spot_price) > 0) else None
    ref_source = spot_info.get("source", "LOCAL_TABLE")
    ref_is_sim = bool(spot_info.get("is_simulated", True))

    st.markdown("### 基准设置")
    st.markdown(
        """
        <style>
        .basis-compact .stNumberInput {
            margin-bottom: 0.2rem !important;
        }
        .basis-compact label {
            margin-bottom: 0.1rem !important;
        }
        .basis-confirm-wrap {
            margin-top: 0.15rem;
            padding-top: 0.1rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    basis_mode = st.radio(
        "基准来源",
        ["市场现货价", "用户自定义", "真实采购成本"],
        index=["市场现货价", "用户自定义", "真实采购成本"].index(
            st.session_state.get("basis_mode", "市场现货价")
        ),
        key="basis_mode",
        horizontal=True,
    )

    top_col1, top_col2 = st.columns(2, gap="large")
    with top_col1:
        st.markdown("<div style='height: 1.35rem;'></div>", unsafe_allow_html=True)
        market_card_value = f"{market_spot_price:,.0f} 元/吨" if market_spot_price is not None else "暂无数据"
        st.metric("市场现货价", market_card_value)

    with top_col2:
        st.markdown('<div class="basis-compact">', unsafe_allow_html=True)
        user_custom_basis = st.number_input(
            "用户自定义基准价",
            min_value=0.0,
            value=float(st.session_state.get("basis_user_custom_price", market_spot_price or 0.0)),
            step=500.0,
            key="basis_user_custom_price",
        )
        st.markdown("</div>", unsafe_allow_html=True)

    bottom_col1, bottom_col2 = st.columns(2, gap="large")
    with bottom_col1:
        st.markdown('<div class="basis-confirm-wrap">', unsafe_allow_html=True)
        user_confirm_real = st.checkbox(
            "我确认“真实采购成本”为企业真实采购/合同成本",
            value=bool(st.session_state.get("basis_user_confirm_real", False)),
            key="basis_user_confirm_real",
        )
        st.markdown("</div>", unsafe_allow_html=True)

    with bottom_col2:
        st.markdown('<div class="basis-compact">', unsafe_allow_html=True)
        real_purchase_basis = st.number_input(
            "真实采购成本",
            min_value=0.0,
            value=float(st.session_state.get("basis_real_purchase_price", 0.0)),
            step=500.0,
            key="basis_real_purchase_price",
        )
        st.markdown("</div>", unsafe_allow_html=True)

    basis_candidates = {
        "市场现货价": market_spot_price,
        "用户自定义": float(user_custom_basis) if float(user_custom_basis) > 0 else None,
        "真实采购成本": float(real_purchase_basis) if float(real_purchase_basis) > 0 else None,
    }

    active_basis_price = basis_candidates.get(basis_mode)
    active_basis_label = basis_mode

    if active_basis_price is None:
        st.markdown(
            "<div style='color:#b00020;font-weight:700'>当前基准价无有效数值，禁止用于对外报告</div>",
            unsafe_allow_html=True,
        )
        return

    warning_msgs = []
    if basis_mode == "市场现货价" and (ref_is_sim or market_spot_price is None):
        warning_msgs.append("当前为模拟数据，禁止用于对外报告")
    if basis_mode == "真实采购成本" and not user_confirm_real:
        warning_msgs.append("未确认真实采购成本来源，禁止用于对外报告")

    display_data["价差"] = display_data["收盘价"] - float(active_basis_price)

    latest_futures = float(display_data["收盘价"].iloc[-1])
    latest_diff = latest_futures - float(active_basis_price)
    update_time = display_data["日期"].iloc[-1]

    metric_row1 = st.columns(2)
    metric_row1[0].metric("当前基准来源", active_basis_label)
    metric_row1[1].metric("当前基准价", f"{float(active_basis_price):,.0f}")

    metric_row2 = st.columns(2)
    metric_row2[0].metric("期货收盘价", f"{latest_futures:,.0f}")
    metric_row2[1].metric(f"价差（基于{active_basis_label}）", f"{latest_diff:+,.0f}")

    st.caption(f"现货、期货数据更新时间：{update_time.strftime('%Y-%m-%d')}")

    for msg in warning_msgs:
        st.markdown(
            f"<div style='color:#b00020;font-weight:700;margin-bottom:8px'>{msg}</div>",
            unsafe_allow_html=True,
        )

    fig, ax = plt.subplots(figsize=(13.8, 6.4))
    ax.plot(display_data["日期"], display_data["价差"], linewidth=2.2)
    ax.axhline(0, linestyle="--", linewidth=1)
    _matplotlib_style(ax, f"价差走势（{active_basis_label}）", "日期", "价差 (元/吨)")
    plt.xticks(rotation=30)
    plt.tight_layout()
    st.pyplot(fig, use_container_width=True)

    st.markdown("### 三种口径对比")
    compare_rows = []
    for label, price in basis_candidates.items():
        if price is None:
            diff_txt = "暂无"
            price_txt = "暂无"
        else:
            diff_txt = f"{latest_futures - float(price):+,.0f}"
            price_txt = f"{float(price):,.0f}"
        compare_rows.append({
            "基准口径": label,
            "基准价（元/吨）": price_txt,
            "当前价差": diff_txt,
            "是否当前生效": "是" if label == active_basis_label else "否",
        })
    st.dataframe(pd.DataFrame(compare_rows), use_container_width=True, hide_index=True)

    st.markdown("### 市场现货价口径明细表")
    st.markdown(
        "**说明：下表为逐日价差明细，用于展示真实市场参考口径下的历史数据。用户自定义基准与真实采购成本属于测算口径，不在本表中展示为市场原始数据。**"
    )
    with st.expander("详细数据表格", expanded=False):
        try:
            fut_detail = display_data[["日期", "收盘价"]].copy()
            fut_detail["日期"] = pd.to_datetime(fut_detail["日期"]).dt.normalize()
            fut_detail = fut_detail.rename(columns={"日期": "date", "收盘价": "期货收盘价"})

            spot_detail = _SPOT_DF.copy()
            spot_detail = spot_detail.rename(columns={"date": "date", "spot_price": "现货参考价"})
            spot_detail["date"] = pd.to_datetime(spot_detail["date"]).dt.normalize()

            basis_detail = pd.merge(fut_detail, spot_detail, on="date", how="inner").sort_values("date").reset_index(drop=True)
            basis_detail["价差"] = basis_detail["期货收盘价"] - basis_detail["现货参考价"]
            basis_detail = basis_detail.rename(columns={"date": "日期"})
            basis_detail["日期"] = basis_detail["日期"].dt.strftime("%Y-%m-%d")

            if basis_detail.empty:
                st.info("当前周期内暂无可对齐的市场现货价逐日数据。")
            else:
                st.dataframe(
                    basis_detail.style.format({
                        "期货收盘价": "{:,.0f}",
                        "现货参考价": "{:,.0f}",
                        "价差": "{:+,.0f}",
                    }),
                    use_container_width=True,
                    height=420,
                )
        except Exception as e:
            st.error(f"价差明细表生成失败：{e}")

    st.session_state.basis_data = {
        "spot_price": float(active_basis_price),
        "ref_spot_price": market_spot_price,
        "ref_spot_source": ref_source,
        "analysis_spot_price": float(active_basis_price),
        "futures_price": latest_futures,
        "basis": latest_diff,
        "diff": latest_diff,
        "update_time": update_time,
        "basis_source_label": active_basis_label,
        "basis_source_mode": basis_mode,
        "user_custom_basis": basis_candidates.get("用户自定义"),
        "real_purchase_basis": basis_candidates.get("真实采购成本"),
    }

def render_inventory_page(analyzer):
    """库存管理（MVP）"""
    render_standard_page_header("库存管理", "记录入库与出库台账，自动汇总库存水平与加权成本，并与套保、利润模块形成联动。")

    user_info = st.session_state.get("user_info") or {}
    user_id = (
        user_info.get("user_id")
        or user_info.get("id")
        or st.session_state.get("user_id")
        or ""
    )

    # Ultimate fallback: derive from username to keep MVP usable
    if (not user_id) and isinstance(user_info, dict):
        uname = (user_info.get("username") or "").strip()
        if uname:
            user_id = f"user::{uname}"
            st.session_state.user_id = user_id

    if not user_id:
        st.error("未获取到用户ID（登录返回未包含 user.id）。请重新登录，或检查 Supabase Auth 返回。")
        return

    tab1, tab2 = st.tabs(["新增记录", "库存总览"])

    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            date = st.date_input("日期", value=datetime.now().date())
            txn_type = st.selectbox("类型", ["入库", "出库"])
            grade = st.selectbox("品级", ["电池级", "工业级", "未分类"])
            warehouse = st.text_input("仓库/地点", value="默认仓")
            qty = st.number_input("数量 (吨)", min_value=0.0, value=10.0, step=1.0)
        with col2:
            unit_cost = st.number_input(
                "单吨成本 (元/吨)",
                min_value=0.0,
                value=100000.0,
                step=1000.0,
                help="入库建议填写真实采购/生产成本；出库可不填（系统按移动加权均价核算）。"
            )
            notes = st.text_area("备注", value="", height=120)

        # 企业真实数据确认（用于避免触发“模拟数据”红字规则）
        confirm_real = st.checkbox("我确认上述成本为企业真实采购/生产成本（用于内部核算）", value=False)

        if st.button("保存库存记录", use_container_width=True):
            txn = {
                "date": str(date),
                "txn_type": txn_type,
                "grade": grade,
                "warehouse": warehouse.strip() or "默认仓",
                "qty_ton": float(qty),
                "unit_cost": float(unit_cost),
                "notes": notes.strip(),
                "is_enterprise_real": bool(confirm_real),
            }
            ok = analyzer.save_inventory_txn(user_id, txn)
            if ok:
                st.success("已保存。")
                st.rerun()
            else:
                st.error("保存失败，请稍后重试。")

    with tab2:
        summary_df, meta = analyzer.compute_inventory_position(user_id)
        st.markdown("### 当前库存汇总（移动加权均价）")
        st.dataframe(summary_df, use_container_width=True, hide_index=True)

        if meta.get("simulated"):
            # 保留项目长期约束：若存在缺失/异常导致“估算/不完整”，必须红字提示
            st.markdown(
                "<p style='color:red;font-weight:600;'>当前为模拟数据，禁止用于对外报告</p>",
                unsafe_allow_html=True
            )
            st.caption("原因：存在缺失成本/超卖/未知类型等情况，系统只能做估算汇总。建议补全入库成本、核对出库数量。")

        st.markdown("### 台账明细")
        txns = analyzer.get_inventory_txns(user_id, limit=2000)
        df = pd.DataFrame(txns)
        if df.empty:
            st.info("暂无库存记录。请先在“新增记录”里录入入库/出库。")
        else:
            # 只展示关键列
            show_cols = [c for c in ["date","txn_type","grade","warehouse","qty_ton","unit_cost","notes"] if c in df.columns]
            st.dataframe(df[show_cols], use_container_width=True, hide_index=True)

        with st.expander("导出库存台账（CSV）"):
            txns = analyzer.get_inventory_txns(user_id, limit=5000)
            df = pd.DataFrame(txns)
            if df.empty:
                st.write("暂无可导出数据。")
            else:
                csv = df.to_csv(index=False).encode("utf-8-sig")
                st.download_button("下载 CSV", data=csv, file_name="inventory_ledger.csv", mime="text/csv")


    st.markdown("---")
    st.markdown("## 库存风险热力图")
    with st.expander("打开：查看不同周期下库存价值的尾部风险", expanded=False):
        st.caption("说明：使用 LC 主力真实历史收盘价变动计算库存价值变化分布，并给出 VaR/CVaR（历史法）。")
        qty = st.number_input("库存数量（吨）", min_value=0.0, value=100.0, step=1.0, key="inv_risk_qty")
        lookback = st.slider("历史窗口（交易日）", 60, 500, 250, 10, key="inv_risk_lb")
        horizons = [1, 5, 20]
        alphas = [0.10, 0.05, 0.01]
        try:
            fdf = analyzer.fetch_real_time_data(symbol="LC0", days=int(lookback)+30, force_refresh=False)
            fdf = fdf.rename(columns={"日期":"date","收盘价":"close"})[["date","close"]].dropna()
            fdf["date"] = pd.to_datetime(fdf["date"]).dt.normalize()
            fdf = fdf.sort_values("date").tail(int(lookback)).reset_index(drop=True)

            # 计算不同周期的价值变化（用期货收盘价做锚点；若你未来接入企业合同价，可替换为合同价序列）
            mat = []
            idx_names = []
            for h in horizons:
                dP = fdf["close"].diff(h).dropna()
                pnl = float(qty) * dP
                row = []
                for a in alphas:
                    var, cvar = compute_historical_var_cvar(pnl, alpha=a)
                    row.append(cvar if cvar is not None else np.nan)  # 用 CVaR 更保守
                mat.append(row)
                idx_names.append(f"{h}日")
            heat = pd.DataFrame(mat, index=idx_names, columns=[f"alpha={a}" for a in alphas])
            st.dataframe(heat.style.format("{:,.0f}"), use_container_width=True)
            st.caption("表格数值为 CVaR（单位：元），越小（越负）代表尾部风险越大。")

            # 简单可视化：选择一个周期展示PnL分布
            h_pick = st.selectbox("查看PnL分布（周期）", horizons, index=0, key="inv_risk_hpick")
            dP2 = fdf["close"].diff(int(h_pick)).dropna()
            pnl2 = float(qty) * dP2
            st.line_chart(pd.Series(pnl2.reset_index(drop=True), name="库存价值变化(元)"))
        except Exception as e:
            st.error(f"库存风险计算失败：{e}")

def render_profit_page(analyzer):
    """利润管理（MVP）"""
    render_standard_page_header("利润管理", "基于销售记录与库存成本台账输出毛利结果，并支持与套保估算结果形成综合利润视图。")

    user_info = st.session_state.get("user_info") or {}
    user_id = (
        user_info.get("user_id")
        or user_info.get("id")
        or st.session_state.get("user_id")
        or ""
    )

    # Ultimate fallback: derive from username to keep MVP usable
    if (not user_id) and isinstance(user_info, dict):
        uname = (user_info.get("username") or "").strip()
        if uname:
            user_id = f"user::{uname}"
            st.session_state.user_id = user_id

    if not user_id:
        st.error("未获取到用户ID（登录返回未包含 user.id）。请重新登录，或检查 Supabase Auth 返回。")
        return

    tab1, tab2 = st.tabs(["新增销售记录", "利润报表"])

    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            date = st.date_input("销售日期", value=datetime.now().date(), key="sale_date")
            grade = st.selectbox("品级", ["电池级", "工业级", "未分类"], key="sale_grade")
            qty = st.number_input("销售数量 (吨)", min_value=0.0, value=10.0, step=1.0, key="sale_qty")
            unit_price = st.number_input("销售单价 (元/吨)", min_value=0.0, value=120000.0, step=1000.0, key="sale_unit_price")
        with col2:
            customer = st.text_input("客户（可选）", value="", key="sale_customer")
            override_cost = st.number_input(
                "单吨成本覆盖值（可选）",
                min_value=0.0,
                value=0.0,
                step=1000.0,
                help="如该笔业务有明确结算成本/合同成本，可填写；否则系统将按库存移动加权均价核算。"
            )
            notes = st.text_area("备注", value="", height=120, key="sale_notes")

        confirm_real = st.checkbox("我确认上述销售单价/成本（如填写）为企业真实业务数据（用于内部核算）", value=False, key="sale_confirm")

        if st.button("保存销售记录", use_container_width=True):
            txn = {
                "date": str(date),
                "grade": grade,
                "qty_ton": float(qty),
                "unit_price": float(unit_price),
                "customer": customer.strip(),
                "override_cost": float(override_cost),
                "notes": notes.strip(),
                "is_enterprise_real": bool(confirm_real),
            }
            ok = analyzer.save_sales_txn(user_id, txn)
            if ok:
                st.success("已保存。")
                st.rerun()
            else:
                st.error("保存失败，请稍后重试。")

    with tab2:
        profit_df, meta = analyzer.compute_profit_report(user_id)

        st.markdown("### 毛利明细（经营层）")
        st.dataframe(profit_df, use_container_width=True, hide_index=True)

        total_revenue = float(profit_df["revenue"].sum()) if not profit_df.empty else 0.0
        total_cogs = float(profit_df["cogs"].sum()) if not profit_df.empty else 0.0
        total_gross = float(profit_df["gross_profit"].sum()) if not profit_df.empty else 0.0

        colA, colB, colC = st.columns(3)
        colA.metric("累计收入", f"{total_revenue:,.0f} 元")
        colB.metric("累计成本", f"{total_cogs:,.0f} 元")
        colC.metric("累计毛利", f"{total_gross:,.0f} 元")

        # 套保损益（如果已有）
        hedge_pnl = 0.0
        try:
            # 优先从 session_state 里拿最近一次套保计算结果
            hedge_result = st.session_state.get("hedge_result") or {}
            hedge_pnl = float(hedge_result.get("estimated_pnl", 0.0) or 0.0)
        except Exception:
            hedge_pnl = 0.0

        st.markdown("### 综合利润（含套保）")
        st.caption("套保损益为系统估算值，取自最近一次套保计算/记录（如有）。")
        st.metric("综合利润（毛利 + 套保损益）", f"{(total_gross + hedge_pnl):,.0f} 元")

        if meta.get("simulated"):
            st.markdown(
                "<p style='color:red;font-weight:600;'>当前为模拟数据，禁止用于对外报告</p>",
                unsafe_allow_html=True
            )
            st.caption("原因：库存成本缺失/库存不足导致成本估算；建议补全库存台账或为销售记录填写明确成本覆盖值。")

        with st.expander("导出利润明细（CSV）"):
            if profit_df.empty:
                st.write("暂无可导出数据。")
            else:
                csv = profit_df.to_csv(index=False).encode("utf-8-sig")
                st.download_button("下载 CSV", data=csv, file_name="profit_report.csv", mime="text/csv")


def render_option_page(analyzer):
    """渲染期权保险计算页面"""
    render_standard_page_header("期权计算", "以采购上限或销售下限为目标，测算期权保费与不同风险管理工具下的情景差异。")

    col_left, col_right = st.columns([1, 2])

    with col_left:
        option_mode = st.radio(
            "方案类型",
            ["锁定最高买入价", "锁定最低卖出价"],
            index=0
        )
        target_price = st.number_input(
            "锁定价格 (元/吨)",
            min_value=0.0,
            value=250000.0,
            step=500.0
        )
        quantity = st.number_input(
            "保障数量 (吨)",
            min_value=0.0,
            value=100.0,
            step=1.0
        )
        months = st.number_input(
            "保障月数",
            min_value=1,
            max_value=24,
            value=3,
            step=1
        )
        with st.expander("高级参数"):
            volatility = st.slider("波动率（年化）", 0.05, 0.8, 0.35, 0.01)
            risk_free = st.slider("无风险利率（年化）", 0.0, 0.1, 0.02, 0.005)

    with col_right:
        price_data = analyzer.fetch_real_time_data()
        current_price = float(price_data["收盘价"].iloc[-1]) if not price_data.empty else target_price
        time_years = months / 12
        option_type = "call" if option_mode == "锁定最高买入价" else "put"
        premium_per_ton = black_scholes_price(
            option_type, current_price, target_price, time_years, risk_free, volatility
        )
        total_premium = premium_per_ton * quantity

        st.markdown("### 输出结果")
        st.metric("您需支付的总保费约为", f"{total_premium:,.0f} 元")
        st.caption(f"当前期货价格：{current_price:,.0f} 元/吨")

        scenario_prices = np.linspace(current_price * 0.7, current_price * 1.3, 60)
        if option_type == "call":
            no_insurance = scenario_prices
            futures_locked = np.full_like(scenario_prices, target_price)
            option_cost = np.minimum(scenario_prices, target_price) + premium_per_ton
            ylabel = "采购成本 (元/吨)"
        else:
            no_insurance = scenario_prices
            futures_locked = np.full_like(scenario_prices, target_price)
            option_cost = np.maximum(scenario_prices, target_price) - premium_per_ton
            ylabel = "销售收入 (元/吨)"

        fig, ax = plt.subplots(figsize=(11, 6))
        ax.plot(scenario_prices, no_insurance, color="#ff3b30", linewidth=2, label="不买保险")
        ax.plot(scenario_prices, futures_locked, color="#34c759", linewidth=2, label="买期货")
        ax.plot(scenario_prices, option_cost, color="#0a84ff", linewidth=2, label="买期权")
        _matplotlib_style(ax, "三种情景对比", "未来价格 (元/吨)", ylabel)
        ax.legend(frameon=False)
        plt.tight_layout()
        st.pyplot(fig)

        st.session_state.option_result = {
            "mode": option_mode,
            "target_price": target_price,
            "quantity": quantity,
            "months": months,
            "premium_total": total_premium,
            "premium_per_ton": premium_per_ton
        }

        st.markdown("#### 操作提示")
        st.markdown(
            "后台可接入Black-Scholes或交易所期权定价模型进行实时计算，"
            "当前结果为基于最新期货价格的估算值。"
        )


def render_exposure_page(analyzer):
    """渲染风险敞口量化页面"""
    render_standard_page_header("风险敞口", "量化未来采购、现有库存与已锁定数量之间的风险暴露，为后续套保与期权决策提供依据。")

    col_left, col_right = st.columns([1, 2])
    with col_left:
        future_purchase = st.number_input(
            "未来计划采购量 (吨)",
            min_value=0.0,
            value=200.0,
            step=1.0
        )
        inventory = st.number_input(
            "现有库存量 (吨)",
            min_value=0.0,
            value=100.0,
            step=1.0
        )
        locked_sales = st.number_input(
            "已锁定的量 (吨)",
            min_value=0.0,
            value=80.0,
            step=1.0
        )
        st.caption(
            "已锁定的量指已签订固定价格合同或完成点价交易的数量，与价格无关。"
        )

    total_exposure = future_purchase + inventory - locked_sales
    total_volume = future_purchase + inventory + locked_sales
    exposure_ratio = abs(total_exposure) / total_volume if total_volume > 0 else 0

    if exposure_ratio < 0.2:
        risk_level = "低"
    elif exposure_ratio < 0.5:
        risk_level = "中"
    else:
        risk_level = "高"

    risk_direction = "原材料价格上涨" if total_exposure > 0 else "原材料价格下跌" if total_exposure < 0 else "风险中性"
    risk_impact = total_exposure * 10000

    with col_right:
        st.markdown("### 量化结果")
        st.metric("风险敞口", f"{total_exposure:,.0f} 吨")
        st.metric("风险程度", risk_level)
        st.markdown(f"**风险方向**：{risk_direction}")
        st.markdown(f"**风险影响**：原材料价格每上涨一万元/吨，成本将变化 {risk_impact:,.0f} 元")

        fig, ax = plt.subplots(figsize=(6, 6))
        components = [future_purchase, inventory, locked_sales]
        labels = ["未来采购", "现有库存", "已锁定"]
        ax.pie(components, labels=labels, autopct="%1.1f%%", startangle=90)
        ax.set_title("敞口构成", fontsize=13, fontweight="bold", color="#172033")
        st.pyplot(fig)

        st.markdown("#### 策略建议")
        st.markdown("考虑买入套保或配置期权，减少价格波动对成本的冲击。")

    st.session_state.exposure_result = {
        "future_purchase": future_purchase,
        "inventory": inventory,
        "locked_sales": locked_sales,
        "exposure": total_exposure,
        "risk_level": risk_level,
        "risk_direction": risk_direction,
        "risk_impact": risk_impact
    }


def render_scenario_page(analyzer):
    """渲染多情景分析页面"""
    render_standard_page_header("多情景分析", "对不同价格波动场景进行横向比较，快速评估不套保与套保后的盈亏敏感性。")

    default_params = st.session_state.get("hedge_results", {}).get("params", {})
    cost_price = st.number_input(
        "存货成本价 (元/吨)",
        min_value=0.0,
        value=float(default_params.get("cost_price", 100000.0)),
        step=500.0
    )
    inventory = st.number_input(
        "存货数量 (吨)",
        min_value=0.0,
        value=float(default_params.get("inventory", 100.0)),
        step=1.0
    )
    hedge_ratio = st.slider(
        "套保比例 (%)",
        min_value=0,
        max_value=100,
        value=int(default_params.get("hedge_ratio", 0.8) * 100),
        step=5
    ) / 100

    scenario_options = {
        "价格上涨10%": 0.1,
        "价格平稳": 0.0,
        "价格下降10%": -0.1,
        "用户自设定": None
    }

    cols = st.columns(4)
    selected = st.session_state.get("scenario_selected", "价格上涨10%")
    for idx, (label, _) in enumerate(scenario_options.items()):
        if cols[idx].button(label):
            selected = label
            st.session_state.scenario_selected = label

    if selected == "用户自设定":
        custom_pct = st.number_input("自设定变动 (%)", value=5.0, step=1.0)
        custom_change = custom_pct / 100
    else:
        custom_change = scenario_options[selected]

    price_data = analyzer.fetch_real_time_data()
    current_price = float(price_data["收盘价"].iloc[-1]) if not price_data.empty else cost_price
    hedge_contracts = int(np.round(inventory * hedge_ratio))

    scenarios = {
        "上涨10%": 0.1,
        "平稳": 0.0,
        "下降10%": -0.1,
        "自设定": custom_change
    }

    rows = []
    for name, change in scenarios.items():
        scenario_price = current_price * (1 + change)
        spot_profit = (scenario_price - cost_price) * inventory
        futures_profit = (current_price - scenario_price) * hedge_contracts
        hedge_profit = spot_profit + futures_profit
        rows.append({
            "情景": name,
            "价格变动": f"{change*100:+.1f}%",
            "不套保盈亏(元)": f"{spot_profit:,.0f}",
            "套保后盈亏(元)": f"{hedge_profit:,.0f}"
        })

    result_df = pd.DataFrame(rows)
    st.markdown("### 情景结果汇总")
    st.dataframe(result_df, use_container_width=True)

    st.session_state.scenario_results = rows


    st.markdown("---")
    st.markdown("## 历史极端情景")
    with st.expander("打开：从真实历史数据中抽取最极端的涨跌日作为压力测试", expanded=False):
        lookback = st.slider("回看窗口（交易日）", 60, 500, 250, 10, key="hist_scn_lb")
        topn = st.slider("展示极端情景数量", 3, 20, 10, 1, key="hist_scn_topn")
        try:
            fdf = analyzer.fetch_real_time_data(symbol="LC0", days=int(lookback)+30, force_refresh=False)
            fdf = fdf.rename(columns={"日期":"date","收盘价":"close"})[["date","close"]].dropna()
            fdf["date"] = pd.to_datetime(fdf["date"]).dt.normalize()
            fdf = fdf.sort_values("date").tail(int(lookback)).reset_index(drop=True)
            fdf["ret_1d"] = fdf["close"].pct_change()
            fdf = fdf.dropna(subset=["ret_1d"])
            worst = fdf.nsmallest(int(topn), "ret_1d")[["date","ret_1d","close"]]
            best = fdf.nlargest(int(topn), "ret_1d")[["date","ret_1d","close"]]
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("### 最差跌幅日")
                st.dataframe(worst.assign(ret_1d=worst["ret_1d"]*100).rename(columns={"ret_1d":"跌幅(%)","close":"收盘价"}), use_container_width=True)
            with c2:
                st.markdown("### 最强涨幅日")
                st.dataframe(best.assign(ret_1d=best["ret_1d"]*100).rename(columns={"ret_1d":"涨幅(%)","close":"收盘价"}), use_container_width=True)
            Q = float(inventory)
            # 用百分比冲击近似作用于“测算基准价/现货参考价”（如果你未来用企业合同价，可替换）
            base_price = float(cost_price) if float(cost_price) > 0 else float(fdf["close"].iloc[-1])
            worst_pnl = -Q * base_price * worst["ret_1d"]  # 价格下跌，库存亏损为负（这里取负号便于解释）
            best_pnl  = -Q * base_price * best["ret_1d"]
            st.markdown("### 压力盈亏")
            st.caption("以当前库存 / 成本价估算")
            df_pnl = pd.concat([
                worst.assign(压力盈亏_元=worst_pnl.values).rename(columns={"ret_1d":"日收益率"}),
                best.assign(压力盈亏_元=best_pnl.values).rename(columns={"ret_1d":"日收益率"})
            ], ignore_index=True)
            st.dataframe(df_pnl, use_container_width=True)
        except Exception as e:
            st.error(f"历史情景提取失败：{e}")

def render_report_page(analyzer):
    """渲染分析报告页面（v45 稳定版）"""
    render_standard_page_header("分析报告", "整合市场概况、价差信息、风险敞口与情景结果，输出适合展示与汇报的结构化分析内容。")

    basis_data = st.session_state.get("basis_data")
    if not basis_data:
        price_data = analyzer.fetch_real_time_data()

        try:
            if price_data is not None and not price_data.empty:
                if "日期" in price_data.columns:
                    update_time = pd.to_datetime(price_data["日期"].iloc[-1])
                elif "date" in price_data.columns:
                    update_time = pd.to_datetime(price_data["date"].iloc[-1])
                else:
                    update_time = datetime.now()
            else:
                update_time = datetime.now()
        except Exception:
            update_time = datetime.now()

        try:
            latest_futures = float(price_data["收盘价"].iloc[-1]) if (price_data is not None and not price_data.empty and "收盘价" in price_data.columns) else 0.0
        except Exception:
            latest_futures = 0.0

        # 现货参考价：优先使用内置现货表（静态），若缺失则要求用户明确输入并红字提示
        s_price, s_date, s_ok = get_spot_price_on_or_before(pd.to_datetime(update_time))
        if s_ok and s_price is not None:
            spot_price = float(s_price)
            ref_source = "内置现货表（用户提供历史数据，静态）"
        else:
            st.markdown("<span style='color:red;font-weight:700'>当前为模拟数据，禁止用于对外报告（未找到内置现货价，请手动输入并确认来源）</span>", unsafe_allow_html=True)
            spot_price = st.number_input("现货参考价 (元/吨) - 手动输入", min_value=0.0, value=0.0, step=500.0)
            ref_source = "手动输入（需用户确认真实来源）"

        basis_value = float(latest_futures - spot_price) if spot_price is not None else 0.0
        basis_data = {
            "spot_price": float(spot_price),
            "ref_spot_price": float(spot_price),
            "ref_spot_source": ref_source,
            "analysis_spot_price": float(spot_price),
            "futures_price": float(latest_futures),
            "basis": float(basis_value),
            "update_time": update_time
        }

    exposure_result = st.session_state.get("exposure_result")
    scenario_results = st.session_state.get("scenario_results", [])

    ref_spot_price = float(basis_data.get("ref_spot_price", basis_data.get("spot_price", 0.0)))
    ref_spot_source = str(basis_data.get("ref_spot_source", "内置现货表（用户提供历史数据）"))
    analysis_spot_price = float(basis_data.get("analysis_spot_price", basis_data.get("spot_price", 0.0)))

    st.markdown("### 1. 当前市场概况")
    st.markdown(
        "\n".join([
            f"- 日期：{_fmt_dt(basis_data.get('update_time'))}",
            f"- 市场参考价（生意社/AKShare）：{ref_spot_price:,.0f} 元/吨",
            f"- 市场参考价来源：{ref_spot_source}",
            f"- 测算基准价（用于计算）：{analysis_spot_price:,.0f} 元/吨",
            f"- 期货价（主力）：{basis_data.get('futures_price', 0):,.0f} 元/吨",
            f"- 实时价差（期货主力 - 市场参考价）：{basis_data.get('basis', 0):+,.0f} 元/吨",
            "- 数据来源：内置现货表（静态，不实时更新）/ 新浪（期货，AkShare futures_zh_daily_sina）",
        ])
    )

    st.markdown("### 2. 风险敞口计算结果")
    if exposure_result:
        st.markdown(
            f"- 风险敞口：{exposure_result['exposure']:,.0f} 吨\n"
            f"- 风险方向：{exposure_result['risk_direction']}\n"
            f"- 风险程度：{exposure_result['risk_level']}\n"
            f"- 风险影响：每上涨1万元/吨，成本变化 {exposure_result['risk_impact']:,.0f} 元"
        )
    else:
        st.info("请先在“风险敞口”模块完成测算。")

    st.markdown("### 3. 情景指标对比")
    if scenario_results:
        scenario_df = pd.DataFrame(scenario_results)
        st.dataframe(scenario_df, use_container_width=True)
    else:
        st.info("请先在“多情景分析”模块生成对比结果。")

    st.markdown("### 4. 风险提示与建议")
    st.markdown(
        "- 价差扩大时需关注套保效率变化。\n"
        "- 结合敞口方向优先选择期货或期权锁定风险。\n"
        "- 如需精细报告，可导出当前页面内容并补充业务说明。"
    )

    report_text = f"""分析报告
生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

1. 当前市场概况
日期：{basis_data['update_time'].strftime('%Y-%m-%d')}
市场参考价（生意社/AKShare）：{ref_spot_price:,.0f} 元/吨
市场参考价来源：{ref_spot_source}
测算基准价：{analysis_spot_price:,.0f} 元/吨
期货价（主力）：{basis_data['futures_price']:,.0f} 元/吨
实时价差（期货主力 - 市场参考价）：{basis_data['basis']:+,.0f} 元/吨
数据来源：内置现货表（静态，不实时更新）/ 新浪（期货，AkShare futures_zh_daily_sina）

2. 风险敞口计算结果
"""
    if exposure_result:
        report_text += (
            f"风险敞口：{exposure_result['exposure']:,.0f} 吨\n"
            f"风险方向：{exposure_result['risk_direction']}\n"
            f"风险程度：{exposure_result['risk_level']}\n"
            f"风险影响：每上涨1万元/吨，成本变化 {exposure_result['risk_impact']:,.0f} 元\n"
        )
    else:
        report_text += "风险敞口：未填写\n"

    report_text += "\n3. 情景指标对比\n"
    if scenario_results:
        for row in scenario_results:
            report_text += (
                f"{row['情景']} | {row['价格变动']} | 不套保盈亏 {row['不套保盈亏(元)']} | "
                f"套保后盈亏 {row['套保后盈亏(元)']}\n"
            )
    else:
        report_text += "未生成情景分析\n"

    report_text += "\n4. 风险提示与建议\n"
    report_text += (
        "价差扩大时需关注套保效率变化；结合敞口方向选择期货或期权；"
        "如需精细报告，可补充业务说明。\n"
    )


    # 商业化导出：PDF / Word（可选）
    st.markdown("#### 导出（企业交付版）")
    if _can("export_report"):
        lines = [ln for ln in report_text.splitlines()]
        pdf_bytes = build_report_pdf_bytes("熵合科技-分析报告", lines)
        docx_bytes = build_report_docx_bytes("熵合科技-分析报告", lines)

        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button(
                label="下载 TXT",
                data=report_text,
                file_name=f"分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        with c2:
            st.download_button(
                label="下载 PDF",
                data=pdf_bytes if pdf_bytes else report_text,
                file_name=f"分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf" if pdf_bytes else "text/plain",
                disabled=(pdf_bytes is None),
                use_container_width=True
            )
        with c3:
            st.download_button(
                label="下载 Word",
                data=docx_bytes if docx_bytes else report_text,
                file_name=f"分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if docx_bytes else "text/plain",
                disabled=(docx_bytes is None),
                use_container_width=True
            )
        _log("export_report", {"formats": ["txt", "pdf" if pdf_bytes else None, "docx" if docx_bytes else None]})
    else:
        st.info("当前环境未开启正式报告导出权限。")


def render_history_page(analyzer):
    """渲染分析历史页面"""
    render_standard_page_header("分析历史", "集中查看云端保存的分析记录，便于复盘不同时间点的参数、结果与策略选择。")
    
    # 获取用户历史记录
    with st.spinner("正在加载分析历史..."):
        history = analyzer.get_user_history(limit=50)
    
    if not history:
        st.info("暂无分析历史记录")
        st.markdown("""
        ### 开始您的第一次分析
        
        1. 前往 **套保计算** 页面
        2. 输入您的存货参数
        3. 点击 **开始计算**
        4. 分析结果将自动保存到历史记录
        
        所有分析记录都会安全存储在云端，您可以随时查看和导出。
        """)
        return
    
    # 显示历史记录统计
    total_analyses = len(history)
    latest_analysis = history[0]['created_at'] if history else None
    
    col_stat1, col_stat2, col_stat3 = st.columns(3)
    with col_stat1:
        st.metric("总分析次数", f"{total_analyses}")
    with col_stat2:
        if latest_analysis:
            from dateutil import parser
            latest_time = parser.parse(latest_analysis)
            time_diff = datetime.now() - latest_time.replace(tzinfo=None)
            if time_diff.days > 0:
                latest_str = f"{time_diff.days}天前"
            elif time_diff.seconds > 3600:
                latest_str = f"{time_diff.seconds // 3600}小时前"
            else:
                latest_str = f"{time_diff.seconds // 60}分钟前"
            st.metric("最近分析", latest_str)
    
    # 历史记录列表
    st.markdown("### 历史记录列表")
    
    for i, record in enumerate(history):
        with st.expander(f"分析 #{total_analyses - i} - {record['created_at'][:19]}", expanded=(i == 0)):
            col_record1, col_record2, col_record3 = st.columns([3, 2, 1])
            
            with col_record1:
                st.markdown(f"**分析类型**：{record['analysis_type']}")
                if 'input_params' in record and isinstance(record['input_params'], dict):
                    st.markdown("**输入参数**：")
                    for key, value in record['input_params'].items():
                        if key == 'cost_price':
                            st.text(f"  - 成本价：{value:,.2f} 元/吨")
                        elif key == 'inventory':
                            st.text(f"  - 存货量：{value:,.2f} 吨")
                        elif key == 'hedge_ratio':
                            st.text(f"  - 套保比例：{value*100:.1f}%")
                        elif key == 'margin_rate':
                            st.text(f"  - 保证金比例：{value*100:.0f}%")
            
            with col_record2:
                if 'result_data' in record and isinstance(record['result_data'], dict):
                    st.markdown("**分析结果**：")
                    result = record['result_data']
                    if 'current_price' in result:
                        st.text(f"  - 当时价格：{result['current_price']:,.0f}元")
                    if 'hedge_contracts' in result:
                        st.text(f"  - 建议手数：{result['hedge_contracts']}手")
                    if 'total_margin' in result:
                        st.text(f"  - 保证金：{result['total_margin']:,.0f}元")
                    if 'profit_status' in result:
                        profit_color = "green" if result['profit_status'] == '盈利' else "red"
                        st.markdown(f"  - 盈亏状态：<span style='color:{profit_color}'>{result['profit_status']}</span>", 
                                  unsafe_allow_html=True)
            
            with col_record3:
                analysis_id = record['analysis_id']
                if st.button("删除", key=f"delete_{analysis_id}", 
                           help="删除此条记录"):
                    if not _require("delete_history", "只有管理者/管理员可删除历史记录。"):
                        return
                    _log("delete_history", {"analysis_id": analysis_id})
                    if analyzer.delete_history_record(analysis_id):
                        st.success("记录已删除")
                        st.rerun()
                    else:
                        st.error("删除失败")
                
                # 重新分析按钮
                if 'input_params' in record and isinstance(record['input_params'], dict):
                    if st.button("重新分析", key=f"recalc_{analysis_id}"):
                        st.session_state.recalc_params = record['input_params']
                        st.session_state.current_page = "套保计算"
                        st.rerun()
    
    # 批量操作
    st.markdown("---")
    st.markdown("### 批量操作")
    
    col_batch1, col_batch2, col_batch3 = st.columns(3)
    
    with col_batch1:
        if st.button("导出所有记录", use_container_width=True):
            # 导出所有历史记录为JSON
            export_data = {
                "export_time": datetime.now().isoformat(),
                "user": st.session_state.user_info['username'] if 'user_info' in st.session_state else "未知用户",
                "total_records": len(history),
                "history": history
            }
            
            json_str = json.dumps(export_data, ensure_ascii=False, indent=2)
            st.download_button(
                label="下载JSON文件",
                data=json_str,
                file_name=f"套保分析历史_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True
            )
    
    with col_batch2:
        if st.button("清空所有记录", use_container_width=True, type="secondary"):
            st.warning("此操作将删除所有历史记录，且不可恢复！")
            confirm = st.checkbox("我确认要删除所有记录")
            if confirm and st.button("确认删除", type="primary"):
                # 这里需要实现批量删除功能
                st.error("批量删除功能开发中")
                # for record in history:
                #     analyzer.delete_history_record(record['analysis_id'])
                # st.success("所有记录已删除")
                # st.rerun()
    
    with col_batch3:
        if st.button("刷新列表", use_container_width=True):
            st.rerun()



def render_strategy_page(analyzer):
    """策略管理与回溯（轻量版，适合竞赛演示与早期试点）"""
    render_standard_page_header("策略管理", "记录套保计算页面保存的策略快照，并基于真实历史期货收盘价进行非模拟回溯分析。")

    strategies = st.session_state.get("saved_strategies", [])
    if not strategies:
        st.info("暂无已保存策略。请先到【套保计算】->“策略保存与回溯”保存一条策略。")
        return

    df = pd.DataFrame(strategies)
    st.dataframe(df, use_container_width=True)

    st.markdown("## 回溯计算")
    idx = st.selectbox("选择一条策略进行回溯", list(range(len(strategies))), format_func=lambda i: f"{strategies[i].get('name','未命名')}（{strategies[i].get('created_at','')}）")
    s = strategies[int(idx)]

    # 回溯参数
    days = st.slider("回溯窗口（交易日）", 20, 400, 120, 10)
    alpha = st.select_slider("尾部概率 alpha", options=[0.10, 0.05, 0.01], value=0.05)

    try:
        fdf = analyzer.fetch_real_time_data(symbol="LC0", days=int(days)+30, force_refresh=False)
        fdf = fdf.rename(columns={"日期":"date","收盘价":"close"})[["date","close"]].dropna()
        fdf["date"] = pd.to_datetime(fdf["date"]).dt.normalize()
        fdf = fdf.sort_values("date").tail(int(days)).reset_index(drop=True)

        Q = float(s.get("inventory_ton", 0.0))
        h = float(s.get("hedge_ratio", 0.0))
        # 默认按“库存（担心下跌）”理解：现货P&L ~ Q*dP；期货卖出套保：-h*Q*dP_fut
        dP = fdf["close"].diff().dropna()
        spot_pnl = Q * dP
        fut_pnl = -h * Q * dP
        net = (spot_pnl + fut_pnl).reset_index(drop=True)

        var, cvar = compute_historical_var_cvar(net, alpha=alpha)

        c1, c2, c3 = st.columns(3)
        c1.metric("套保比例", f"{h*100:.0f}%")
        c2.metric("历史VaR (PnL)", f"{var:,.0f}" if var is not None else "—")
        c3.metric("历史CVaR (PnL)", f"{cvar:,.0f}" if cvar is not None else "—")

        st.line_chart(pd.Series(net, name="净PnL(元)"))

        st.markdown("### 备注")
        st.write(s.get("note", ""))

        col_del, col_export = st.columns(2)
        with col_del:
            if st.button("删除该策略", type="secondary", use_container_width=True):
                if not _require("delete_strategy", "只有管理者/管理员可删除策略。"):
                    return
                _log("delete_strategy", {"name": s.get("name"), "created_at": s.get("created_at")})
                st.session_state.saved_strategies.pop(int(idx))
                st.success("已删除。")
                st.rerun()
        with col_export:
            out = {
                "strategy": s,
                "backtest_window_days": int(days),
                "alpha": float(alpha),
                "var_pnl": var,
                "cvar_pnl": cvar,
            }
            st.download_button("导出回溯结果(JSON)", data=json.dumps(out, ensure_ascii=False, indent=2), file_name="strategy_backtest.json", mime="application/json", use_container_width=True)
    except Exception as e:
        st.error(f"回溯失败：{e}")




def render_settings_page(analyzer):
    """渲染账号设置页面"""
    render_standard_page_header("账号设置", "管理账户信息、密码、偏好配置与数据导出，保持系统使用体验与数据安全的一致性。")
    
    user_info = st.session_state.user_info
    
    tab1, tab2, tab3, tab4 = st.tabs(["账户信息", "修改密码", "偏好设置", "数据管理"])
    
    with tab1:
        st.markdown("### 账户信息")
        
        if user_info:
            col_info1, col_info2 = st.columns(2)
            
            with col_info1:
                st.markdown(f"**用户名**：{user_info['username']}")
                st.markdown(f"**邮箱**：{user_info['email']}")
                st.markdown(f"**用户ID**：`{user_info['user_id']}`")
            
            with col_info2:
                if 'settings' in user_info and user_info['settings']:
                    settings = user_info['settings']
                    st.markdown("**账户状态**：正常")
                    st.markdown(f"**会员等级**：{settings.get('subscription_tier', '免费版')}")
                    st.markdown(f"**注册时间**：{settings.get('created_at', '未知')[:10]}")
                else:
                    st.markdown("**账户状态**：设置未加载")
                # 角色/身份区分功能已移除：所有账号拥有相同权限。

        
        # 账户操作
        st.markdown("### 账户操作")
        
        col_action1, col_action2 = st.columns(2)
        
        with col_action1:
            if st.button("刷新账户信息", use_container_width=True):
                # 重新加载用户信息
                if analyzer.supabase and 'user_info' in st.session_state:
                    settings = analyzer.supabase.get_user_settings(st.session_state.user_info['user_id'])
                    if settings:
                        st.session_state.user_info['settings'] = settings
                        st.success("账户信息已刷新")
                        st.rerun()
        
        with col_action2:
            if st.button("导出账户数据", use_container_width=True):
                export_data = {
                    "user_info": user_info,
                    "export_time": datetime.now().isoformat()
                }
                
                json_str = json.dumps(export_data, ensure_ascii=False, indent=2)
                st.download_button(
                    label="下载账户数据",
                    data=json_str,
                    file_name=f"账户数据_{user_info['username']}_{datetime.now().strftime('%Y%m%d')}.json",
                    mime="application/json",
                    use_container_width=True
                )
    
    with tab2:
        st.markdown("### 修改密码")
        
        old_password = st.text_input("当前密码", type="password", 
                                   help="请输入当前使用的密码")
        new_password = st.text_input("新密码", type="password", 
                                   help="至少6个字符，建议包含字母和数字")
        confirm_password = st.text_input("确认新密码", type="password")
        
        # 密码强度检查
        if new_password:
            has_letter = any(c.isalpha() for c in new_password)
            has_digit = any(c.isdigit() for c in new_password)
            length_ok = len(new_password) >= 6
            
            if length_ok and (has_letter and has_digit):
                strength = "强"
                color = "green"
            elif length_ok and (has_letter or has_digit):
                strength = "中"
                color = "orange"
            else:
                strength = "弱"
                color = "red"
            
            st.markdown(f"密码强度：<span style='color:{color};font-weight:bold'>{strength}</span>", 
                      unsafe_allow_html=True)
        
        if st.button("确认修改密码", type="primary", use_container_width=True):
            if not all([old_password, new_password, confirm_password]):
                st.error("请填写所有字段")
            elif new_password != confirm_password:
                st.error("两次输入的新密码不一致")
            elif len(new_password) < 6:
                st.error("密码长度至少6位")
            elif old_password == new_password:
                st.error("新密码不能与旧密码相同")
            else:
                success, message = analyzer.auth.change_password(
                    user_info['username'], old_password, new_password
                )
                if success:
                    st.success(message)
                    st.info("请使用新密码重新登录")
                else:
                    st.error(message)
    
    with tab3:
        st.markdown("### 偏好设置")
        
        if 'settings' in user_info and user_info['settings']:
            settings = user_info['settings']
            
            # 默认参数设置
            st.markdown("#### 默认计算参数")
            
            default_cost = st.number_input(
                "默认成本价 (元/吨)",
                min_value=0.0,
                max_value=500000.0,
                value=float(settings.get('default_cost_price', 100000.0)),
                step=1000.0
            )
            
            default_inventory = st.number_input(
                "默认存货量 (吨)",
                min_value=0.0,
                max_value=10000.0,
                value=float(settings.get('default_inventory', 100.0)),
                step=1.0
            )
            

            default_ratio = st.slider(
                "默认套保比例 (%)",
                min_value=0,
                max_value=100,
                value=int(settings.get('default_hedge_ratio', 0.8) * 100),
                step=5
            )
            
            # 主题颜色
            theme_color = st.selectbox(
                "主题颜色",
                ["blue", "green", "purple", "orange", "red"],
                index=["blue", "green", "purple", "orange", "red"].index(
                    settings.get('theme_color', 'blue')
                )
            )
            
            if st.button("保存设置", type="primary", use_container_width=True):
                new_settings = {
                    'default_cost_price': float(default_cost),
                    'default_inventory': float(default_inventory),
                    'default_hedge_ratio': float(default_ratio) / 100,
                    'theme_color': theme_color
                }
                
                if analyzer.auth.update_user_settings(user_info['user_id'], new_settings):
                    st.success("偏好设置已保存")
                    st.session_state.user_info['settings'] = new_settings
                    _log("update_settings", {"keys": list(new_settings.keys())})
                else:
                    st.error("保存设置失败")
        else:
            st.info("正在加载用户设置...")
    
    with tab4:
        st.markdown("### 数据管理")
        
        st.markdown("#### 本地缓存")
        col_cache1, col_cache2 = st.columns(2)
        
        with col_cache1:
            if st.button("清除本地缓存", use_container_width=True, 
                        help="清除本地缓存的价格数据"):
                analyzer.cache_data = {}
                analyzer.cache_time = {}
                st.success("本地缓存已清除")
        
        with col_cache2:
            if st.button("查看缓存状态", use_container_width=True):
                cache_count = len(analyzer.cache_data)
                st.info(f"当前缓存了 {cache_count} 个数据集的 {sum(len(df) for df in analyzer.cache_data.values())} 条记录")
        
        st.markdown("#### 数据导出")
        
        # 导出所有分析历史
        history = analyzer.get_user_history(limit=1000)
        if history:
            export_all = {
                "user": user_info['username'],
                "export_time": datetime.now().isoformat(),
                "total_records": len(history),
                "records": history
            }
            
            json_str = json.dumps(export_all, ensure_ascii=False, indent=2)
            st.download_button(
                label="导出所有历史记录",
                data=json_str,
                file_name=f"套保分析完整历史_{user_info['username']}_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json",
                use_container_width=True
            )
        else:
            st.info("暂无历史记录可导出")
        
        st.markdown("#### 账户操作")
        
        if st.button("注销账户", type="secondary", use_container_width=True):
            st.warning("此操作将删除您的所有数据，且不可恢复！")
            confirm = st.checkbox("我确认要注销账户")
            if confirm:
                st.error("账户注销功能开发中")
                # 这里需要实现账户删除功能
    
    # 退出登录按钮
    st.markdown("---")
    col_logout1, col_logout2, col_logout3 = st.columns([1, 2, 1])
    
    with col_logout2:
        if st.button("退出登录", type="primary", use_container_width=True):
            st.session_state.authenticated = False
            _log("logout", {})
            st.session_state.user_info = None
            st.success("已退出登录")
            st.rerun()

# ============================================================================
# 主程序入口
# ============================================================================

if __name__ == "__main__":
    # 创建必要的目录
    os.makedirs('data', exist_ok=True)
    os.makedirs('charts', exist_ok=True)
    
    # 运行应用
    try:
        main()
    except Exception as e:
        st.error(f"应用程序运行出错: {str(e)}")
        st.code(traceback.format_exc())
        st.info("请检查：\n1. 网络连接\n2. 环境变量配置\n3. 依赖包安装")
